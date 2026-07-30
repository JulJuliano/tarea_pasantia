import os
import subprocess
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ================================================================
#  FUNCIONES DE CONFIGURACIÓN Y FORMATO (ESTÁTICAS)
# ================================================================

def setup_page(doc):
    """Márgenes Art. 6 (4cm Izquierdo para encuadernación, 3cm los demás)"""
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

def add_paragraph_normado(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)

    run_l = p.add_run(label)
    run_l.font.name = 'Times New Roman'
    run_l.font.size = Pt(12)
    run_l.font.bold = True

    run_v = p.add_run(value)
    run_v.font.name = 'Times New Roman'
    run_v.font.size = Pt(12)
    return p

def set_cell_format(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        borders.append(border)
    tblPr.append(borders)

# ================================================================
#  PROCESADOR CORE: COMPILA UNA SEMANA ESPECÍFICA
# ================================================================

def generar_documento_semana(datos):
    doc = Document()
    setup_page(doc)

    # 1. Título
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_after = Pt(12)
    run_t = p_titulo.add_run("PLAN SEMANAL DE PASANTÍAS")
    run_t.font.name = 'Times New Roman'
    run_t.font.size = Pt(14)
    run_t.font.bold = True

    # 2. Datos Informativos
    add_label_value(doc, "EMPRESA: ", "Empresa Mixta Petrolera Venangocupet, S.A.")
    add_label_value(doc, "DEPARTAMENTO: ", "Presidencia / Alta Dirección")
    add_label_value(doc, "ESPECIALIDAD: ", "Informática")
    add_label_value(doc, "SEMANA N°: ", datos["num_semana"])
    add_label_value(doc, "PERÍODO: ", datos["periodo"])
    add_label_value(doc, "PASANTE: ", "Juliano Cardona  |  C.I.: 32.281.199")
    add_label_value(doc, "TUTOR INDUSTRIAL: ", "Ing. Yasmin Sabaneta  |  C.I.: 14.187.924")

    doc.add_paragraph()

    # 3. Objetivo de la semana
    p_obj_t = doc.add_paragraph()
    p_obj_t.paragraph_format.space_after = Pt(4)
    run_ot = p_obj_t.add_run("OBJETIVO DE LA SEMANA:")
    run_ot.font.name = 'Times New Roman'
    run_ot.font.size = Pt(12)
    run_ot.font.bold = True

    add_paragraph_normado(doc, datos["objetivo"])

    doc.add_paragraph()

    # 4. Tabla de Actividades Planificadas (Diseño Fijo Estricto)
    p_cuarto_t = doc.add_paragraph()
    p_cuarto_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cuarto_t.paragraph_format.space_after = Pt(6)
    run_ct = p_cuarto_t.add_run(f"CUADRO 1. Actividades Planificadas de la Semana {datos['num_semana']}")
    run_ct.font.name = 'Times New Roman'
    run_ct.font.size = Pt(10)
    run_ct.font.bold = True

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    # Inyectar propiedad XML para congelar dimensiones de columnas
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    table._tbl.tblPr.append(tblLayout)

    # Encabezados con sombreado
    headers = ["Nro", "Actividad", "Descripción", "Recursos"]
    for i, h in enumerate(headers):
        set_cell_format(table.rows[0].cells[i], h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), "D9E2F3")
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_elm)

    # Inyección de filas de actividades
    for nro, act, desc, rec in datos["actividades"]:
        row = table.add_row()
        set_cell_format(row.cells[0], nro, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_format(row.cells[1], act, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_format(row.cells[2], desc, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_cell_format(row.cells[3], rec, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Forzar el reparto de los 14.59 cm de ancho de página útil
    widths = [Cm(1.2), Cm(3.5), Cm(6.5), Cm(3.39)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.add_paragraph()

    # 5. Entregables
    p_ent_t = doc.add_paragraph()
    p_ent_t.paragraph_format.space_after = Pt(4)
    run_et = p_ent_t.add_run("ENTREGABLES DE LA SEMANA:")
    run_et.font.name = 'Times New Roman'
    run_et.font.size = Pt(12)
    run_et.font.bold = True

    add_paragraph_normado(doc, datos["entregables"])

    doc.add_paragraph()

    # 6. Observaciones (Mantenemos tu estructura exacta de 7 líneas a 75 de ancho)
    p_obs_t = doc.add_paragraph()
    p_obs_t.paragraph_format.space_after = Pt(4)
    run_ot2 = p_obs_t.add_run("OBSERVACIONES DEL TUTOR INDUSTRIAL:")
    run_ot2.font.name = 'Times New Roman'
    run_ot2.font.size = Pt(12)
    run_ot2.font.bold = True

    p_obs = doc.add_paragraph()
    p_obs.paragraph_format.line_spacing = 1.5
    run_obs = p_obs.add_run("(Espacio reservado para la evaluación cualitativa del tutor sobre puntualidad, actitud, comprensión técnica y cumplimiento de normas internas)")
    run_obs.font.name = 'Times New Roman'
    run_obs.font.size = Pt(11)
    run_obs.font.italic = True

    for _ in range(4):
        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_before = Pt(8)
        p_line.paragraph_format.space_after = Pt(0)
        run_line = p_line.add_run("_" * 75)
        run_line.font.name = 'Times New Roman'
        run_line.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # 7. Firmas en Tabla Invisible
    sign_table = doc.add_table(rows=2, cols=2)
    remove_table_borders(sign_table)
    sign_table.allow_autofit = False
    for row in sign_table.rows:
        row.cells[0].width = Cm(7.29)
        row.cells[1].width = Cm(7.30)

    set_cell_format(sign_table.rows[0].cells[0], "_________________________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_format(sign_table.rows[0].cells[1], "_________________________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_format(sign_table.rows[1].cells[0], "Firma del Pasante:\nJuliano Cardona | C.I.: 32.281.199", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_format(sign_table.rows[1].cells[1], "Firma del Tutor Industrial:\nIng. Yasmin Sabaneta | C.I.: 14.187.924", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # 8. Bloque Sello Corporativo (Título limpio)
    p_sello = doc.add_paragraph()
    p_sello.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sello.paragraph_format.space_before = Pt(12)
    run_s = p_sello.add_run("SELLO DE LA EMPRESA")
    run_s.font.name = 'Times New Roman'
    run_s.font.size = Pt(10)
    run_s.font.bold = True

    # Salvar y Compilar
    fn_word = f"Cronograma_Informatica_Semana{datos['num_semana']}_IUTECP.docx"
    fn_pdf = fn_word.replace(".docx", ".pdf")
    doc.save(fn_word)
    print(f"✔ Generado Word: {fn_word}")

    try:
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', fn_word], check=True, stdout=subprocess.DEVNULL)
        print(f"✔ Convertido PDF: {fn_pdf}")
    except FileNotFoundError:
        print(f"⚠️  Word creado. LibreOffice headless no se ejecutó.")

# ================================================================
#  COLECCIÓN DE DATOS MAESTRA (SEMANAS 1 A 4 UPDATED)
# ================================================================

reportes_pasantias = [
    # ── SEMANA 1 ────────────────────────────────────────────────
    {
        "num_semana": "1",
        "periodo": "Del 15 al 19 de junio de 2026",
        "objetivo": "Realizar el diagnóstico técnico del flujo documental en el área de Presidencia para identificar las deficiencias de integridad, redundancia y trazabilidad en el sistema actual basado en hojas de cálculo, y definir los requerimientos estructurales y funcionales de la nueva Base de Datos de control documental.",
        "entregables": "Informe técnico-diagnóstico que evidencia la descripción del flujo procedimental, la identificación de inconsistencias en el sistema actual de registro, y la definición estructurada de los campos, relaciones y reglas de validación requeridos para la Base de Datos de Trazabilidad Documental del Área de Presidencia.",
        "actividades": [
            ("1", "Inducción institucional", "Recepción por la Gerencia de Recursos Humanos, recorrido por las instalaciones del área administrativa, presentación del equipo de trabajo y revisión de protocolos de seguridad integral, confidencialidad y manejo de información clasificada.", "Bienvenida, lineamientos generales de la corporación"),
            ("2", "Análisis del proceso documental actual", "Observación directa y mapeo del ciclo de vida de los documentos: recepción, registro, derivación, firma y archivo físico o digital. Estudio de las hojas de cálculo vigentes utilizadas para el control.", "Formatos de hojas de cálculo actuales, bitácora de observación"),
            ("3", "Detección de fallas procedimentales", "Identificación de puntos críticos: redundancia de datos, actualización manual de estatus, falta de validaciones automáticas y pérdida de trazabilidad en movimientos interdepartamentales.", "Registros históricos, entrevistas al personal administrativo"),
            ("4", "Levantamiento de campos y atributos", "Determinación de los metadatos esenciales que debe capturar la base de datos: fecha de entrada o salida, remitente, destinatario, tipo de documento, estatus de firma, observaciones y alertas de vencimiento.", "Plantilla de requerimientos técnicos, herramientas de diseño básico"),
            ("5", "Elaboración del informe diagnóstico", "Redacción técnica del documento que consolida los hallazgos, propone la arquitectura inicial de la base de datos y establece la línea base para el desarrollo semanal.", "Computador, material de apoyo recabado")
        ]
    },
    # ── SEMANA 2 (REESTRUCTURADA CON LABORES DE OFICINA) ──────────
    {
        "num_semana": "2",
        "periodo": "Del 22 al 26 de junio de 2026",
        "objetivo": "Diseñar el modelo lógico y físico de la Base de Datos de Trazabilidad Documental y ejecutar el registro rutinario de expedientes administrativos en los sistemas de control preexistentes de la organización.",
        "entregables": "Diagrama Entidad-Relación validado, modelo físico normalizado, diccionario de datos completo y actualización de expedientes en la hoja de cálculo.",
        "actividades": [
            ("1", "Revisión de requerimientos y diseño de Entidad Relación", "Validación de los requerimientos analizados con el tutor industrial y construcción del diagrama Entidad-Relación que representa las entidades principales del sistema (Documento, Observaciones, Estatus).", "Informe diagnóstico, esbozos lógicos, cuadernos de notas"),
            ("2", "Registro de expedientes en base de datos previa", "Transmisión, clasificación y registro manual de expedientes físicos correspondientes a la gestión de Planificación y Control de documentos en la base de datos de Excel preexistente de la oficina.", "Hojas de cálculo institucionales (Excel), expedientes físicos"),
            ("3", "Normalización del modelo relacional", "Aplicación de las dos primeras formas normales sobre el esquema de base de datos propuesto para eliminar redundancias operativas y asegurar la integridad referencial.", "Esquema relacional, documentación técnica de normalización"),
            ("4", "Actualización y vaciado continuo en Excel", "Continuación de las labores correspondientes a la carga, y actualización de datos de nuevos expedientes en la plantilla de Excel matriz de la organización.", "Base de datos Excel preexistente, documentos recibidos"),
            ("5", "Diccionario de datos y archivos de creación", "Elaboración formal del diccionario de datos (campos, tipos de datos, restricciones de clave) y redacción de las instrucciones estructuradas para la inicialización física en SQLite.", "Plantilla de diccionario de datos, DB Browser for SQLite")
        ]
    },
    # ── SEMANA 3 (REESTRUCTURADA CON LABORES DE OFICINA) ──────────
    {
        "num_semana": "3",
        "periodo": "Del 29 de junio al 3 de julio de 2026",
        "objetivo": "Implementar la base de datos relacional del proyecto, desarrollar los módulos esenciales de consulta y ejecutar los flujos de transcripción y control de expedientes de la oficina en soporte digital.",
        "entregables": "Módulos funcionales de registro y seguimiento de documentos en el entorno de desarrollo, y procesamiento al día de los expedientes físicos en el sistema Excel.",
        "actividades": [
            ("1", "Despliegue físico, configuración e inicialización de la base de datos relacional", "Ejecución sistemática de la secuencia de comandos de arquitectura estructural en el motor de base de datos SQLite, procediendo con la migración de los registros históricos previamente validados para garantizar la correcta puesta en marcha del sistema de trazabilidad.", "Motor de base de datos SQLite, registros y datos históricos"),
            ("2", "Procesamiento y transcripción de expedientes en Excel", "Apoyo administrativo mediante la actualización del estatus de expedientes recibidos, dentro del archivo Excel de la oficina.", "Computador de la oficina, hojas de cálculo preexistentes"),
            ("3", "Diseño y desarrollo de interfaces de usuario", "Construcción de los formularios de captura de datos y configuración de las lógicas de actualización para el control del historial de movimientos documentales.", "Entorno de LibreOffice Base, base de datos SQLite"),
            ("4", "Control de documentos y validación de consultas", "Carga de metadatos para el registro de movimientos en el archivo Excel, aprovechando de forma simultánea estos flujos de información real para ejecutar pruebas, depuración y validación de las consultas de trazabilidad SQL en el entorno del proyecto.", "Sistema Excel, documentos, motor de base de datos SQLite"),
            ("5", "Consultas de trazabilidad y evaluación técnica", "Programación de consultas estructuradas de búsqueda por descripción o fecha, y detección de limitaciones técnicas en la escalabilidad de las macros del entorno LibreOffice Base.", "Editor de base de datos")
        ]
    },
    # ── SEMANA 4 (CON CHARLA SIAHO Y HOJA DE RECOBROS REALES) ────
    {
        "num_semana": "4",
        "periodo": "Del 6 al 10 de julio de 2026",
        "objetivo": "Diseñar la interfaz gráfica del sistema tomando como referencia el modelo funcional previo, asistir a las capacitaciones de seguridad laboral de la empresa y gestionar el llenado de expedientes.",
        "entregables": "Código fuente optimizado del entorno visual autónomo.",
        "actividades": [
            ("1", "Desarrollo de interfaz de escritorio", "Programación del entorno visual utilizando tecnologías de marcado y estilos empaquetados localmente.", "Editor de código fuente, repositorio github"),
            ("2", "Capacitación en seguridad y prevención", "Asistencia a la charla técnica sobre Identificación y Notificación de Peligros y Riesgos en Instalaciones y Puestos de Trabajo dictada por el departamento de SIAHO.", "Material formativo de SIAHO, sala de conferencias"),
            ("3", "Actualización de expedientes", "Carga de expedientes recibidos en la hoja de cálculo.", "Archivo central Excel"),
            ("4", "Diseño conceptual de esquemas remanentes", "Planificación estructural e identificación de entidades para los fondos documentales faltantes.", "Herramientas de diseño lógico, documentación interna"),
            ("5", "Consolidación de código y documentación", "Sincronización de las versiones funcionales de la interfaz gráfica en la plataforma de control de código en línea y redacción de documentación.", "Dispositivo móvil, plataforma de control de versiones")
        ]
    },
    # ── SEMANA 5 ────────────────────────────────────────────────
    {
        "num_semana": "5",
        "periodo": "Del 13 al 17 de julio de 2026",
        "objetivo": "Completar los módulos visuales de consulta e historial de movimientos del sistema, e integrar la lógica de visualización dinámica de expedientes por proceso administrativo, asistiendo en paralelo en las labores de revisión y canalización de expedientes.",
        "entregables": "Módulos de consulta e historial de movimientos funcionales e integrados en el entorno de escritorio, con filtrado operativo por proceso administrativo y rango de fechas.",
        "actividades": [
            ("1", "Desarrollo del módulo de consulta de expedientes", "Programación de la interfaz de búsqueda con filtros dinámicos por número de proceso, tipo de documento y rango de fechas, conectada directamente a las consultas SQL del motor SQLite.", "Editor de código fuente, motor SQLite, repositorio GitHub"),
            ("2", "Revisión y canalización de correspondencia presidencial", "Apoyo administrativo en la revisión de expedientes recibidos, canalización para su firma y posterior despacho al departamento destinatario.", "Expedientes físicos, hojas de control de la oficina"),
            ("3", "Integración del historial de movimientos", "Implementación del panel de historial que muestra la línea de tiempo completa de cada expediente, con indicación del responsable de cada etapa.", "Entorno de desarrollo, base de datos SQLite con datos reales"),
            ("4", "Pruebas de usabilidad del módulo de consulta", "Ejecución de casos de prueba sobre el módulo de búsqueda utilizando datos reales de expedientes registrados en semanas anteriores, verificando la integridad de los resultados.\n\n", "Datos históricos de la base de datos"),
            ("5", "Documentación técnica de los módulos desarrollados", "Redacción de la descripción funcional de los módulos de consulta e historial para su incorporación en el informe técnico de pasantías.", "Computador, editor de texto")
        ]
    },
    # ── SEMANA 6 ────────────────────────────────────────────────
    {
        "num_semana": "6",
        "periodo": "Del 20 al 24 de julio de 2026",
        "objetivo": "Empaquetar la aplicación de escritorio mediante el Marco de trabajo Wails con integración de WebView2 portable, logrando un ejecutable sin dependencias de instalación, y colaborar en las labores regulares.",
        "entregables": "Ejecutable autónomo del sistema (.exe) empaquetado con Wails y WebView2 portable, funcional en el equipo corporativo.",
        "actividades": [
            ("1", "Configuración del entorno de empaquetado con Wails", "Instalación y configuración del Marco de trabajo Wails en el entorno de desarrollo, revisión de compatibilidad con el sistema operativo del equipo corporativo y definición de los parámetros de compilación.", "Marco de trabajo Wails, entorno de ejecución Go, editor de código fuente"),
            ("2", "Control de expedientes firmados", "Colaboración en el registro de entrada y salida de expedientes firmados durante el período y apoyo con actualización del control en la hoja de cálculo.", "Expedientes firmados, hoja de cálculo"),
            ("3", "Integración de WebView2 portable en el ejecutable", "Incorporación del componente WebView2 portable al proceso de compilación de Wails para garantizar la renderización de la interfaz.", "Marco de trabajo Wails"),
            ("4", "Compilación y prueba del ejecutable", "Generación del archivo ejecutable, prueba de su funcionamiento sobre el equipo corporativo de la oficina y verificación de que todos los módulos responden correctamente sin instalaciones adicionales.\n", "Equipo corporativo de la oficina, ejecutable compilado"),
            ("5", "Corrección de errores post-empaquetado", "Identificación y corrección de los errores detectados durante las pruebas del ejecutable: rutas relativas del archivo de base de datos y ajustes de permisos de lectura y escritura.", "Editor de código fuente, repositorio GitHub")
        ]
    },
    # ── SEMANA 7 ────────────────────────────────────────────────
    {
        "num_semana": "7",
        "periodo": "Del 27 al 31 de julio de 2026",
        "objetivo": "Ejecutar las pruebas de validación funcional del sistema sobre datos reales del departamento, depurar los errores detectados y verificar el correcto funcionamiento de los módulos de registro, trazabilidad y generación de reportes.",
        "entregables": "Informe de pruebas funcionales con registro de casos de prueba, errores detectados, correcciones aplicadas y confirmación del funcionamiento integral del sistema sobre el entorno corporativo.",
        "actividades": [
            ("1", "Diseño del plan de pruebas funcionales", "Elaboración de la matriz de casos de prueba que cubre los flujos críticos del sistema: actualización de estatus, consulta por filtros y generación de reportes.", "Plantilla de casos de prueba"),
            ("2", "Verificación de estatus documental", "Revisión de los expedientes registrados durante la semana en el sistema, para contrastarlo con la información en la hoja de cálculo, identificando pendientes y documentos antes de continuar con las pruebas del prototipo.", "Hoja de cálculo, sistema ejecutable"),
            ("3", "Ejecución de pruebas sobre datos reales", "Corrida sistemática de los casos de prueba definidos utilizando los expedientes registrados, documentando el resultado de cada caso y clasificando los errores encontrados.", "Sistema ejecutable, base de datos con datos reales"),
            ("4", "Depuración y corrección de errores detectados", "Análisis de los errores registrados durante las pruebas, corrección en el código fuente y ejecución de pruebas de para confirmar que las correcciones no introducen nuevas fallas.\n", "Editor de código fuente, repositorio GitHub"),
            ("5", "Redacción del informe de pruebas funcionales", "Consolidación de los resultados de las pruebas en un informe técnico que documenta los casos ejecutados, los errores corregidos y la validación final de los módulos del sistema.", "Computador, editor de texto")
        ]
    },
    # ── SEMANA 8 ────────────────────────────────────────────────
    {
        "num_semana": "8",
        "periodo": "Del 3 al 7 de agosto de 2026",
        "objetivo": "Redactar el informe técnico de pasantías, elaborar los manuales de usuario del sistema y estructurar definitivamente el código fuente para la entrega institucional, participando en las actividades regulares de cierre administrativo del período.",
        "entregables": "Informe técnico de pasantías completo en su versión preliminar, manual de usuario del sistema y código fuente estructurado y comentado en el repositorio de control de versiones.",
        "actividades": [
            ("1", "Redacción del informe técnico de pasantías", "Desarrollo de los capítulos I al V del informe académico: realidad organizacional, planificación, marco teórico, actividades realizadas, conclusiones y recomendaciones, siguiendo las normas del IUTECP.", "Computador, editor de texto, normativas institucionales"),
            ("2", "Organización documental", "Apoyo en el registro de expedientes en la hoja de cálculo y actividades regulares de soporte.", "Expedientes físicos, hoja de cálculo"),
            ("3", "Elaboración del manual de usuario del sistema", "Redacción del manual de usuario que describe paso a paso el uso de los módulos de registro, consulta, historial y generación de reportes, con capturas de pantalla de la interfaz.", "Sistema ejecutable, editor de texto, capturas de pantalla"),
            ("4", "Estructuración y comentado del código fuente", "Revisión integral del código fuente, incorporación de comentarios explicativos en las funciones críticas, organización de los archivos del proyecto y sincronización de la versión final en el repositorio.\n\n", "Editor de código fuente, repositorio GitHub"),
            ("5", "Revisión preliminar del informe con el tutor académico", "Envío de la versión preliminar del informe al tutor académico para su revisión, recepción de observaciones iniciales e incorporación de las correcciones formales indicadas.", "Informe preliminar, medios de comunicación institucional")
        ]
    },
    # ── SEMANA 9 (FINAL) ────────────────────────────────────────
    {
        "num_semana": "9",
        "periodo": "Del 10 al 14 de agosto de 2026",
        "objetivo": "Consolidar y revisar el informe académico en su versión final, incorporar las observaciones del tutor industrial, presentar el prototipo funcional ante la tutora industrial para su evaluación y preparar los anexos definitivos para la entrega institucional.",
        "entregables": "Informe técnico de pasantías en versión final con todos los anexos, carta de aprobación del tutor industrial firmada y prototipo funcional del sistema debidamente validado y entregado.",
        "actividades": [
            ("1", "Incorporación de observaciones del tutor industrial", "Revisión de las correcciones y sugerencias indicadas por la tutora industrial sobre el borrador del informe y el funcionamiento del sistema, e incorporación de los ajustes pertinentes.", "Borrador del informe, observaciones escritas de la tutora"),
            ("2", "Presentación del prototipo funcional ante la tutora industrial", "Demostración formal del sistema ante la Ing. Yasmin Sabaneta, recorriendo los módulos de registro, consulta, historial y generación de reportes con expedientes reales del departamento.", "Sistema ejecutable, equipo informático de la oficina"),
            ("3", "Preparación de los anexos definitivos", "Organización y presentación formal de los anexos del informe: diagrama Entidad-Relación y capturas de la interfaz del sistema.", "Impresora, archivos de imagen, editor de texto"),
            ("4", "Revisión y consolidación final del informe académico", "Corrección ortográfica y de formato del informe completo, verificación del cumplimiento de las normas del IUTECP, numeración de páginas, índices y lista de referencias bibliográficas.\n", "Informe completo, normativas del IUTECP, computador"),
            ("5", "Entrega institucional y firma de carta de aprobación", "Presentación oficial del informe técnico final ante las autoridades académicas e institucionales correspondientes y obtención de la firma de aprobación de la tutora industrial para su entrega al IUTECP.", "Informe impreso y encuadernado, carta de aprobación")
        ]
    },
]

# ================================================================
#  EJECUCIÓN GENERAL
# ================================================================

if __name__ == "__main__":
    print("🚀 Ejecutando motor de automatización documental...")
    for idx, reporte in enumerate(reportes_pasantias):
        print(f"\n[Procesando {idx+1}/{len(reportes_pasantias)}] Generando entregables Semana {reporte['num_semana']}...")
        generar_documento_semana(reporte)
    print("\n✨ ¡Listo! Revisa tu directorio, ya tienes los 8 archivos listos para imprimir.")
