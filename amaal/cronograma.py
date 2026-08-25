import os
import subprocess
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ================================================================
#  DATOS DEL PASANTE (MODIFIQUE ESTOS VALORES CON SUS DATOS)
# ================================================================
PASANTE_NOMBRE = "Alrifaai Alrifaaie, Amaal"
PASANTE_CI = "31.985.792"
TUTOR_INDUSTRIAL = "Lic. Lenny Mata"
TUTOR_INDUSTRIAL_CI = "8969750"
FUENTE_CUADRO_SEMANAL = "Alrifaai A. (2026)."

# ================================================================
#  FUNCIONES DE CONFIGURACIÓN Y FORMATO (ESTÁTICAS - NORMA IUTECP)
# ================================================================

def setup_page(doc):
    """Márgenes Art. 6 (4cm Izquierdo para encuadernación, 3cm los demás)"""
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

def add_paragraph_normado(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)

    run_l = p.add_run(label)
    run_l.font.name = 'Times New Roman'
    run_l.font.size = Pt(12)
    run_l.font.bold = True

    run_v = p.add_run(value)
    run_v.font.name = 'Times New Roman'
    run_v.font.size = Pt(12)
    return p

def add_fuente(doc, fuente):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    etiqueta = p.add_run("Fuente: ")
    etiqueta.font.name = 'Times New Roman'
    etiqueta.font.size = Pt(10)
    etiqueta.font.italic = True
    texto = p.add_run(fuente)
    texto.font.name = 'Times New Roman'
    texto.font.size = Pt(10)
    return p

def set_cell_format(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        borders.append(border)
    tblPr.append(borders)

# ================================================================
#  PROCESADOR CORE: COMPILA UNA SEMANA ESPECÍFICA
# ================================================================

def generar_documento_semana(datos):
    doc = Document()
    setup_page(doc)

    # 1. Título Centralizado
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_after = Pt(12)
    run_t = p_titulo.add_run("PLAN SEMANAL DE PASANTÍAS")
    run_t.font.name = 'Times New Roman'
    run_t.font.size = Pt(14)
    run_t.font.bold = True

    # 2. Datos Informativos de la Empresa y la Pasante
    add_label_value(doc, "EMPRESA: ", "Ingeniería de Telecomunicaciones, C.A.")
    add_label_value(doc, "DEPARTAMENTO: ", "Departamento de Administración")
    add_label_value(doc, "ESPECIALIDAD: ", "Administración")
    add_label_value(doc, "SEMANA N°: ", datos["num_semana"])
    add_label_value(doc, "PERÍODO: ", datos["periodo"])
    add_label_value(doc, "PASANTE: ", f"{PASANTE_NOMBRE} |  C.I.: {PASANTE_CI}")
    add_label_value(doc, "TUTOR INDUSTRIAL: ", f"{TUTOR_INDUSTRIAL}  |  C.I.: {TUTOR_INDUSTRIAL_CI}")

    doc.add_paragraph()

    # 3. Objetivo de la semana
    p_obj_t = doc.add_paragraph()
    p_obj_t.paragraph_format.space_after = Pt(4)
    run_ot = p_obj_t.add_run("OBJETIVO DE LA SEMANA:")
    run_ot.font.name = 'Times New Roman'
    run_ot.font.size = Pt(12)
    run_ot.font.bold = True

    add_paragraph_normado(doc, datos["objetivo"])

    doc.add_paragraph()

    # 4. Tabla de Actividades Planificadas (Dimensiones fijas para evitar desbordes)
    p_cuarto_t = doc.add_paragraph()
    p_cuarto_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cuarto_t.paragraph_format.space_after = Pt(6)
    run_ct = p_cuarto_t.add_run(f"CUADRO 1. Actividades Planificadas de la Semana {datos['num_semana']}")
    run_ct.font.name = 'Times New Roman'
    run_ct.font.size = Pt(10)
    run_ct.font.bold = True

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    table._tbl.tblPr.append(tblLayout)

    # Encabezados con Sombreado Azul Pastel
    headers = ["Nro", "Actividad", "Descripción", "Recursos"]
    for i, h in enumerate(headers):
        set_cell_format(table.rows[0].cells[i], h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), "D9E2F3")
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_elm)

    # Inyección de actividades semanales
    for nro, act, desc, rec in datos["actividades"]:
        row = table.add_row()
        set_cell_format(row.cells[0], nro, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_format(row.cells[1], act, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_format(row.cells[2], desc, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_cell_format(row.cells[3], rec, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Ancho total útil: 14.59 cm
    widths = [Cm(1.2), Cm(3.5), Cm(6.5), Cm(3.39)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    add_fuente(doc, FUENTE_CUADRO_SEMANAL)

    # 5. Entregables
    p_ent_t = doc.add_paragraph()
    p_ent_t.paragraph_format.space_after = Pt(4)
    run_et = p_ent_t.add_run("ENTREGABLES DE LA SEMANA:")
    run_et.font.name = 'Times New Roman'
    run_et.font.size = Pt(12)
    run_et.font.bold = True

    add_paragraph_normado(doc, datos["entregables"])

    doc.add_paragraph()

    # 6. Observaciones del Tutor Industrial (Inyección cualitativa real)
    p_obs_t = doc.add_paragraph()
    p_obs_t.paragraph_format.space_after = Pt(4)
    p_obs_t.paragraph_format.keep_with_next = True
    run_ot2 = p_obs_t.add_run("OBSERVACIONES DEL TUTOR INDUSTRIAL:")
    run_ot2.font.name = 'Times New Roman'
    run_ot2.font.size = Pt(12)
    run_ot2.font.bold = True

    observaciones = datos.get("observaciones_tutor", "").strip()
    if observaciones:
        add_paragraph_normado(doc, observaciones)
    else:
        p_obs = doc.add_paragraph()
        p_obs.paragraph_format.line_spacing = 1.5
        run_obs = p_obs.add_run("(Espacio reservado para la evaluación cualitativa del tutor sobre control de procesos, puntualidad, manejo de solicitudes y cumplimiento de directrices contables)")
        run_obs.font.name = 'Times New Roman'
        run_obs.font.size = Pt(11)
        run_obs.font.italic = True

        for _ in range(4):
            p_line = doc.add_paragraph()
            p_line.paragraph_format.space_before = Pt(8)
            p_line.paragraph_format.space_after = Pt(0)
            run_line = p_line.add_run("_" * 75)
            run_line.font.name = 'Times New Roman'
            run_line.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # 7. Firmas (Estructura de Tabla Invisible)
    sign_table = doc.add_table(rows=2, cols=2)
    remove_table_borders(sign_table)
    sign_table.allow_autofit = False
    for row in sign_table.rows:
        row.cells[0].width = Cm(7.29)
        row.cells[1].width = Cm(7.30)

    # Insertar firma del pasante si existe
    ruta_firma_amaal = os.path.join("imagenes", "firma_amaal.png")
    if os.path.exists(ruta_firma_amaal):
        cell_fk = sign_table.rows[0].cells[0]
        cell_fk.text = ""
        cell_fk.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_fk = cell_fk.paragraphs[0].add_run()
        run_fk.add_picture(ruta_firma_amaal, width=Cm(3.5))
    else:
        set_cell_format(sign_table.rows[0].cells[0], "_________________________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    ruta_firma_tutor = os.path.join("imagenes", "firma_tutor_amaal.png")
    if os.path.exists(ruta_firma_tutor):
        cell_ft = sign_table.rows[0].cells[1]
        cell_ft.text = ""
        cell_ft.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_ft = cell_ft.paragraphs[0].add_run()
        run_ft.add_picture(ruta_firma_tutor, width=Cm(3.5))
    else:
        set_cell_format(sign_table.rows[0].cells[1], "_________________________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_format(sign_table.rows[1].cells[0], f"Firma del Pasante:\n{PASANTE_NOMBRE} |  C.I.: {PASANTE_CI}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_format(sign_table.rows[1].cells[1], f"Firma del Tutor Industrial:\n{TUTOR_INDUSTRIAL}  |  C.I.: {TUTOR_INDUSTRIAL_CI}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # 8. Bloque Sello Corporativo
    ruta_sello = os.path.join("imagenes", "sello_amaal.png")
    if os.path.exists(ruta_sello):
        p_sello = doc.add_paragraph()
        p_sello.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sello.paragraph_format.space_before = Pt(12)
        run_s = p_sello.add_run()
        run_s.add_picture(ruta_sello, width=Cm(4))
    else:
        p_sello = doc.add_paragraph()
        p_sello.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sello.paragraph_format.space_before = Pt(12)
        run_s = p_sello.add_run("SELLO DE LA EMPRESA")
        run_s.font.name = 'Times New Roman'
        run_s.font.size = Pt(10)
        run_s.font.bold = True

    fn_word = f"Cronograma_Administracion_Semana{datos['num_semana']}_IUTECP.docx"
    doc.save(fn_word)
    print(f"✔ Generado Word: {fn_word}")

    try:
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', fn_word], check=True, stdout=subprocess.DEVNULL)
        print(f"✔ Convertido PDF: {fn_word.replace('.docx', '.pdf')}")
    except Exception:
        print(f"⚠️ Word creado. LibreOffice headless no se ejecutó.")

# ================================================================
#  COLECCIÓN DE DATOS MAESTRA (CRONOGRAMA ADAPTADO A TSU - 10 SEMANAS)
# ================================================================

reportes_pasantias = [
    {
        'num_semana': '1',
        'periodo': 'Del 01 al 05 de junio de 2026',
        'objetivo': 'Realizar la inducción institucional y reconocer el funcionamiento de Atención al Cliente, '
                    'identificando cómo se reciben pagos y solicitudes de los suscriptores.',
        'entregables': 'Registro de inducción y esquema preliminar del flujo observado en Atención al Cliente.',
        'observaciones_tutor': 'La pasante mostró buena disposición durante el proceso de inducción, adaptándose adecuadamente a las normas internas y demostrando interés en conocer el funcionamiento de Atención al Cliente y el manejo de las solicitudes de los suscriptores.',
        'actividades': [
            ('1',
             'Inducción y recorrido',
             'Conocer las instalaciones, normas internas y áreas vinculadas con la atención de suscriptores.',
             'Guía de inducción, cuaderno'),
            ('2',
             'Reconocimiento de Atención al Cliente',
             'Identificar las funciones del puesto y los tipos de requerimientos recibidos.',
             'Registros del área'),
            ('3',
             'Registro de pagos',
             'Observar y apoyar el registro de pagos reportados por los suscriptores.',
             'Comprobantes, PC'),
            ('4',
             'Recepción de solicitudes',
             'Reconocer cómo se reciben solicitudes de afiliación e incidencias.',
             'Formatos y registros de casos'),
            ('5',
             'Bitácora inicial',
             'Documentar el recorrido básico de la información desde el contacto con el suscriptor.',
             'Cuaderno de campo')
        ]
    },

    {
        'num_semana': '2',
        'periodo': 'Del 08 al 12 de junio de 2026',
        'objetivo': 'Apoyar las actividades de Atención al Cliente y levantar el recorrido de las solicitudes de '
                    'afiliación e incidencias.',
        'entregables': 'Registro de actividades de atención y mapa preliminar del proceso de solicitudes.',
        'observaciones_tutor': 'La pasante cumplió satisfactoriamente con las actividades asignadas, demostrando responsabilidad y organización en el registro de pagos, revisión de expedientes y seguimiento inicial de las solicitudes recibidas.',
        'actividades': [
            ('1',
             'Recepción de pagos',
             'Apoyar la verificación y registro de pagos informados por clientes.',
             'Comprobantes, PC'),
            ('2',
             'Atención de solicitudes',
             'Registrar datos básicos de afiliaciones e incidencias reportadas.',
             'Registro de solicitudes'),
            ('3',
             'Revisión de expedientes',
             'Verificar documentación básica asociada a solicitudes de servicio.',
             'Carpetas de clientes'),
            ('4',
             'Mapeo del flujo',
             'Observar hacia qué áreas se remiten los diferentes tipos de solicitudes.',
             'Hojas de flujo'),
            ('5',
             'Registro de hallazgos',
             'Anotar puntos donde la información cambia de responsable o pierde continuidad.',
             'Bitácora')
        ]
    },

    {
        'num_semana': '3',
        'periodo': 'Del 15 al 19 de junio de 2026',
        'objetivo': 'Continuar el apoyo en Atención al Cliente y diseñar la guía de entrevista estructurada para '
                    'evaluar el seguimiento administrativo de las solicitudes.',
        'entregables': 'Guía de entrevista estructurada y registro actualizado del flujo observado.',
        'observaciones_tutor': 'La pasante evidenció iniciativa y capacidad de organización al continuar el seguimiento de los casos y preparar la guía de entrevista, mostrando receptividad ante las recomendaciones realizadas para mejorar el instrumento.',
        'actividades': [
            ('1',
             'Registro de pagos y solicitudes',
             'Continuar el apoyo en recepción de pagos y requerimientos de clientes.',
             'PC, comprobantes'),
            ('2',
             'Seguimiento inicial de casos',
             'Verificar el estado de solicitudes remitidas a otras áreas.',
             'Registros de casos'),
            ('3',
             'Diseño de guía de entrevista',
             'Preparar preguntas sobre recepción, actualización, comunicación y cierre de solicitudes.',
             'Procesador de textos'),
            ('4',
             'Revisión con tutor',
             'Revisar la pertinencia de la guía y realizar ajustes antes de aplicarla.',
             'Guía de entrevista'),
            ('5',
             'Planificación de aplicación',
             'Definir las personas y momentos adecuados para aplicar la entrevista.',
             'Agenda')
        ]
    },

    {
        'num_semana': '4',
        'periodo': 'Del 22 al 26 de junio de 2026',
        'objetivo': 'Iniciar la rotación por Administración, apoyar pagos y facturación, y aplicar entrevistas '
                    'estructuradas para identificar deficiencias de control.',
        'entregables': 'Guías de entrevista aplicadas y matriz preliminar de hallazgos.',
        'observaciones_tutor': 'La pasante se integró adecuadamente al Departamento de Administración, mostrando responsabilidad en las actividades de pagos y facturación. Aplicó las entrevistas de manera organizada y registró correctamente los principales hallazgos identificados.',
        'actividades': [
            ('1',
             'Rotación a Administración',
             'Reconocer las funciones administrativas relacionadas con pagos, facturación y seguimiento de casos.',
             'Documentos del área'),
            ('2',
             'Procesamiento de pagos',
             'Apoyar la verificación y registro administrativo de pagos del período.',
             'Comprobantes, PC'),
            ('3',
             'Apoyo en facturación',
             'Colaborar en tareas de facturación y verificación de datos.',
             'Facturas, sistema de registro'),
            ('4',
             'Aplicación de entrevistas',
             'Aplicar la guía estructurada al personal vinculado con el proceso de solicitudes.',
             'Guía de entrevista'),
            ('5',
             'Matriz de hallazgos',
             'Organizar las respuestas sobre demoras, comunicación y trazabilidad.',
             'Hoja de cálculo')
        ]
    },

    {
        'num_semana': '5',
        'periodo': 'Del 29 de junio al 03 de julio de 2026',
        'objetivo': 'Completar el diagnóstico de las deficiencias del control administrativo y elaborar el diagrama de '
                    'Ishikawa.',
        'entregables': 'Diagnóstico situacional, diagrama de Ishikawa y avance del marco teórico.',
        'observaciones_tutor': 'La pasante demostró capacidad de análisis al identificar y organizar las causas asociadas con las deficiencias del control administrativo, desarrollando de forma adecuada el diagrama de Ishikawa y relacionándolo con los hallazgos obtenidos.',
        'actividades': [
            ('1',
             'Apoyo administrativo',
             'Colaborar en facturación, actualización de registros y seguimiento de casos.',
             'PC, documentos administrativos'),
            ('2',
             'Revisión de registros',
             'Contrastar solicitudes e incidencias para detectar retrasos y datos faltantes.',
             'Registros históricos'),
            ('3',
             'Diagrama de Ishikawa',
             'Representar las causas asociadas con procedimiento, comunicación, registro y seguimiento.',
             'Herramienta de diagramación'),
            ('4',
             'Análisis de causas',
             'Relacionar las causas con las principales manifestaciones observadas.',
             'Matriz de hallazgos'),
            ('5',
             'Bases teóricas',
             'Iniciar la redacción del Capítulo III sobre control administrativo y gestión de solicitudes.',
             'Bibliografía, procesador de textos')
        ]
    },

    {
        'num_semana': '6',
        'periodo': 'Del 06 al 10 de julio de 2026',
        'objetivo': 'Analizar el impacto de las deficiencias detectadas y consolidar las bases teóricas del informe.',
        'entregables': 'Análisis de impacto y Capítulo III en versión preliminar.',
        'observaciones_tutor': 'La pasante desarrolló las actividades asignadas con responsabilidad, mostrando capacidad para relacionar las deficiencias identificadas con sus efectos sobre la trazabilidad y atención de las solicitudes, manteniendo además un buen desempeño en las labores administrativas.',
        'actividades': [
            ('1',
             'Procesamiento de pagos',
             'Apoyar la conciliación y registro de pagos del período.',
             'Comprobantes, hoja de cálculo'),
            ('2',
             'Facturación',
             'Colaborar en la preparación y verificación de documentos de facturación.',
             'Facturas, PC'),
            ('3',
             'Seguimiento de casos',
             'Revisar incidencias y solicitudes pendientes para comparar tiempos de atención.',
             'Registros de casos'),
            ('4',
             'Análisis de impacto',
             'Relacionar los retrasos con problemas de trazabilidad y atención al suscriptor.',
             'Hoja de cálculo, matriz de análisis'),
            ('5',
             'Marco teórico',
             'Completar y revisar las bases teóricas y referencias utilizadas.',
             'Procesador de textos')
        ]
    },

    {
        'num_semana': '7',
        'periodo': 'Del 13 al 17 de julio de 2026',
        'objetivo': 'Diseñar el flujo estandarizado de gestión de solicitudes y continuar el apoyo en las tareas '
                    'administrativas del área.',
        'entregables': 'Flujograma preliminar con etapas, responsables y puntos de control.',
        'observaciones_tutor': 'La pasante mostró iniciativa en el diseño del flujo estandarizado, identificando de manera coherente las etapas, responsables y puntos de control necesarios. Asimismo, mantuvo una participación constante en las actividades administrativas del departamento.',
        'actividades': [
            ('1',
             'Seguimiento administrativo',
             'Actualizar registros de solicitudes y casos activos.',
             'PC, registros'),
            ('2',
             'Apoyo en reportes',
             'Colaborar en la preparación de reportes administrativos del período.',
             'Hoja de cálculo'),
            ('3',
             'Diseño del flujo',
             'Definir la secuencia propuesta desde la recepción hasta el cierre del caso.',
             'Herramienta de diagramación'),
            ('4',
             'Responsables por etapa',
             'Asignar el área responsable de registrar o actualizar cada fase.',
             'Organigrama, procesador de textos'),
            ('5',
             'Puntos de control',
             'Definir verificaciones y tiempos de referencia para el seguimiento.',
             'Hoja de trabajo')
        ]
    },

    {
        'num_semana': '8',
        'periodo': 'Del 20 al 24 de julio de 2026',
        'objetivo': 'Diseñar los formatos estandarizados y el procedimiento de control que acompaña al flujo propuesto.',
        'entregables': 'Formatos de afiliación e incidencias y procedimiento escrito en versión preliminar.',
        'observaciones_tutor': 'La pasante desarrolló adecuadamente los formatos de afiliación e incidencias, demostrando orden y criterio administrativo en la selección de los datos necesarios para facilitar el registro, actualización y seguimiento de las solicitudes.',
        'actividades': [
            ('1',
             'Apoyo en facturación',
             'Colaborar en la facturación y verificación de documentos del período.',
             'PC, facturas'),
            ('2',
             'Seguimiento de incidencias',
             'Actualizar información de casos administrativos y revisar pendientes.',
             'Registro de casos'),
            ('3',
             'Formato de afiliaciones',
             'Diseñar un formato uniforme para el registro y seguimiento de nuevas afiliaciones.',
             'Procesador de textos'),
            ('4',
             'Formato de incidencias',
             'Diseñar un formato uniforme para registrar incidencias y su estatus.',
             'Procesador de textos'),
            ('5',
             'Procedimiento escrito',
             'Redactar las reglas básicas de uso de los formatos y actualización del flujo.',
             'Procesador de textos')
        ]
    },

    {
        'num_semana': '9',
        'periodo': 'Del 27 al 31 de julio de 2026',
        'objetivo': 'Formular y validar la propuesta de mejora del control administrativo, incorporando el flujograma y '
                    'los formatos diseñados.',
        'entregables': 'Propuesta de mejora revisada con flujograma y formatos de control.',
        'observaciones_tutor': 'La pasante presentó la propuesta de mejora de manera organizada, demostrando dominio de los aspectos trabajados durante las pasantías. Recibió las observaciones realizadas con actitud receptiva e incorporó las correcciones pertinentes al documento.',
        'actividades': [
            ('1',
             'Apoyo operativo',
             'Participar en las actividades administrativas y de seguimiento de solicitudes del período.',
             'PC, registros'),
            ('2',
             'Redacción de propuesta',
             'Integrar diagnóstico, flujo, formatos y mecanismos de seguimiento.',
             'Procesador de textos'),
            ('3',
             'Flujograma definitivo',
             'Completar el diagrama del proceso propuesto.',
             'Herramienta de diagramación'),
            ('4',
             'Validación con tutor',
             'Presentar la propuesta al tutor industrial y registrar observaciones.',
             'Documento de propuesta'),
            ('5',
             'Ajustes',
             'Incorporar las correcciones pertinentes al documento.',
             'Procesador de textos')
        ]
    },

    {
        'num_semana': '10',
        'periodo': 'Del 03 al 07 de agosto de 2026',
        'objetivo': 'Consolidar el informe final, presentar los resultados y completar los recaudos de cierre de las '
                    'pasantías.',
        'entregables': 'Informe final, propuesta presentada y recaudos institucionales de cierre.',
        'observaciones_tutor': 'La pasante culminó satisfactoriamente su período de pasantías, demostrando responsabilidad, puntualidad, disposición para recibir orientaciones y cumplimiento de las actividades asignadas. Se evidenció un progreso favorable en el manejo de los procesos administrativos y en la formulación de propuestas de mejora.',
        'actividades': [
            ('1',
             'Consolidación del informe',
             'Integrar capítulos, referencias y anexos en el documento definitivo.',
             'PC, procesador de textos'),
            ('2',
             'Revisión normativa',
             'Verificar márgenes, interlineado, paginación, índices y referencias conforme a IUTECP.',
             'Normativa institucional'),
            ('3',
             'Presentación de resultados',
             'Presentar el diagnóstico y la propuesta ante la supervisión del área.',
             'Presentación, laptop'),
            ('4',
             'Revisión final',
             'Corregir ortografía y coherencia entre objetivos, cronograma, actividades y conclusiones.',
             'Procesador de textos'),
            ('5',
             'Recaudos de cierre',
             'Gestionar firmas, sello y documentación de culminación.',
             'Formatos institucionales')
        ]
    }
]

# ================================================================
#  EJECUCIÓN GENERAL
# ================================================================

if __name__ == "__main__":
    print("🚀 Iniciando motor de automatización de cronogramas para Administración (Amaal)...")
    for idx, reporte in enumerate(reportes_pasantias):
        print(f"\n[Compilando {idx+1}/{len(reportes_pasantias)}] Semana {reporte['num_semana']}...")
        generar_documento_semana(reporte)
    print("\n✨ ¡Proceso culminado con éxito! Se han generado los 10 archivos independientes en Word.")
