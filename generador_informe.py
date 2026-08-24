import os
import math
import re
import subprocess
import platform
from docx import Document
from docx.shared import Cm, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Importamos el contenido
import contenido as c

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ================================================================
#  CONSTANTES DE FORMATO (NORMATIVA IUTECP)
# ================================================================

# Fuente y tamaño (Art. 5)
FUENTE = 'Times New Roman'
TAMANO_BASE = 12           # Pt
TAMANO_TABLA = 10          # Pt (cuerpo de tablas)
TAMANO_TABLA_CHICO = 10    # Pt (encabezados y celdas compactas Gantt)
RESUMEN_MAX_PALABRAS = 300

# Página (Art. 6)
PAG_ANCHO = Cm(21.59)     # Carta
PAG_ALTO  = Cm(27.94)     # Carta
MARGEN_IZQ = Cm(4)        # Encuadernación
MARGEN_DER = Cm(3)
MARGEN_SUP = Cm(3)
MARGEN_INF = Cm(3)
MARGEN_SUP_CAP = Cm(5)    # Primera página de capítulo/parte

# Tipografía
INTERLINEADO = 1.5
SANGRIA_LINEA = Cm(1.25)          # Sangría de primera línea (Art. 7)
SANGRIA_CITA = Cm(1.25)           # Citas largas (Art. 22)
SANGRIA_NIVEL45 = Cm(1.27)        # Niveles 4 y 5 (Art. 9)
SANGRIA_REF = Cm(0.75)            # Sangría francesa referencias (Art. 25)

# Espaciados (Art. 8)
ESP_DOBLE = Pt(24)
ESP_SENCILLO = Pt(12)
ESP_TRIPLE = Pt(36)
ESP_CITA_ANTES = Pt(36)           # 3 espacios antes de cita larga
ESP_CITA_DESPUES = Pt(12)

# Colores de tablas
COLOR_ENCABEZADO = '4472C4'       # Azul corporativo
COLOR_FILA_PAR = 'D9E1F2'         # Gris azulado
COLOR_GANTT_VERDE = '70AD47'      # Barras activas Gantt
COLOR_TEXTO_CLARO = 'FFFFFF'
COLOR_TEXTO_OSCURO = '000000'

# Producto: nombres de archivo de salida
DOCX_SALIDA = "Informe_Pasantia_IUTECP.docx"
PDF_SALIDA  = "Informe_Pasantia_IUTECP.pdf"

# Tamaño de la portadilla de ANEXOS (Art. 26)
TAMANO_PORTADILLA_ANEXOS = 12

# Títulos por defecto de Cuadros y Gráficos (configurables desde contenido.py)
CUADRO_POBLACION_TITULO_DEF = "Cuadro 1. Población de los trabajadores de la empresa."
CUADRO_PLANIFICACION_TITULO_DEF = "Cuadro 2. Planificación integral de objetivos específicos."
CUADRO_CRONOGRAMA_TITULO_DEF = "Cuadro 3. Cronograma de actividades administrativas."

ESTILO_CAPITULO = 'IUTECP Capítulo'
ESTILO_PRELIMINAR = 'IUTECP Preliminar'
ESTILO_TITULO_NIVEL2 = 'IUTECP Título 2'
ESTILO_TITULO_NIVEL3 = 'IUTECP Título 3'
ESTILO_SUBTITULO_CENTRADO = 'IUTECP Subtítulo centrado'

# Parágrafo posterior a la cita (configurable desde contenido.py)
POST_CITA_TEXTO_DEF = (
    "De acuerdo a la cita previa, se comprende la relevancia del control sistemático "
    "y la inmutabilidad de los registros en los departamentos estratégicos de la empresa."
)

# ================================================================
#  FUNCIONES AUXILIARES DE FORMATO (NORMATIVA IUTECP 2025)
# ================================================================

def setup_iutecp_document():
    """Configura el documento base: Tamaño Carta, Márgenes Art. 6, Fuente Art. 5"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = PAG_ANCHO
    section.page_height = PAG_ALTO
    section.top_margin = MARGEN_SUP
    section.bottom_margin = MARGEN_INF
    section.left_margin = MARGEN_IZQ      # Encuadernación
    section.right_margin = MARGEN_DER

    style = doc.styles['Normal']
    style.font.name = FUENTE
    style.font.size = Pt(TAMANO_BASE)
    style.paragraph_format.line_spacing = INTERLINEADO
    style.paragraph_format.space_after = Pt(0)

    estilos_titulo = (
        (ESTILO_CAPITULO, WD_ALIGN_PARAGRAPH.CENTER, 0),
        (ESTILO_PRELIMINAR, WD_ALIGN_PARAGRAPH.CENTER, 0),
        (ESTILO_TITULO_NIVEL2, WD_ALIGN_PARAGRAPH.LEFT, 1),
        (ESTILO_TITULO_NIVEL3, WD_ALIGN_PARAGRAPH.LEFT, 2),
        (ESTILO_SUBTITULO_CENTRADO, WD_ALIGN_PARAGRAPH.CENTER, 1),
    )
    for nombre, alineacion, nivel_esquema in estilos_titulo:
        estilo = doc.styles.add_style(nombre, WD_STYLE_TYPE.PARAGRAPH)
        estilo.base_style = style
        estilo.font.name = FUENTE
        estilo.font.size = Pt(TAMANO_BASE)
        estilo.font.bold = True
        estilo.paragraph_format.alignment = alineacion
        estilo.paragraph_format.keep_with_next = True
        estilo.paragraph_format.keep_together = True
        p_pr = estilo.element.get_or_add_pPr()
        outline = p_pr.find(qn('w:outlineLvl'))
        if outline is None:
            outline = OxmlElement('w:outlineLvl')
            p_pr.append(outline)
        outline.set(qn('w:val'), str(nivel_esquema))

    update_fields = doc.settings.element.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        doc.settings.element.append(update_fields)
    update_fields.set(qn('w:val'), 'true')

    return doc

def _ruta_desde_generador(ruta):
    """Resuelve rutas configuradas sin depender del directorio de ejecución."""
    if not ruta or os.path.isabs(ruta):
        return ruta
    return os.path.join(BASE_DIR, ruta)

def _numero_romano(numero):
    valores = (
        (1000, 'm'), (900, 'cm'), (500, 'd'), (400, 'cd'),
        (100, 'c'), (90, 'xc'), (50, 'l'), (40, 'xl'),
        (10, 'x'), (9, 'ix'), (5, 'v'), (4, 'iv'), (1, 'i'),
    )
    resultado = []
    for valor, simbolo in valores:
        while numero >= valor:
            resultado.append(simbolo)
            numero -= valor
    return ''.join(resultado)

def _mantener_con_siguiente(parrafo):
    parrafo.paragraph_format.keep_with_next = True
    parrafo.paragraph_format.keep_together = True
    return parrafo

def _repetir_encabezado(fila):
    """Marca una fila para repetirla cuando una tabla continúa en otra página."""
    tr_pr = fila._tr.get_or_add_trPr()
    if tr_pr.find(qn('w:tblHeader')) is None:
        tr_pr.append(OxmlElement('w:tblHeader'))

def _agregar_fuente(doc, fuente, centrada=False):
    if not str(fuente or '').strip():
        raise ValueError("Todo cuadro, gráfico o anexo debe declarar una fuente.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centrada else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    etiqueta = p.add_run('Fuente: ')
    etiqueta.font.name = FUENTE
    etiqueta.font.size = Pt(TAMANO_TABLA)
    etiqueta.font.italic = True
    texto = p.add_run(str(fuente).removeprefix('Fuente: ').strip())
    texto.font.name = FUENTE
    texto.font.size = Pt(TAMANO_TABLA)
    return p

def agregar_parrafo_normado(doc, texto, cursiva=False, sangria=True):
    """Párrafo justificado, 1.5 interlineado, sangría 1.25cm (Art. 7, 8)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.line_spacing = INTERLINEADO
    p_format.first_line_indent = SANGRIA_LINEA if sangria else Cm(0)
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)

    run = p.add_run(texto)
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_BASE)
    run.font.italic = cursiva
    return p

def agregar_parrafo_resumen(doc, texto):
    """Agrega el cuerpo del resumen con interlineado sencillo y máximo normativo."""
    palabras = str(texto).split()
    if len(palabras) > RESUMEN_MAX_PALABRAS:
        raise ValueError(
            f"El resumen excede el máximo de {RESUMEN_MAX_PALABRAS} palabras "
            f"({len(palabras)})."
        )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.line_spacing = 1.0
    p_format.first_line_indent = SANGRIA_LINEA
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

    run = p.add_run(str(texto))
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_BASE)
    return p

def agregar_item_lista(doc, numero, texto, negrita_inicio=""):
    """Agrega un elemento enumerado con sangría de primera línea (Art. 10)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.line_spacing = INTERLINEADO
    p_format.first_line_indent = SANGRIA_LINEA

    run_num = p.add_run(f"{numero}. ")
    run_num.font.name = FUENTE
    run_num.font.size = Pt(TAMANO_BASE)

    if negrita_inicio:
        run_bold = p.add_run(f"{negrita_inicio}: ")
        run_bold.font.name = FUENTE
        run_bold.font.size = Pt(TAMANO_BASE)
        run_bold.font.bold = True

    run_text = p.add_run(texto)
    run_text.font.name = FUENTE
    run_text.font.size = Pt(TAMANO_BASE)
    return p

def agregar_titulo_nivel2(doc, texto, bookmark_id=None):
    """Nivel 2: Alineado a la izquierda, Negrita (Art. 9 - IUTECP)"""
    p = doc.add_paragraph(style=ESTILO_TITULO_NIVEL2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = ESP_DOBLE
    p.paragraph_format.space_after = ESP_DOBLE
    p.paragraph_format.line_spacing = INTERLINEADO
    _mantener_con_siguiente(p)

    run = p.add_run(texto)
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_BASE)
    run.font.bold = True
    if bookmark_id:
        _agregar_bookmark(p, bookmark_id)
    return p

def agregar_titulo_nivel3(doc, texto, bookmark_id=None):
    """Nivel 3: Alineado a la izquierda, Negrita y Cursiva (Art. 9 - IUTECP)"""
    p = doc.add_paragraph(style=ESTILO_TITULO_NIVEL3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = ESP_SENCILLO
    p.paragraph_format.space_after = ESP_SENCILLO
    p.paragraph_format.line_spacing = INTERLINEADO
    p.paragraph_format.first_line_indent = SANGRIA_LINEA
    _mantener_con_siguiente(p)

    run = p.add_run(texto)
    run.font.name = FUENTE
    run.font.size = Pt(TAMANO_BASE)
    run.font.bold = True
    run.font.italic = True
    if bookmark_id:
        _agregar_bookmark(p, bookmark_id)
    return p

def agregar_titulo_nivel4(doc, texto, texto_parrafo=''):
    """
    Nivel 4: Izquierda · Negrita · Sangría ½ pulgada (1.27 cm) · Punto final
    El encabezado termina en punto y el texto del párrafo continúa en la misma línea.
    (Art. 9 - IUTECP)
    Parámetros:
        texto         : El título del nivel 4 (sin punto, se agrega automáticamente).
        texto_parrafo : Texto que continúa en la misma línea después del título.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = ESP_SENCILLO
    p.paragraph_format.space_after = ESP_SENCILLO
    p.paragraph_format.line_spacing = INTERLINEADO
    p.paragraph_format.left_indent = SANGRIA_NIVEL45
    p.paragraph_format.first_line_indent = Pt(0)

    # Encabezado en negrita con punto final
    run_header = p.add_run(texto.rstrip('.') + '. ')
    run_header.font.name = FUENTE
    run_header.font.size = Pt(TAMANO_BASE)
    run_header.font.bold = True

    # Texto del párrafo en la misma línea (formato normal)
    if texto_parrafo:
        run_body = p.add_run(texto_parrafo)
        run_body.font.name = FUENTE
        run_body.font.size = Pt(TAMANO_BASE)
    return p

def agregar_titulo_nivel5(doc, texto, texto_parrafo=''):
    """
    Nivel 5: Izquierda · Negrita · Cursiva · Sangría ½ pulgada (1.27 cm) · Punto final
    El encabezado termina en punto y el texto del párrafo continúa en la misma línea.
    (Art. 9 - IUTECP)
    Parámetros:
        texto         : El título del nivel 5 (sin punto, se agrega automáticamente).
        texto_parrafo : Texto que continúa en la misma línea después del título.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Pt(0)

    # Encabezado en negrita + cursiva con punto final
    run_header = p.add_run(texto.rstrip('.') + '. ')
    run_header.font.name = 'Times New Roman'
    run_header.font.size = Pt(12)
    run_header.font.bold = True
    run_header.font.italic = True

    # Texto del párrafo en la misma línea (formato normal)
    if texto_parrafo:
        run_body = p.add_run(texto_parrafo)
        run_body.font.name = 'Times New Roman'
        run_body.font.size = Pt(12)
    return p

def _config_seccion(sec, sup=None, ocultar_primera_pagina=False):
    """Aplica márgenes y, solo cuando corresponde, oculta el pie de la primera página."""
    sec.top_margin = sup if sup is not None else MARGEN_SUP
    sec.bottom_margin = MARGEN_INF
    sec.left_margin = MARGEN_IZQ
    sec.right_margin = MARGEN_DER
    sec.different_first_page_header_footer = ocultar_primera_pagina
    v_align = sec._sectPr.find(qn('w:vAlign'))
    if v_align is not None:
        sec._sectPr.remove(v_align)

def iniciar_capitulo(doc, numero_romano, titulo, bookmark_id=None):
    """
    Crea un nuevo capítulo con espacio extra de 2cm en el título (simula Art. 6),
    centrado, negrita, mayúsculas (Art. 9).
    """
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _config_seccion(sec, ocultar_primera_pagina=True)

    # Título CAPÍTULO X — con espacio extra para simular margen superior de 5cm
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(57)
    p1.paragraph_format.space_after = ESP_DOBLE
    _mantener_con_siguiente(p1)
    r1 = p1.add_run(f"CAPÍTULO {numero_romano}")
    r1.font.name = FUENTE
    r1.font.size = Pt(TAMANO_BASE)
    r1.font.bold = True

    # Título del Capítulo
    p2 = doc.add_paragraph(style=ESTILO_CAPITULO)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = ESP_DOBLE
    r2 = p2.add_run(titulo.upper())
    r2.font.name = FUENTE
    r2.font.size = Pt(TAMANO_BASE)
    r2.font.bold = True
    if bookmark_id:
        _agregar_bookmark(p2, bookmark_id)

def iniciar_seccion_preliminar(doc, titulo, bookmark_id=None, ocultar_primera_pagina=False):
    """Para secciones preliminares que inician en página nueva (Art. 9, 12)"""
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _config_seccion(sec, ocultar_primera_pagina=ocultar_primera_pagina)

    p = doc.add_paragraph(style=ESTILO_PRELIMINAR)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(57)
    p.paragraph_format.space_after = ESP_DOBLE
    _mantener_con_siguiente(p)
    r = p.add_run(titulo.upper())
    r.font.name = FUENTE
    r.font.size = Pt(TAMANO_BASE)
    r.font.bold = True
    if bookmark_id:
        _agregar_bookmark(p, bookmark_id)

def iniciar_seccion_resumen(doc, contenido, bookmark_id=None):
    """Construye la página de resumen con membrete, autor y fecha."""
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _config_seccion(sec)

    for linea in contenido.MEMBRETE:
        p_membrete = doc.add_paragraph()
        p_membrete.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_membrete.paragraph_format.line_spacing = 1.0
        p_membrete.paragraph_format.space_after = Pt(0)
        run_membrete = p_membrete.add_run(linea)
        run_membrete.font.name = FUENTE
        run_membrete.font.size = Pt(TAMANO_BASE)
        run_membrete.font.bold = True

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.line_spacing = 1.0
    p_titulo.paragraph_format.space_after = Pt(6)
    run_titulo = p_titulo.add_run(contenido.TITULO_PROYECTO.upper())
    run_titulo.font.name = FUENTE
    run_titulo.font.size = Pt(TAMANO_BASE)
    run_titulo.font.bold = True

    fecha = getattr(contenido, 'FECHA_LUGAR', '')
    fecha_resumen = fecha.split(',', 1)[1].strip() if ',' in fecha else fecha
    p_autor = doc.add_paragraph()
    p_autor.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_autor.paragraph_format.left_indent = Cm(0)
    p_autor.paragraph_format.line_spacing = 1.0
    p_autor.paragraph_format.space_after = Pt(6)
    for indice, texto in enumerate((
        f"Autor: {contenido.NOMBRE_PASANTE}",
        f"C.I.: {contenido.CI_PASANTE}",
        fecha_resumen,
    )):
        run_autor = p_autor.add_run(texto)
        run_autor.font.name = FUENTE
        run_autor.font.size = Pt(TAMANO_BASE)
        if indice == 0:
            run_autor.font.bold = True
        if indice < 2:
            run_autor.add_break()

    p_resumen = doc.add_paragraph()
    p_resumen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_resumen.paragraph_format.line_spacing = 1.0
    p_resumen.paragraph_format.space_before = Pt(6)
    p_resumen.paragraph_format.space_after = Pt(6)
    run_resumen = p_resumen.add_run("RESUMEN")
    run_resumen.font.name = FUENTE
    run_resumen.font.size = Pt(TAMANO_BASE)
    run_resumen.font.bold = True
    if bookmark_id:
        _agregar_bookmark(p_resumen, bookmark_id)

def agregar_cita_larga(doc, texto, cita):
    """
    Citas textuales de más de 40 palabras:
    Párrafo separado, sangría 1.25cm (5 espacios) a ambos lados, interlineado sencillo,
    sin comillas, distancia de 3 espacios de separación del párrafo anterior (Art. 7, 22).
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = SANGRIA_CITA
    p.paragraph_format.right_indent = SANGRIA_CITA
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = ESP_CITA_ANTES
    p.paragraph_format.space_after = ESP_CITA_DESPUES

    run_t = p.add_run(texto)
    run_t.font.name = FUENTE
    run_t.font.size = Pt(TAMANO_BASE)

    run_c = p.add_run(f" {cita}")
    run_c.font.name = FUENTE
    run_c.font.size = Pt(TAMANO_BASE)

def agregar_referencia(doc, referencia):
    """
    Entrada en referencias bibliográficas (Art. 7, 25 IUTECP):
    - Interlineado sencillo dentro de cada entrada.
    - Sangría francesa de 3 espacios (~0.75cm) hacia la derecha.
    - Entre una referencia y otra: dos (2) espacios sencillos de separación.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = SANGRIA_REF
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.line_spacing = 1.0
    # 2 espacios sencillos de separación entre referencias (Art. 25)
    p.paragraph_format.space_after = ESP_DOBLE
    p.paragraph_format.space_before = Pt(0)

    if isinstance(referencia, dict):
        partes = (
            (referencia.get('titulo', ''), True),
            (referencia.get('texto', referencia.get('detalle', '')), False),
        )
    elif isinstance(referencia, (list, tuple)):
        partes = [(str(parte), indice == 0) for indice, parte in enumerate(referencia)]
    else:
        texto_referencia = str(referencia)
        coincidencia = re.match(
            r'^(.*?\(\d{4}[a-z]?\)\.\s+)(.+?)(\.\s+.+)$',
            texto_referencia,
        )
        if coincidencia:
            partes = [
                (coincidencia.group(1), False),
                (coincidencia.group(2), True),
                (coincidencia.group(3), False),
            ]
        else:
            partes = [(texto_referencia, False)]
    for texto, negrita in partes:
        run = p.add_run(texto)
        run.font.name = FUENTE
        run.font.size = Pt(TAMANO_BASE)
        run.font.bold = negrita
    return p

def set_cell_border(cell, **kwargs):
    """Permite definir bordes de celda personalizados mediante XML"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        val = kwargs.get(side, {'val': 'single', 'sz': '4', 'color': '000000'})
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   val.get('val', 'single'))
        el.set(qn('w:sz'),    val.get('sz', '4'))
        el.set(qn('w:color'), val.get('color', '000000'))
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_cell_shading(cell, fill='4472C4'):
    """Establece el color de fondo de una celda mediante XML"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)

def set_table_fixed_layout(tabla):
    """Establece layout de tabla fijo en XML"""
    tbl = tabla._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

def set_cell_width(cell, width_cm):
    """Establece el ancho de una celda en XML (en dxa/twips)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    twips = int(width_cm * 567)
    tcW.set(qn('w:w'), str(twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_table_grid_widths_xml(tabla, lista_anchos_cm):
    """Aplica el gridCol XML para garantizar renderizado exacto en LibreOffice/Word"""
    tbl = tabla._tbl
    tblGrid = tbl.tblGrid
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.append(tblGrid)
    else:
        for existing in tblGrid.findall(qn('w:gridCol')):
            tblGrid.remove(existing)

    for ancho_cm in lista_anchos_cm:
        gridCol = OxmlElement('w:gridCol')
        twips = int(ancho_cm * 567)
        gridCol.set(qn('w:w'), str(twips))
        tblGrid.append(gridCol)

def aplicar_formato_tabla_xml(tabla, lista_anchos_cm):
    """Aplica el ancho fijo a la tabla y a cada celda de forma inmutable"""
    set_table_fixed_layout(tabla)
    set_table_grid_widths_xml(tabla, lista_anchos_cm)
    for fila in tabla.rows:
        for col_idx, ancho in enumerate(lista_anchos_cm):
            set_cell_width(fila.cells[col_idx], ancho)

def _celda(cell, texto, negrita=False, centrado=False, tamaño=Pt(TAMANO_TABLA), color_texto=COLOR_TEXTO_OSCURO):
    """Inserta texto formateado en una celda de tabla limpiando párrafos vacíos"""
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centrado else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = 1.0

    run = p.add_run(texto)
    run.font.name  = FUENTE
    run.font.size  = tamaño
    run.font.bold  = negrita
    run.font.color.rgb = RGBColor(
        int(color_texto[0:2], 16), int(color_texto[2:4], 16), int(color_texto[4:6], 16)
    )

def agregar_titulo_cuadro(doc, texto, bookmark_id=None):
    """
    Inserta el título de un cuadro ANTES de la tabla (Art. 6, 8, 13 IUTECP).
    - Título arriba del cuadro.
    - El número del cuadro va en negrita, la descripción en cursiva (Art. 6).
    - Espacio doble antes y después (Art. 8).
    Formato esperado del texto: 'Cuadro X. Descripción del cuadro'
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Espacio doble antes y después de cuadros titulados (Art. 8)
    p.paragraph_format.space_before = ESP_DOBLE
    p.paragraph_format.space_after = ESP_SENCILLO
    p.paragraph_format.line_spacing = INTERLINEADO
    _mantener_con_siguiente(p)

    # Separar 'Cuadro X' de la descripción para aplicar formatos distintos
    partes = texto.split('. ', 1)
    if len(partes) == 2:
        # "Cuadro X" en negrita
        run_num = p.add_run(partes[0] + '. ')
        run_num.font.name = FUENTE
        run_num.font.size = Pt(TAMANO_TABLA)
        run_num.font.bold = True
        # Descripción en cursiva (Art. 6)
        run_desc = p.add_run(partes[1])
        run_desc.font.name = FUENTE
        run_desc.font.size = Pt(TAMANO_TABLA)
        run_desc.font.italic = True
    else:
        run = p.add_run(texto)
        run.font.name = FUENTE
        run.font.size = Pt(TAMANO_TABLA)
        run.font.bold = True
    if bookmark_id:
        _agregar_bookmark(p, bookmark_id)

def agregar_tabla_planificacion(doc, datos, titulo_cuadro=None, bookmark_id=None, fuente=None):
    """Genera la tabla de planificación de objetivos con 5 columnas"""
    if titulo_cuadro:
        agregar_titulo_cuadro(doc, titulo_cuadro, bookmark_id=bookmark_id)

    ENCABEZADOS = ['Objetivo', 'Variable', 'Actividades', 'Técnica', 'Instrumento']
    ANCHOS_CM   = [2.9, 2.9, 2.9, 2.9, 2.99]

    tabla = doc.add_table(rows=1 + len(datos), cols=5)
    tabla.style = 'Table Grid'
    aplicar_formato_tabla_xml(tabla, ANCHOS_CM)
    _repetir_encabezado(tabla.rows[0])

    for col, enc in enumerate(ENCABEZADOS):
        cell = tabla.cell(0, col)
        set_cell_shading(cell, COLOR_ENCABEZADO)
        _celda(cell, enc, negrita=True, centrado=True, tamaño=Pt(TAMANO_TABLA), color_texto=COLOR_TEXTO_CLARO)

    for fila, datos_fila in enumerate(datos, start=1):
        if len(datos_fila) == 5:
            contenidos = list(datos_fila)
        else:
            contenidos = list(datos_fila) + ['']
        fondo = COLOR_FILA_PAR if fila % 2 == 0 else COLOR_TEXTO_CLARO
        for col, texto in enumerate(contenidos):
            cell = tabla.cell(fila, col)
            set_cell_shading(cell, fondo)
            _celda(cell, texto, tamaño=Pt(TAMANO_TABLA))
    _agregar_fuente(doc, fuente)

def agregar_gantt(doc, semanas, titulo_cuadro=None, bookmark_id=None, fuente=None):
    """Genera la tabla Gantt con 2 filas de encabezado (meses y semanas)"""
    if titulo_cuadro:
        agregar_titulo_cuadro(doc, titulo_cuadro, bookmark_id=bookmark_id)

    num_sem = max(len(s[1]) for s in semanas)

    section = doc.sections[-1]
    ancho_util_emu = section.page_width - section.left_margin - section.right_margin
    ancho_util_cm = Emu(ancho_util_emu).cm

    meses = [("JUNIO", 4), ("JULIO", 4), ("AGOSTO", 2)]
    COL_SEM_CM = Pt(3 * TAMANO_TABLA_CHICO * 0.55).cm + 0.4
    min_act_cm = 3.0
    max_s_por_col = (ancho_util_cm - min_act_cm) / num_sem if num_sem > 0 else ancho_util_cm
    COL_SEM_CM = min(COL_SEM_CM, max_s_por_col)

    # Reserva ancho suficiente para que el nombre de un mes no se parta cuando
    # el último mes solo contiene una semana (por ejemplo, en cronogramas de 9 semanas).
    anchos_semanas = []
    col_start = 1
    for nombre, ncols in meses:
        if col_start > num_sem:
            break
        ncols = min(ncols, num_sem - col_start + 1)
        ancho_grupo = max(ncols * COL_SEM_CM, 2.0)
        anchos_semanas.extend([ancho_grupo / ncols] * ncols)
        col_start += ncols
    while len(anchos_semanas) < num_sem:
        anchos_semanas.append(COL_SEM_CM)
    COL_ACT_CM = ancho_util_cm - sum(anchos_semanas)

    tabla = doc.add_table(rows=2 + len(semanas), cols=1 + num_sem)
    tabla.style = 'Table Grid'

    anchos_gantt = [COL_ACT_CM] + anchos_semanas
    aplicar_formato_tabla_xml(tabla, anchos_gantt)
    _repetir_encabezado(tabla.rows[0])
    _repetir_encabezado(tabla.rows[1])

    # ── Fila 0: meses (JUNIO, JULIO, AGOSTO) ──
    cell_label = tabla.cell(0, 0)
    set_cell_shading(cell_label, COLOR_ENCABEZADO)
    _celda(cell_label, 'Semana', negrita=True, centrado=True, tamaño=Pt(TAMANO_TABLA), color_texto=COLOR_TEXTO_CLARO)

    col_start = 1
    for nombre, ncols in meses:
        if col_start > num_sem:
            break
        ncols = min(ncols, num_sem - col_start + 1)
        c1 = tabla.cell(0, col_start)
        c2 = tabla.cell(0, col_start + ncols - 1)
        merged = c1.merge(c2)
        set_cell_shading(merged, COLOR_ENCABEZADO)
        _celda(merged, nombre, negrita=True, centrado=True, tamaño=Pt(TAMANO_TABLA), color_texto=COLOR_TEXTO_CLARO)
        col_start += ncols

    # ── Fila 1: semanas (1..n) ──
    cell_act = tabla.cell(1, 0)
    set_cell_shading(cell_act, COLOR_ENCABEZADO)
    _celda(cell_act, 'Actividad', negrita=True, centrado=True, tamaño=Pt(TAMANO_TABLA), color_texto=COLOR_TEXTO_CLARO)

    semana_meses = [4, 4, 2]  # weeks per month
    col_actual = 1
    for ncols in semana_meses:
        for w in range(1, ncols + 1):
            if col_actual > num_sem:
                break
            cell_s = tabla.cell(1, col_actual)
            set_cell_shading(cell_s, COLOR_ENCABEZADO)
            _celda(cell_s, str(w), negrita=True, centrado=True, tamaño=Pt(TAMANO_TABLA_CHICO), color_texto=COLOR_TEXTO_CLARO)
            col_actual += 1

    # ── Filas de datos ──
    for fila, (desc, activas) in enumerate(semanas, start=2):
        fondo_fila = COLOR_FILA_PAR if fila % 2 == 0 else COLOR_TEXTO_CLARO
        cell_a = tabla.cell(fila, 0)
        set_cell_shading(cell_a, fondo_fila)
        _celda(cell_a, desc, tamaño=Pt(TAMANO_TABLA))

        for s in range(num_sem):
            cell_s = tabla.cell(fila, s + 1)
            activa = activas[s] if s < len(activas) else False
            set_cell_shading(cell_s, COLOR_GANTT_VERDE if activa else fondo_fila)
            _celda(cell_s, '✓' if activa else '', centrado=True, tamaño=Pt(TAMANO_TABLA_CHICO), color_texto=COLOR_TEXTO_CLARO if activa else COLOR_TEXTO_OSCURO)
    _agregar_fuente(doc, fuente)

def _insertar_campo_pagina(run, formato_pagina='PAGE'):
    """Inserta un campo de número de página con el formato especificado en un run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = formato_pagina
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def _establecer_tipo_numeracion_seccion(section, formato, inicio=None):
    """Configura el formato y, solo cuando corresponde, el inicio de página de una sección."""
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn('w:pgNumType'))
    if pg_num_type is None:
        pg_num_type = OxmlElement('w:pgNumType')
        sect_pr.append(pg_num_type)

    pg_num_type.set(qn('w:fmt'), formato)
    if inicio is None:
        pg_num_type.attrib.pop(qn('w:start'), None)
    else:
        pg_num_type.set(qn('w:start'), str(inicio))

def _limpiar_pie(pie):
    """Elimina campos heredados del pie y deja un único párrafo vacío."""
    for parrafo in list(pie.paragraphs):
        parrafo._element.getparent().remove(parrafo._element)
    pie.add_paragraph()

def agregar_numeracion_pie(doc, idx_inicio_arabigo):
    """
    Configura la paginación institucional usando propiedades de sección y campos PAGE:
    - la portada se cuenta sin mostrar número y la contraportada muestra ii;
    - todas las demás preliminares muestran romanos minúsculos, sin ocultar sus primeras páginas;
    - Introducción reinicia la secuencia arábiga en 1 y no imprime su primera página;
    - capítulos y Referencias ocultan únicamente su primera página;
    - la secuencia arábiga continúa hasta los anexos.
    """
    if idx_inicio_arabigo is None or not 0 <= idx_inicio_arabigo < len(doc.sections):
        raise ValueError("No se encontró una sección válida para iniciar la numeración arábiga.")

    for sec_idx, section in enumerate(doc.sections):
        formato = 'lowerRoman' if sec_idx < idx_inicio_arabigo else 'decimal'
        inicio = 1 if sec_idx in (0, idx_inicio_arabigo) else None
        _establecer_tipo_numeracion_seccion(section, formato, inicio=inicio)

        footer = section.footer
        footer.is_linked_to_previous = False
        _limpiar_pie(footer)

        # Las secciones marcadas como primera página distinta solo tienen pie regular
        # después de esa página; su pie de primera página debe quedar vacío.
        if section.different_first_page_header_footer:
            first_page_footer = section.first_page_footer
            first_page_footer.is_linked_to_previous = False
            _limpiar_pie(first_page_footer)

        if sec_idx == 0:
            continue

        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        # El formato lowerRoman/decimal lo determina w:pgNumType de la sección.
        _insertar_campo_pagina(run, ' PAGE ')

def buscar_imagen_por_numero(carpeta, numero, extensiones=None):
    """Busca una imagen por su nombre numérico en una carpeta específica"""
    if extensiones is None:
        extensiones = ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']
    carpeta = _ruta_desde_generador(carpeta)
    if not os.path.exists(carpeta):
        print(f"⚠ Advertencia: La carpeta '{carpeta}' no existe")
        return None
    numero_str = str(numero)
    for archivo in os.listdir(carpeta):
        nombre, ext = os.path.splitext(archivo)
        if nombre == numero_str and ext.lower()[1:] in extensiones:
            ruta_completa = os.path.join(carpeta, archivo)
            print(f"✓ Imagen encontrada: {ruta_completa}")
            return ruta_completa
    print(f"⚠ No se encontró imagen con número {numero} en {carpeta}")
    return None

def buscar_imagen_por_referencia(carpeta, referencia, extensiones=None):
    """Busca un anexo por nombre de archivo o conserva la búsqueda numérica."""
    if isinstance(referencia, str) and os.path.splitext(referencia)[1]:
        ruta = os.path.join(_ruta_desde_generador(carpeta), referencia)
        if os.path.isfile(ruta):
            print(f"✓ Imagen encontrada: {ruta}")
            return ruta
        print(f"⚠ No se encontró el archivo de anexo '{referencia}' en {carpeta}")
        return None
    return buscar_imagen_por_numero(carpeta, referencia, extensiones=extensiones)

def _calcular_tamano_anexo_proporcional(ruta_imagen, configuracion):
    """Devuelve una sola dimensión para insertar un anexo sin deformarlo."""
    # python-docx ya incluye lectores de encabezados para los formatos de imagen
    # que el generador inserta, por lo que no hace falta añadir Pillow.
    from docx.image.image import Image as DocxImage

    imagen = DocxImage.from_file(ruta_imagen)
    ancho_real_cm = float(imagen.width.cm)
    alto_real_cm = float(imagen.height.cm)
    if ancho_real_cm <= 0 or alto_real_cm <= 0:
        raise ValueError(f"Dimensiones inválidas para la imagen {ruta_imagen}")

    def limite_positivo_cm(valor):
        try:
            valor = float(valor)
        except (TypeError, ValueError):
            return None
        return valor if math.isfinite(valor) and valor > 0 else None

    max_ancho_cm = None
    max_alto_cm = None
    if isinstance(configuracion, dict):
        max_ancho_cm = limite_positivo_cm(configuracion.get('width_cm'))
        max_alto_cm = limite_positivo_cm(configuracion.get('height_cm'))
    elif configuracion is not None:
        # Formato anterior: el cuarto elemento era únicamente el alto máximo.
        max_alto_cm = limite_positivo_cm(configuracion)

    if max_ancho_cm is None and max_alto_cm is None:
        # Conserva el ancho histórico por defecto, pero como límite para no
        # ampliar imágenes que ya caben.
        max_ancho_cm = 14.0

    factor_ancho = max_ancho_cm / ancho_real_cm if max_ancho_cm is not None else float('inf')
    factor_alto = max_alto_cm / alto_real_cm if max_alto_cm is not None else float('inf')
    factor = min(1.0, factor_ancho, factor_alto)
    ancho_final_cm = ancho_real_cm * factor
    alto_final_cm = alto_real_cm * factor

    # add_picture conserva la proporción cuando recibe una sola dimensión.
    # Elegir la dimensión limitante evita confiar en dos valores
    # independientes dentro del XML del DOCX.
    if max_alto_cm is not None and (max_ancho_cm is None or factor_alto <= factor_ancho):
        return {'height': Cm(alto_final_cm)}
    return {'width': Cm(ancho_final_cm)}

def agregar_imagen(doc, ruta_imagen, titulo, ancho=Cm(12), fuente=None, bookmark_id=None):
    """
    Agrega una imagen (gráfico) y su título descriptivo DEBAJO (Art. 13 IUTECP).
    - Espacio doble antes y después de gráficos titulados (Art. 8).
    - El número del gráfico en cursiva, la descripción en negrita (Art. 6).
    - Opcionalmente agrega la línea 'Fuente:' (Art. 13).
    Formato esperado del titulo: 'Gráfico X. Descripción del gráfico'
    """
    if not ruta_imagen or not os.path.exists(ruta_imagen):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[IMAGEN NO ENCONTRADA: {titulo}]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        return

    section = doc.sections[-1]
    max_width = Emu(section.page_width - section.left_margin - section.right_margin)
    if ancho > max_width:
        ancho = max_width

    # Espacio doble antes del gráfico (Art. 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    _mantener_con_siguiente(p)
    run = p.add_run()
    try:
        run.add_picture(ruta_imagen, width=ancho)
    except Exception as e:
        print(f"❌ Error al agregar imagen {ruta_imagen}: {e}")
        return

    # Título del gráfico debajo de la imagen (Art. 13)
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_before = Pt(6)
    p_titulo.paragraph_format.space_after = Pt(24)
    _mantener_con_siguiente(p_titulo)

    # Separar 'Gráfico X' de la descripción para formatos distintos
    partes = titulo.split('. ', 1)
    if len(partes) == 2:
        # "Gráfico X." en cursiva
        run_num = p_titulo.add_run(partes[0] + '. ')
        run_num.font.name = 'Times New Roman'
        run_num.font.size = Pt(TAMANO_TABLA)
        run_num.font.italic = True
        # Descripción en negrita
        run_desc = p_titulo.add_run(partes[1])
        run_desc.font.name = 'Times New Roman'
        run_desc.font.size = Pt(TAMANO_TABLA)
        run_desc.font.bold = True
    else:
        run_titulo = p_titulo.add_run(titulo)
        run_titulo.font.name = 'Times New Roman'
        run_titulo.font.size = Pt(TAMANO_TABLA)
        run_titulo.font.bold = True
    if bookmark_id:
        _agregar_bookmark(p_titulo, bookmark_id)

    _agregar_fuente(doc, fuente, centrada=True)

# ================================================================
# CONSTRUCCIÓN DEL DOCUMENTO (PORTADA COMPATIBLE LIBREOFFICE)
# ================================================================
def construir_portada(doc, solo_autor=False, idx_seccion=0, bookmark_id=None):
    """
    Construye la portada distribuyendo los 4 bloques de forma proporcional
    al área útil de la página, sin valores fijos de puntos.

    Las posiciones de cada bloque se definen como porcentajes del área útil:
      POS_TITULO: dónde empieza el título (defecto 42% = ligeramente sobre el centro)
      POS_AUTOR : dónde empiezan los datos del autor (defecto 67% = tercio inferior)
      POS_FECHA : dónde empieza la fecha (defecto 90% = cerca del margen inferior)

    solo_autor: si es True, solo muestra el nombre y CI del autor (sin tutores),
                usado en la primera página (portada). La contraportada muestra todo.
    """
    # ------------------------------------------------------------------
    # 1. Leer dimensiones reales de la página (en puntos, 1 pt = 1/72 in)
    # ------------------------------------------------------------------
    section = doc.sections[idx_seccion]
    EMU_PER_PT = 12700  # 1 pt = 12700 EMU

    usable_h = (section.page_height - section.top_margin - section.bottom_margin) / EMU_PER_PT
    usable_w = (section.page_width  - section.left_margin - section.right_margin ) / EMU_PER_PT

    # ------------------------------------------------------------------
    # 2. Estimar la altura de cada bloque en puntos
    #    (fuente 12pt × interlineado 1.15 = ~13.8pt por línea)
    # ------------------------------------------------------------------
    LINE_PT = 12 * 1.15

    membrete_lines = len(c.MEMBRETE)

    lineas_autor_completo = c.AUTOR_DATOS if isinstance(c.AUTOR_DATOS, list) else [c.AUTOR_DATOS]
    if solo_autor:
        # Portada: solo nombre y CI (todo antes del primer string vacío)
        lineas_autor = []
        for linea in lineas_autor_completo:
            if not linea.strip():
                break
            lineas_autor.append(linea)
    else:
        lineas_autor = lineas_autor_completo
    autor_lines  = sum(1 for l in lineas_autor if l.strip())

    # Times New Roman 12pt: ancho promedio ~6pt por carácter
    chars_per_line = max(1, int(usable_w / 6.0))
    titulo_lines   = max(1, math.ceil(len(c.TITULO_PROYECTO) / chars_per_line))

    h_membrete = membrete_lines * LINE_PT
    h_titulo   = titulo_lines   * LINE_PT
    h_autor    = autor_lines    * LINE_PT
    h_fecha    = LINE_PT

    # ------------------------------------------------------------------
    # 3. Calcular gaps en puntos entre bloques
    #    El gap datos-fecha es la mitad del gap normal
    # ------------------------------------------------------------------
    h_membrete = membrete_lines * LINE_PT
    h_titulo   = titulo_lines   * LINE_PT
    h_autor    = autor_lines    * LINE_PT
    h_fecha    = LINE_PT

    # Logo height estimate (solo portada, 3cm de ancho ≈ 80pt de alto)
    logo_path = _ruta_desde_generador(os.path.join("compartido", "iutecp.png"))
    tiene_logo = solo_autor and os.path.exists(logo_path)
    h_logo = 80.0 if tiene_logo else 0

    h_contenido = h_membrete + h_titulo + h_autor + h_fecha
    gap = max(12.0, (usable_h - h_contenido - 30) / 2.7)
    gap_mt = gap * 1.2   # membrete-título más grande
    gap_td = gap          # título-datos
    gap_df = max(6.0, gap * 0.5)  # datos-fecha mitad

    if tiene_logo:
        logo_mitad = max(0, gap_mt / 4.0)
        logo_resto = max(0, gap_mt - logo_mitad - h_logo)
        before_titulo = 0
    else:
        logo_mitad = 0
        logo_resto = 0
        before_titulo = gap_mt
    before_autor  = gap_td
    before_fecha  = gap_df

    # La contraportada necesita una distribución más compacta porque incluye
    # la descripción institucional y los datos de ambos tutores.
    if not solo_autor:
        before_titulo = 64
        before_autor = 105
        before_fecha = 40

    # ------------------------------------------------------------------
    # 4. Renderizar cada bloque
    # ------------------------------------------------------------------

    # BLOQUE 1: MEMBRETE (Alineado al margen superior)
    p_memb = doc.add_paragraph()
    p_memb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_memb.paragraph_format.line_spacing = 1.0 if not solo_autor else 1.15
    p_memb.paragraph_format.space_before = Pt(0)
    p_memb.paragraph_format.space_after  = Pt(0)
    for i, linea in enumerate(c.MEMBRETE):
        r = p_memb.add_run(linea)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = True
        if i < len(c.MEMBRETE) - 1:
            r.add_break()
    if bookmark_id:
        _agregar_bookmark(p_memb, bookmark_id)

    # BLOQUE 1.5: LOGO IUTECP (solo portada, debajo del membrete)
    if solo_autor and os.path.exists(logo_path):
        pic = doc.add_picture(logo_path, width=Cm(3.0))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last_p.paragraph_format.space_before = Pt(logo_mitad)
        last_p.paragraph_format.space_after  = Pt(logo_resto)

    # BLOQUE 2: TÍTULO DEL PROYECTO (Centrado proporcionalmente)
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.line_spacing = 1.15
    p_titulo.paragraph_format.space_before = Pt(before_titulo)
    p_titulo.paragraph_format.space_after  = Pt(0)
    r_title = p_titulo.add_run(c.TITULO_PROYECTO)
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(12)
    r_title.font.bold = True

    if not solo_autor:
        p_desc = doc.add_paragraph()
        p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_desc.paragraph_format.line_spacing = 1.15
        p_desc.paragraph_format.space_before = Pt(22)
        p_desc.paragraph_format.space_after = Pt(0)
        r_desc = p_desc.add_run(
            "Informe de pasantías para obtener el título de Técnico Superior Universitario "
            f"en la especialidad de: {getattr(c, 'ESPECIALIDAD', '')}"
        )
        r_desc.font.name = 'Times New Roman'
        r_desc.font.size = Pt(12)
        r_desc.font.bold = True

    # BLOQUE 3: DATOS DEL AUTOR Y TUTORES
    if solo_autor:
        # Portada: todo el bloque alineado a la derecha
        p_datos = doc.add_paragraph()
        p_datos.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_datos.paragraph_format.line_spacing = 1.15
        p_datos.paragraph_format.space_before = Pt(before_autor)
        p_datos.paragraph_format.space_after  = Pt(0)
        for i, linea in enumerate(lineas_autor):
            r = p_datos.add_run(linea)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.font.bold = True
            if i < len(lineas_autor) - 1:
                r.add_break()
    else:
        # Contraportada: tutores a la izquierda, autor a la derecha, misma fila
        split_idx = len(lineas_autor)
        for i, linea in enumerate(lineas_autor):
            if not linea.strip():
                split_idx = i
                break
        autor_block = lineas_autor[:split_idx]
        tutor_start = split_idx
        while tutor_start < len(lineas_autor) and not lineas_autor[tutor_start].strip():
            tutor_start += 1
        tutor_block = lineas_autor[tutor_start:]

        # Espaciador antes de la tabla (para posicionar a la altura de POS_AUTOR)
        p_before = doc.add_paragraph()
        p_before.paragraph_format.space_before = Pt(before_autor)
        p_before.paragraph_format.space_after = Pt(0)
        p_before.paragraph_format.line_spacing = Pt(1)

        # Tabla invisible de 1 fila × 2 columnas
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False

        # Quitar bordes
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            borders.append(border)
        tblPr.append(borders)

        # Cell 0: tutores (izquierda)
        cell_tutor = table.rows[0].cells[0]
        cell_tutor.width = Cm(7.29)
        cell_tutor.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p_t = cell_tutor.paragraphs[0]
        p_t.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_t.paragraph_format.line_spacing = 1.0
        for i, linea in enumerate(tutor_block):
            run = p_t.add_run(linea)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            if i < len(tutor_block) - 1:
                run.add_break()

        # Cell 1: autor (derecha)
        cell_autor = table.rows[0].cells[1]
        cell_autor.width = Cm(7.30)
        cell_autor.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p_a = cell_autor.paragraphs[0]
        p_a.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_a.paragraph_format.line_spacing = 1.0
        for i, linea in enumerate(autor_block):
            run = p_a.add_run(linea)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            if i < len(autor_block) - 1:
                run.add_break()

    # BLOQUE 4: CIUDAD Y FECHA (Proporcional al margen inferior)
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie.paragraph_format.line_spacing = 1.15
    p_pie.paragraph_format.space_before = Pt(before_fecha)
    p_pie.paragraph_format.space_after  = Pt(0)
    r_pie = p_pie.add_run(c.FECHA_LUGAR)
    r_pie.font.name = 'Times New Roman'
    r_pie.font.size = Pt(12)

def agregar_parada_tabulacion_puntos(parrafo, posicion_emu):
    """Agrega una parada de tabulación derecha con puntos de relleno usando la API oficial de python-docx."""
    parrafo.paragraph_format.tab_stops.add_tab_stop(
        Emu(posicion_emu),
        alignment=WD_TAB_ALIGNMENT.RIGHT,
        leader=WD_TAB_LEADER.DOTS
    )

def agregar_tabulaciones_lista(parrafo, pos_tab_texto_emu, pos_tab_pag_emu):
    """Agrega dos paradas de tabulación oficiales: izquierda (texto) y derecha con puntos (página)."""
    parrafo.paragraph_format.tab_stops.add_tab_stop(
        Emu(pos_tab_texto_emu),
        alignment=WD_TAB_ALIGNMENT.LEFT
    )
    parrafo.paragraph_format.tab_stops.add_tab_stop(
        Emu(pos_tab_pag_emu),
        alignment=WD_TAB_ALIGNMENT.RIGHT,
        leader=WD_TAB_LEADER.DOTS
    )

def agregar_tabulacion_derecha(parrafo, pos_tab_pag_emu):
    """Agrega una única parada de tabulación derecha oficial (para alinear pp. en la cabecera)."""
    parrafo.paragraph_format.tab_stops.add_tab_stop(
        Emu(pos_tab_pag_emu),
        alignment=WD_TAB_ALIGNMENT.RIGHT
    )

def agregar_fila_indice_general_nativa(
    doc, titulo, pagina='', sangria_cm=0.0, negrita=False, bookmark_id=None,
):
    """Agrega una línea del índice general usando tabulaciones nativas de Word para alinear al extremo derecho de forma absoluta.

    Si se pasa `bookmark_id`, en lugar del número `pagina` se inserta un campo
    PAGEREF que Word/LibreOffice rellena automáticamente con la página real
    del título marcado por ese bookmark.
    """
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(sangria_cm)
    p.paragraph_format.keep_together = True
    
    section = doc.sections[-1]
    ancho_util_emu = section.page_width - section.left_margin - section.right_margin
    # La posición del tabulador en Word es absoluta respecto a los márgenes, por lo que usamos directamente ancho_util_emu
    agregar_parada_tabulacion_puntos(p, ancho_util_emu)
    
    run_desc = p.add_run(f"{titulo}\t")
    run_desc.font.name = 'Times New Roman'
    run_desc.font.size = Pt(12)
    run_desc.font.bold = negrita
    
    if bookmark_id:
        _agregar_campo_pageref(
            p,
            bookmark_id,
            negrita=negrita,
        )
    else:
        run_pag = p.add_run(pagina)
        run_pag.font.name = 'Times New Roman'
        run_pag.font.size = Pt(12)
        run_pag.font.bold = negrita
    return p

def _agregar_bookmark(parrafo, bookmark_id):
    """Inserta un marcador (bookmark) XML oculto en el párrafo indicado.

    El bookmark rodea el contenido del párrafo, permitiendo que un campo
    PAGEREF del índice apunte a la página real donde cae este párrafo.
    """
    p_elem = parrafo._element
    ids_existentes = p_elem.getroottree().xpath(
        '//*[local-name()="bookmarkStart"]/@*[local-name()="id"]'
    )
    ids_numericos = [int(valor) for valor in ids_existentes if str(valor).isdigit()]
    bookmark_num = str(max(ids_numericos, default=-1) + 1)
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), bookmark_num)
    start.set(qn('w:name'), bookmark_id)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bookmark_num)
    # Insertar al inicio y al final del párrafo para rodear todo su contenido
    p_elem.insert(0, start)
    p_elem.append(end)

def _agregar_campo_pageref(parrafo, bookmark_id, negrita=False):
    """Inserta un campo PAGEREF que Word/LibreOffice evalúa al actualizar campos."""
    run = parrafo.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = negrita

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' PAGEREF {bookmark_id} \\h '

    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')

    # Texto de respaldo que se muestra si los campos no se actualizan
    t = OxmlElement('w:t')
    t.text = '?'

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._element.append(fldChar_begin)
    run._element.append(instrText)
    run._element.append(fldChar_sep)
    run._element.append(t)
    run._element.append(fldChar_end)

def agregar_fila_lista_preliminar_nativa(doc, col1_text, col2_text, col3_text, bookmark_id=None):
    """Agrega una entrada a una lista descriptiva (cuadro/gráfico/anexo) con sangría colgante y tabulación nativa absoluta.

    Si se pasa `bookmark_id`, la tercera columna (página) se reemplaza por un
    campo PAGEREF que Word/LibreOffice actualiza automáticamente.
    """
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.8)
    p.paragraph_format.first_line_indent = Cm(-1.8)
    p.paragraph_format.keep_together = True

    section = doc.sections[-1]
    ancho_util_emu = section.page_width - section.left_margin - section.right_margin
    # La posición del tabulador en Word es absoluta respecto a los márgenes, por lo que usamos directamente ancho_util_emu
    agregar_tabulaciones_lista(p, Cm(1.8).emu, ancho_util_emu)

    # Escribir número y descripción con formato
    run = p.add_run(f"{col1_text}\t{col2_text}\t")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    if bookmark_id:
        _agregar_campo_pageref(p, bookmark_id, negrita=False)
    else:
        run_pag = p.add_run(col3_text)
        run_pag.font.name = 'Times New Roman'
        run_pag.font.size = Pt(12)
    return p

# ================================================================
#  INSERCIÓN DE GRÁFICOS DATA-DRIVEN
# ================================================================
# Anclas válidas donde pueden insertarse gráficos:
#   "ubicacion"  -> tras la sección 1.1.7 Ubicación geográfica
#   "estructura" -> tras la sección 1.1.9 Estructura Organizativa
ANCLAS_VALIDAS = {"logo_empresa", "ubicacion", "estructura"}

def _resolver_ruta_imagen(carpeta_imagenes, cfg):
    """Resuelve una imagen por nombre explícito (`archivo`) o por número, manteniendo compatibilidad."""
    archivo = cfg.get("archivo") if isinstance(cfg, dict) else None
    if archivo:
        ruta = os.path.join(_ruta_desde_generador(carpeta_imagenes), archivo)
        if os.path.exists(ruta):
            return ruta
        print(f"⚠ No se encontró imagen '{archivo}' en {carpeta_imagenes}")
        return None
    numero = cfg.get("numero") if isinstance(cfg, dict) else None
    return buscar_imagen_por_numero(carpeta_imagenes, numero)

def _insertar_graficos_por_ancla(doc, carpeta_imagenes, ancla):
    """Inserta todos los gráficos de contenido.py marcados con `tras = ancla`."""
    graficos = getattr(c, 'GRAFICOS', [])
    for g in graficos:
        if g.get("tras") != ancla:
            continue
        numero = g.get("numero")
        titulo = g.get("titulo", f"Gráfico {numero}.")
        ancho = g.get("ancho_cm", 12)
        ruta = _resolver_ruta_imagen(carpeta_imagenes, g)
        bookmark_id = f"bm_grafico{numero}" if numero else None
        agregar_imagen(doc, ruta, titulo, ancho=Cm(ancho), fuente=g.get("fuente"), bookmark_id=bookmark_id)

def _insertar_logo_empresa(doc, carpeta_imagenes):
    """Inserta el logotipo configurado en GRAFICOS o mediante LOGO_EMPRESA."""
    logo_grafico = next(
        (g for g in getattr(c, 'GRAFICOS', []) if g.get("tras") == "logo_empresa"),
        None,
    )
    if logo_grafico:
        numero = logo_grafico.get("numero")
        ruta = _resolver_ruta_imagen(carpeta_imagenes, logo_grafico)
        agregar_imagen(
            doc,
            ruta,
            logo_grafico.get("titulo", f"Gráfico {numero}."),
            ancho=Cm(logo_grafico.get("ancho_cm", 12)),
            fuente=logo_grafico.get("fuente"),
            bookmark_id=f"bm_grafico{numero}" if numero else None,
        )
        return

    cfg = getattr(c, 'LOGO_EMPRESA', None)
    if not cfg:
        return
    if isinstance(cfg, str):
        cfg = {"archivo": cfg}
    ruta = _resolver_ruta_imagen(carpeta_imagenes, cfg)
    if not ruta:
        return
    ancho = Cm(cfg.get("ancho_cm", 4.0))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    try:
        p.add_run().add_picture(ruta, width=ancho)
    except Exception as e:
        print(f"❌ Error al agregar logo empresarial {ruta}: {e}")

def _texto_referencia(referencia):
    if isinstance(referencia, dict):
        return f"{referencia.get('titulo', '')}{referencia.get('texto', referencia.get('detalle', ''))}"
    if isinstance(referencia, (list, tuple)):
        return ''.join(str(parte) for parte in referencia)
    return str(referencia)

def _registro_cuadros(incluir_capitulo2=True):
    """Única fuente para títulos, numeración, marcadores y lista de cuadros."""
    titulos = [
        getattr(c, 'CUADRO_POBLACION_TITULO', CUADRO_POBLACION_TITULO_DEF),
    ]
    if incluir_capitulo2:
        titulos.extend((
            getattr(c, 'CUADRO_PLANIFICACION_TITULO', CUADRO_PLANIFICACION_TITULO_DEF),
            getattr(c, 'CUADRO_CRONOGRAMA_TITULO', CUADRO_CRONOGRAMA_TITULO_DEF),
        ))
    return [
        {
            'numero': numero,
            'titulo': titulo,
            'descripcion': str(titulo).split('. ', 1)[-1].rstrip('.'),
            'bookmark': f'bm_cuadro{numero}',
        }
        for numero, titulo in enumerate(titulos, 1)
    ]

def _iterar_textos(valor):
    if isinstance(valor, str):
        yield valor
    elif isinstance(valor, dict):
        for contenido in valor.values():
            yield from _iterar_textos(contenido)
    elif isinstance(valor, (list, tuple)):
        for contenido in valor:
            yield from _iterar_textos(contenido)

def _normalizar_busqueda(texto):
    import unicodedata

    texto = unicodedata.normalize('NFKD', str(texto).casefold())
    texto = ''.join(caracter for caracter in texto if not unicodedata.combining(caracter))
    return re.sub(r'[^a-z0-9]+', ' ', texto).strip()

def _datos_referencia(referencia):
    texto = _texto_referencia(referencia).strip()
    coincidencia = re.match(r'^(.*?)\s+\((\d{4})\)\.\s+(.+?)(?:\.\s+|$)', texto)
    if not coincidencia:
        return None
    autor, ano, titulo = coincidencia.groups()
    autor_norm = _normalizar_busqueda(autor)
    titulo_norm = _normalizar_busqueda(titulo)
    claves = {titulo_norm}
    if ',' in autor:
        claves.update(
            _normalizar_busqueda(apellido)
            for apellido in re.findall(r'(?:^|\by\s+)([^,]+),', autor)
        )
    else:
        claves.add(autor_norm)
        sigla = ''.join(
            palabra[0]
            for palabra in autor.split()
            if palabra and palabra[0].isupper()
        ).casefold()
        if len(sigla) >= 2:
            claves.add(sigla)
    claves.update(titulo_norm.split()[:1])
    return {'texto': texto, 'ano': ano, 'claves': {clave for clave in claves if len(clave) >= 3}}

def _validar_correspondencia_citas(referencias):
    """Valida citas explícitas de los bloques teóricos sin interpretar años narrativos."""
    bloques = [
        getattr(c, 'BASES_TEORICAS', []),
        getattr(c, 'BASES_TEORICAS_PARRAFOS', []),
        getattr(c, 'CITA_LARGA_TEXTO', ''),
        getattr(c, 'CITA_LARGA_AUTOR', ''),
    ]
    corpus = ' '.join(texto for bloque in bloques for texto in _iterar_textos(bloque))
    datos = [dato for referencia in referencias if (dato := _datos_referencia(referencia))]
    citadas = set()
    citas_sin_referencia = []
    for coincidencia in re.finditer(r'\([^()]*?\b((?:19|20)\d{2})\b[^()]*?\)', corpus):
        ano = coincidencia.group(1)
        inicio = max(0, coincidencia.start() - 120)
        contexto = _normalizar_busqueda(corpus[inicio:coincidencia.end()])
        candidatas = [dato for dato in datos if dato['ano'] == ano]
        emparejadas = [
            dato for dato in candidatas
            if any(clave in contexto for clave in dato['claves'])
        ]
        if emparejadas:
            citadas.update(dato['texto'] for dato in emparejadas)
        else:
            cita = corpus[coincidencia.start():coincidencia.end()]
            citas_sin_referencia.append(cita)
    no_citadas = [dato['texto'] for dato in datos if dato['texto'] not in citadas]
    return no_citadas, list(dict.fromkeys(citas_sin_referencia))

def validar_contenido():
    """Falla antes de construir el DOCX si el contenido estructural es inconsistente."""
    errores = []
    resumen = str(getattr(c, 'RESUMEN_TEXTO', ''))
    if len(resumen.split()) > RESUMEN_MAX_PALABRAS:
        errores.append(f"el resumen tiene {len(resumen.split())} palabras (máximo {RESUMEN_MAX_PALABRAS})")

    carpeta_imagenes = _ruta_desde_generador(getattr(c, 'CARPETA_IMAGENES', 'imagenes'))
    graficos = getattr(c, 'GRAFICOS', [])
    numeros_graficos = [g.get('numero') for g in graficos]
    if numeros_graficos != list(range(1, len(graficos) + 1)):
        errores.append("los gráficos deben estar numerados correlativamente desde 1")
    for grafico in graficos:
        numero = grafico.get('numero')
        if grafico.get('tras') not in ANCLAS_VALIDAS:
            errores.append(f"el gráfico {numero} tiene un ancla de inserción inválida")
        if not str(grafico.get('titulo', '')).startswith(f"Gráfico {numero}."):
            errores.append(f"el título del gráfico {numero} no coincide con su número")
        if not str(grafico.get('lista', '')).strip():
            errores.append(f"el gráfico {numero} no tiene descripción para la lista")
        if not str(grafico.get('fuente', '')).strip():
            errores.append(f"el gráfico {numero} no tiene fuente")
        if not _resolver_ruta_imagen(carpeta_imagenes, grafico):
            errores.append(f"no existe la imagen del gráfico {numero}")

    figuras = getattr(c, 'FIGURAS', None)
    if figuras is None:
        errores.append("falta la configuración explícita FIGURAS (use [] cuando no aplique)")
    elif figuras:
        errores.append(
            "FIGURAS contiene elementos, pero el generador no tiene un punto de inserción "
            "para crear sus marcadores; use GRAFICOS o implemente primero su ancla física"
        )

    for cuadro in _registro_cuadros():
        numero = cuadro['numero']
        titulo = cuadro['titulo']
        if not str(titulo).startswith(f"Cuadro {numero}."):
            errores.append(f"el título del cuadro {numero} no coincide con su número")
    for atributo in ('POBLACION_TABLA', 'POBLACION_FUENTE', 'CUADRO_PLANIFICACION_FUENTE', 'CUADRO_CRONOGRAMA_FUENTE'):
        if not getattr(c, atributo, None):
            errores.append(f"falta {atributo}")
    for fila in getattr(c, 'POBLACION_TABLA', []):
        if len(fila) != 5:
            errores.append("cada fila de POBLACION_TABLA debe tener cinco columnas")
    for fila in getattr(c, 'PLANIFICACION_DATOS', []):
        if len(fila) != 5:
            errores.append("cada fila de PLANIFICACION_DATOS debe tener cinco columnas")
    longitudes_cronograma = {len(fila[1]) for fila in getattr(c, 'CRONOGRAMA_DATOS', []) if len(fila) == 2}
    if len(longitudes_cronograma) > 1 or any(len(fila) != 2 for fila in getattr(c, 'CRONOGRAMA_DATOS', [])):
        errores.append("CRONOGRAMA_DATOS tiene filas o cantidades de semanas inconsistentes")

    referencias = getattr(c, 'REFERENCIAS_LISTA', [])
    textos_referencias = [_texto_referencia(ref).strip() for ref in referencias]
    normalizadas = [re.sub(r'\s+', ' ', texto).casefold() for texto in textos_referencias]
    if not all(textos_referencias):
        errores.append("hay referencias vacías")
    if len(normalizadas) != len(set(normalizadas)):
        errores.append("hay referencias duplicadas")
    if normalizadas != sorted(normalizadas):
        errores.append("las referencias no están en orden alfabético")
    referencias_no_citadas, citas_sin_referencia = _validar_correspondencia_citas(referencias)
    for referencia in referencias_no_citadas:
        errores.append(f"referencia no citada en las bases teóricas: {referencia}")
    for cita in citas_sin_referencia:
        errores.append(f"cita sin referencia reconocible: {cita}")

    anexos = getattr(c, 'ANEXOS_LISTA', [])
    codigos_anexos = [anexo[0] for anexo in anexos if anexo]
    codigos_esperados = [f"ANEXO {chr(65 + indice)}" for indice in range(len(anexos))]
    if codigos_anexos != codigos_esperados:
        errores.append("los anexos deben estar identificados correlativamente desde ANEXO A")
    for anexo in anexos:
        codigo = anexo[0] if anexo else "sin código"
        referencia = anexo[2] if len(anexo) > 2 else None
        fuente = anexo[4] if len(anexo) > 4 else None
        if not fuente:
            errores.append(f"{codigo} no tiene fuente propia")
        referencias_imagen = referencia if isinstance(referencia, (list, tuple)) else [referencia]
        for imagen in referencias_imagen:
            imagen_ref = imagen.get('archivo', imagen.get('referencia')) if isinstance(imagen, dict) else imagen
            if imagen_ref is not None and not buscar_imagen_por_referencia(carpeta_imagenes, imagen_ref):
                errores.append(f"no existe la imagen {imagen_ref!r} de {codigo}")

    if errores:
        raise ValueError("Validación previa fallida:\n- " + "\n- ".join(errores))
    print("✔ Validación previa de contenido completada.")

def agregar_pagina_aprobacion(doc, titulo, texto_parrafo, pie_firma, nombre_tutor, ci_tutor):
    """Agrega página de aprobación con membrete, título centrado, firma y datos del tutor."""
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    _config_seccion(sec)
    
    for linea in c.MEMBRETE:
        p_m = doc.add_paragraph()
        p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_m.paragraph_format.line_spacing = 1.15
        p_m.paragraph_format.space_before = Pt(0)
        p_m.paragraph_format.space_after = Pt(0)
        run_m = p_m.add_run(linea)
        run_m.font.name = FUENTE
        run_m.font.size = Pt(12)
        run_m.font.bold = True
    
    doc.add_paragraph()
    
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t.paragraph_format.space_before = Pt(24)
    p_t.paragraph_format.space_after = Pt(24)
    run_t = p_t.add_run(titulo)
    run_t.font.name = FUENTE
    run_t.font.size = Pt(12)
    run_t.font.bold = True
    bookmark_id = "bm_aprob_ind" if "INDUSTRIAL" in titulo else "bm_aprob_acad"
    _agregar_bookmark(p_t, bookmark_id)
    
    p_ap = doc.add_paragraph()
    p_ap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_ap.paragraph_format.line_spacing = 1.5
    p_ap.paragraph_format.first_line_indent = Cm(0)
    run_ap = p_ap.add_run(texto_parrafo)
    run_ap.font.name = FUENTE
    run_ap.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p_f = doc.add_paragraph()
    p_f.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_f.paragraph_format.line_spacing = 1.5
    run_f = p_f.add_run(pie_firma)
    run_f.font.name = FUENTE
    run_f.font.size = Pt(12)
    
    for _ in range(1):
        doc.add_paragraph()
    
    # Línea de firma (o imagen de firma si está configurada en contenido.py)
    if "INDUSTRIAL" in titulo:
        firma_img = getattr(c, 'FIRMA_TUTOR_INDUSTRIAL', None)
    else:
        firma_img = getattr(c, 'FIRMA_TUTOR_ACADEMICO', None)

    ruta_firma = None
    if firma_img:
        ruta_firma = os.path.join(_ruta_desde_generador(getattr(c, 'CARPETA_IMAGENES', 'imagenes')), firma_img)
        if not os.path.exists(ruta_firma):
            ruta_firma = None

    if ruta_firma:
        p_firma = doc.add_paragraph()
        p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_firma = p_firma.add_run()
        run_firma.add_picture(ruta_firma, width=Cm(4))
    else:
        p_linea = doc.add_paragraph()
        p_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_linea = p_linea.add_run("_" * 35)
        run_linea.font.name = FUENTE
        run_linea.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Nombre del tutor centrado
    p_nom = doc.add_paragraph()
    p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nom.paragraph_format.line_spacing = 1.5
    run_nom = p_nom.add_run(nombre_tutor)
    run_nom.font.name = FUENTE
    run_nom.font.size = Pt(12)
    run_nom.font.bold = True
    
    # Cédula del tutor centrada
    p_ci = doc.add_paragraph()
    p_ci.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ci.paragraph_format.space_before = Pt(0)
    run_ci = p_ci.add_run(ci_tutor)
    run_ci.font.name = FUENTE
    run_ci.font.size = Pt(12)

def construir_cuerpo_documento(doc, modo="completo"):
    """Escribe secuencialmente todas las secciones del informe de pasantía.
    
    modo: "completo" | "borrador1" (solo Cap I) | "borrador2" (Cap I+II) |
          "borrador3" (Cap I+II+III) | "borrador4" (Cap I+II+III+IV)
    """
    tiene_cap2 = modo in ("completo", "borrador2", "borrador3", "borrador4")
    tiene_cap3 = modo in ("completo", "borrador3", "borrador4")
    tiene_cap4 = modo in ("completo", "borrador4")
    tiene_cap5 = modo == "completo"

    # LibreOffice 26.2 ignora el modificador romano de PAGEREF y devuelve
    # arábigos. Solo estas entradas preliminares conservan cálculo estático;
    # todos los destinos del cuerpo y listas usan PAGEREF real.
    pagina_romana = 5
    paginas_preliminares = {
        'bm_contraportada': 'ii',
        'bm_aprob_ind': 'iii',
        'bm_aprob_acad': 'iv',
    }

    def registrar_preliminar(bookmark_id, presente=True, paginas=1):
        nonlocal pagina_romana
        if not presente:
            return
        paginas_preliminares[bookmark_id] = _numero_romano(pagina_romana)
        pagina_romana += paginas

    registrar_preliminar('bm_agradecimientos', bool(getattr(c, 'AGRADECIMIENTOS', None)))
    registrar_preliminar('bm_dedicatoria', bool(getattr(c, 'DEDICATORIA', None)))
    pagina_romana += 2 if modo == 'completo' else 1  # índice de contenido
    registrar_preliminar('bm_lista_cuadros')
    registrar_preliminar('bm_lista_figuras', bool(getattr(c, 'FIGURAS', [])))
    registrar_preliminar('bm_lista_graficos', bool(getattr(c, 'GRAFICOS', [])))
    registrar_preliminar('bm_lista_anexos', bool(getattr(c, 'ANEXOS_LISTA', [])))
    registrar_preliminar('bm_resumen', bool(getattr(c, 'RESUMEN_TEXTO', None)))

    idx_inicio_arabigo = None

    if True:
        # --- APROBACIÓN DEL TUTOR INDUSTRIAL ---
        nom_pas = getattr(c, 'NOMBRE_PASANTE', '[Nombre del Pasante]')
        ci_pas = getattr(c, 'CI_PASANTE', 'XX.XXX.XXX')
        esp = getattr(c, 'ESPECIALIDAD', '[Especialidad]')
        tit_proy = getattr(c, 'TITULO_PROYECTO', '[Título del Proyecto]')
        ciudad = getattr(c, 'CIUDAD_FECHA', 'El Tigre').split(",")[0] if "," in getattr(c, 'CIUDAD_FECHA', 'El Tigre') else 'El Tigre'
        
        texto_base = (
            f"En mi car\u00e1cter de tutor industrial del informe de pasant\u00edas presentado por: "
            f"{nom_pas}, de C\u00e9dula de Identidad V-{ci_pas}; para optar al grado de "
            f"T\u00e9cnico Superior Universitario en la especialidad de: {esp}, cuyo t\u00edtulo es; "
            f"\u201c{tit_proy}\u201d, manifiesto que cumple con los requisitos exigidos por el "
            f"Instituto Universitario de Tecnolog\u00eda \u201cEl\u00edas Calixto Pompa\u201d (IUTECP); "
            f"y que, por lo tanto, considero que re\u00fane los m\u00e9ritos suficientes para ser "
            f"evaluado por el jurado que se decida designar a tal fin."
        )
        # Extraer nombres y CIs de tutores desde AUTOR_DATOS
        autor_datos = getattr(c, 'AUTOR_DATOS', [])
        tut_ind_nom = ""
        tut_ind_ci = ""
        tut_acad_nom = ""
        tut_acad_ci = ""
        for i, linea in enumerate(autor_datos):
            if "Tutor Industrial" in linea and ":" in linea:
                partes = linea.split(":", 1)
                if len(partes) > 1 and partes[1].strip():
                    tut_ind_nom = partes[1].strip()
                elif i + 2 < len(autor_datos):
                    tut_ind_nom = autor_datos[i + 1].strip()
                    if "C.I.:" in autor_datos[i + 2]:
                        tut_ind_ci = autor_datos[i + 2].split(":", 1)[1].strip()
            if "Tutor Académico" in linea and ":" in linea:
                partes = linea.split(":", 1)
                if len(partes) > 1 and partes[1].strip():
                    tut_acad_nom = partes[1].strip()
                elif i + 2 < len(autor_datos):
                    tut_acad_nom = autor_datos[i + 1].strip()
                    if "C.I.:" in autor_datos[i + 2]:
                        tut_acad_ci = autor_datos[i + 2].split(":", 1)[1].strip()
            if "C.I.:" in linea and i > 0:
                prev = autor_datos[i - 1].strip()
                if prev == tut_ind_nom:
                    tut_ind_ci = linea.split(":", 1)[1].strip()
                elif prev == tut_acad_nom:
                    tut_acad_ci = linea.split(":", 1)[1].strip()

        fecha_aprobacion_ind = getattr(
            c,
            'TEXTO_FECHA_APROBACION_TUTOR_INDUSTRIAL',
            'a los ___ días del mes de _______',
        )
        coincidencia_ano = re.search(r'\b(\d{4})\b', getattr(c, 'FECHA_LUGAR', ''))
        ano_aprobacion = coincidencia_ano.group(1) if coincidencia_ano else '____'
        fecha_aprobacion_ind = re.sub(r'\b\d{4}\b', ano_aprobacion, fecha_aprobacion_ind)
        if ano_aprobacion not in fecha_aprobacion_ind:
            fecha_aprobacion_ind = f"{fecha_aprobacion_ind.rstrip()} de {ano_aprobacion}"
        agregar_pagina_aprobacion(doc, "APROBACIÓN DEL TUTOR INDUSTRIAL", texto_base,
            f"En la Ciudad de {ciudad}, {fecha_aprobacion_ind}",
            tut_ind_nom, f"C.I.: {tut_ind_ci}" if tut_ind_ci else "")

        # --- APROBACIÓN DEL TUTOR ACADÉMICO ---

        # --- APROBACIÓN DEL TUTOR ACADÉMICO ---
        texto_base_acad = (
            f"En mi car\u00e1cter de tutor acad\u00e9mico del informe de pasant\u00edas presentado por: "
            f"{nom_pas}, de C\u00e9dula de Identidad V-{ci_pas}; para optar al grado de "
            f"T\u00e9cnico Superior Universitario en la especialidad de: {esp}, cuyo t\u00edtulo es; "
            f"\u201c{tit_proy}\u201d, manifiesto que cumple con los requisitos exigidos por el "
            f"Instituto Universitario de Tecnolog\u00eda \u201cEl\u00edas Calixto Pompa\u201d (IUTECP); "
            f"y que, por lo tanto, considero que re\u00fane los m\u00e9ritos suficientes para ser "
            f"evaluado por el jurado que se decida designar a tal fin."
        )
        agregar_pagina_aprobacion(doc, "APROBACIÓN DEL TUTOR ACADÉMICO", texto_base_acad,
            f"En la Ciudad de {ciudad}, a los ___ días del mes de _______ de {ano_aprobacion}",
            tut_acad_nom, f"C.I.: {tut_acad_ci}" if tut_acad_ci else "")

        # --- PÁGINAS PRELIMINARES ---
        if hasattr(c, 'AGRADECIMIENTOS') and c.AGRADECIMIENTOS:
            iniciar_seccion_preliminar(doc, "AGRADECIMIENTOS", bookmark_id="bm_agradecimientos")
            agregar_parrafo_normado(doc, c.AGRADECIMIENTOS)
    
        if hasattr(c, 'DEDICATORIA') and c.DEDICATORIA:
            iniciar_seccion_preliminar(doc, "DEDICATORIA", bookmark_id="bm_dedicatoria")
            agregar_parrafo_normado(doc, c.DEDICATORIA)
    
        # --- ÍNDICE DE CONTENIDO ---
        idx_indice_inicio = len(doc.sections)
        iniciar_seccion_preliminar(doc, "ÍNDICE DE CONTENIDO", bookmark_id="bm_indice")
        p_header_ind = doc.add_paragraph()
        p_header_ind.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_header_ind.paragraph_format.space_after = Pt(12)
        run_h_ind = p_header_ind.add_run("pp.")
        run_h_ind.font.name = 'Times New Roman'
        run_h_ind.font.size = Pt(12)
        run_h_ind.font.bold = True
    
        preliminares_indice = [
            ("CONTRAPORTADA", "bm_contraportada"),
            ("APROBACIÓN DEL TUTOR INDUSTRIAL", "bm_aprob_ind"),
            ("APROBACIÓN DEL TUTOR ACADÉMICO", "bm_aprob_acad"),
        ]
        if getattr(c, 'AGRADECIMIENTOS', None):
            preliminares_indice.append(("AGRADECIMIENTOS", "bm_agradecimientos"))
        if getattr(c, 'DEDICATORIA', None):
            preliminares_indice.append(("DEDICATORIA", "bm_dedicatoria"))
        preliminares_indice.append(("LISTA DE CUADROS", "bm_lista_cuadros"))
        if getattr(c, 'FIGURAS', []):
            preliminares_indice.append(("LISTA DE FIGURAS", "bm_lista_figuras"))
        if getattr(c, 'GRAFICOS', []):
            preliminares_indice.append(("LISTA DE GRÁFICOS", "bm_lista_graficos"))
        if getattr(c, 'ANEXOS_LISTA', []):
            preliminares_indice.append(("LISTA DE ANEXOS", "bm_lista_anexos"))
        if getattr(c, 'RESUMEN_TEXTO', None):
            preliminares_indice.append(("RESUMEN", "bm_resumen"))
        for titulo_preliminar, bookmark_preliminar in preliminares_indice:
            agregar_fila_indice_general_nativa(
                doc,
                titulo_preliminar,
                pagina=paginas_preliminares[bookmark_preliminar],
            )
    
        agregar_fila_indice_general_nativa(doc, "INTRODUCCIÓN", "", bookmark_id="bm_introduccion")

        idx_indice_fin = len(doc.sections)

        # Capítulos
        p_cap_lbl = doc.add_paragraph()
        p_cap_lbl.paragraph_format.space_before = Pt(12)
        p_cap_lbl.paragraph_format.space_after = Pt(6)
        _mantener_con_siguiente(p_cap_lbl)
        p_cap_lbl.add_run("CAPÍTULOS").font.bold = True
    
        # Capítulo I
        agregar_fila_indice_general_nativa(doc, "CAPÍTULO I: REALIDAD ORGANIZACIONAL", "", sangria_cm=0, negrita=True, bookmark_id="bm_cap1")
        agregar_fila_indice_general_nativa(doc, "Identificación de la empresa", "", sangria_cm=0.5, bookmark_id="bm_cap1_ident")
        agregar_fila_indice_general_nativa(doc, "Reseña histórica", "", sangria_cm=0.5, bookmark_id="bm_cap1_resena")
        agregar_fila_indice_general_nativa(doc, "Misión", "", sangria_cm=0.5, bookmark_id="bm_cap1_mision")
        agregar_fila_indice_general_nativa(doc, "Visión", "", sangria_cm=0.5, bookmark_id="bm_cap1_vision")
        agregar_fila_indice_general_nativa(doc, "Valores", "", sangria_cm=0.5, bookmark_id="bm_cap1_valores")
        agregar_fila_indice_general_nativa(doc, "Objetivos Organizacionales", "", sangria_cm=0.5, bookmark_id="bm_cap1_obj")
        agregar_fila_indice_general_nativa(doc, "Ubicación geográfica", "", sangria_cm=0.5, bookmark_id="bm_cap1_ubic")
        agregar_fila_indice_general_nativa(doc, "Población de los trabajadores de la empresa", "", sangria_cm=0.5, bookmark_id="bm_cap1_pobla")
        agregar_fila_indice_general_nativa(doc, "Estructura organizacional de la empresa (organigrama)", "", sangria_cm=0.5, bookmark_id="bm_cap1_estruct")
        agregar_fila_indice_general_nativa(doc, "Descripción del departamento donde realizó la pasantía", "", sangria_cm=0.5, bookmark_id="bm_cap1_departamento")
    
        # Capítulo II
        if tiene_cap2:
            agregar_fila_indice_general_nativa(doc, "CAPÍTULO II: DIAGNÓSTICO SITUACIONAL", "", sangria_cm=0, negrita=True, bookmark_id="bm_cap2")
            agregar_fila_indice_general_nativa(doc, "Identificación de la situación problemática", "", sangria_cm=0.5, bookmark_id="bm_cap2_sit")
            agregar_fila_indice_general_nativa(doc, "Objetivo General", "", sangria_cm=0.5, bookmark_id="bm_cap2_objg")
            agregar_fila_indice_general_nativa(doc, "Objetivos Específicos", "", sangria_cm=0.5, bookmark_id="bm_cap2_obje")
            agregar_fila_indice_general_nativa(doc, "Planificación integral de objetivos", "", sangria_cm=0.5, bookmark_id="bm_cap2_planif")
            agregar_fila_indice_general_nativa(doc, "Cronograma de actividades", "", sangria_cm=0.5, bookmark_id="bm_cap2_crono")
    
        # Capítulo III
        if tiene_cap3:
            agregar_fila_indice_general_nativa(doc, "CAPÍTULO III: MARCO TEÓRICO", "", sangria_cm=0, negrita=True, bookmark_id="bm_cap3")
            agregar_fila_indice_general_nativa(doc, "Bases teóricas referenciales", "", sangria_cm=0.5, bookmark_id="bm_cap3_bases")
    
        # Capítulo IV
        if tiene_cap4:
            agregar_fila_indice_general_nativa(doc, "CAPÍTULO IV: ACTIVIDADES REALIZADAS", "", sangria_cm=0, negrita=True, bookmark_id="bm_cap4")
            agregar_fila_indice_general_nativa(doc, "Descripción de actividades ejecutadas por semana", "", sangria_cm=0.5, bookmark_id="bm_cap4_desc")
    
        # Capítulo V
        if tiene_cap5:
            agregar_fila_indice_general_nativa(doc, "CAPÍTULO V: CONCLUSIONES Y RECOMENDACIONES", "", sangria_cm=0, negrita=True, bookmark_id="bm_cap5")
            agregar_fila_indice_general_nativa(doc, "Conclusiones", "", sangria_cm=0.5, bookmark_id="bm_cap5_concl")
            agregar_fila_indice_general_nativa(doc, "Recomendaciones", "", sangria_cm=0.5, bookmark_id="bm_cap5_recom")
    
        # Referencias y Anexos
        agregar_fila_indice_general_nativa(doc, "REFERENCIAS", "", sangria_cm=0, negrita=True, bookmark_id="bm_referencias")
        agregar_fila_indice_general_nativa(doc, "ANEXOS", "", sangria_cm=0, negrita=True, bookmark_id="bm_anexos")
    
        # --- LISTA DE CUADROS ---
        iniciar_seccion_preliminar(doc, "LISTA DE CUADROS", bookmark_id="bm_lista_cuadros")
        p_header_c = doc.add_paragraph()
        p_header_c.paragraph_format.space_before = Pt(12)
        p_header_c.paragraph_format.space_after = Pt(12)
        _mantener_con_siguiente(p_header_c)
    
        section_c = doc.sections[-1]
        ancho_c_emu = section_c.page_width - section_c.left_margin - section_c.right_margin
        agregar_tabulacion_derecha(p_header_c, ancho_c_emu)
    
        run_c_lbl = p_header_c.add_run("CUADRO")
        run_c_lbl.font.name = 'Times New Roman'
        run_c_lbl.font.size = Pt(12)
        run_c_lbl.font.bold = True
    
        p_header_c.add_run("\t")  # Salto de tabulador explícito para OXML
    
        run_pp_c = p_header_c.add_run("pp.")
        run_pp_c.font.name = FUENTE
        run_pp_c.font.size = Pt(TAMANO_BASE)
        run_pp_c.font.bold = True
    
        for cuadro in _registro_cuadros(incluir_capitulo2=tiene_cap2):
            agregar_fila_lista_preliminar_nativa(
                doc,
                str(cuadro['numero']),
                cuadro['descripcion'],
                '',
                bookmark_id=cuadro['bookmark'],
            )
    
        # --- LISTA DE FIGURAS ---
        figuras_cfg = getattr(c, 'FIGURAS', [])
        if figuras_cfg:
            iniciar_seccion_preliminar(doc, "LISTA DE FIGURAS", bookmark_id="bm_lista_figuras")
            p_header_f = doc.add_paragraph()
            p_header_f.paragraph_format.space_before = Pt(12)
            p_header_f.paragraph_format.space_after = Pt(12)
            _mantener_con_siguiente(p_header_f)

            section_f = doc.sections[-1]
            ancho_f_emu = section_f.page_width - section_f.left_margin - section_f.right_margin
            agregar_tabulacion_derecha(p_header_f, ancho_f_emu)

            run_f_lbl = p_header_f.add_run("FIGURA")
            run_f_lbl.font.name = FUENTE
            run_f_lbl.font.size = Pt(TAMANO_BASE)
            run_f_lbl.font.bold = True
            p_header_f.add_run("\t")
            run_pp_f = p_header_f.add_run("pp.")
            run_pp_f.font.name = FUENTE
            run_pp_f.font.size = Pt(TAMANO_BASE)
            run_pp_f.font.bold = True

            for figura in figuras_cfg:
                num = str(figura.get("numero", ""))
                desc = figura.get("lista", figura.get("titulo", "").split('. ', 1)[-1])
                pag = str(figura.get("pagina", ""))
                bookmark_id = f"bm_figura{num}" if num else None
                agregar_fila_lista_preliminar_nativa(doc, num, desc, pag, bookmark_id=bookmark_id)

        # --- LISTA DE GRÁFICOS ---
        graficos_cfg = getattr(c, 'GRAFICOS', [])
        if graficos_cfg:
            iniciar_seccion_preliminar(doc, "LISTA DE GRÁFICOS", bookmark_id="bm_lista_graficos")
            p_header_g = doc.add_paragraph()
            p_header_g.paragraph_format.space_before = Pt(12)
            p_header_g.paragraph_format.space_after = Pt(12)
            _mantener_con_siguiente(p_header_g)

            section_g = doc.sections[-1]
            ancho_g_emu = section_g.page_width - section_g.left_margin - section_g.right_margin
            agregar_tabulacion_derecha(p_header_g, ancho_g_emu)

            run_g_lbl = p_header_g.add_run("GRÁFICO")
            run_g_lbl.font.name = FUENTE
            run_g_lbl.font.size = Pt(TAMANO_BASE)
            run_g_lbl.font.bold = True
            p_header_g.add_run("\t")
            run_pp_g = p_header_g.add_run("pp.")
            run_pp_g.font.name = FUENTE
            run_pp_g.font.size = Pt(TAMANO_BASE)
            run_pp_g.font.bold = True

            for g in graficos_cfg:
                num = str(g.get("numero", ""))
                desc = g.get("lista", g.get("titulo", "").split('. ', 1)[-1] if '. ' in g.get("titulo", "") else "")
                bookmark_id = f"bm_grafico{num}" if num else None
                agregar_fila_lista_preliminar_nativa(doc, num, desc, '', bookmark_id=bookmark_id)
    
        # --- LISTA DE ANEXOS ---
        if getattr(c, 'ANEXOS_LISTA', []):
            iniciar_seccion_preliminar(doc, "LISTA DE ANEXOS", bookmark_id="bm_lista_anexos")
            p_header_a = doc.add_paragraph()
            p_header_a.paragraph_format.space_before = Pt(12)
            p_header_a.paragraph_format.space_after = Pt(12)
            _mantener_con_siguiente(p_header_a)
        
            section_a = doc.sections[-1]
            ancho_a_emu = section_a.page_width - section_a.left_margin - section_a.right_margin
            agregar_tabulacion_derecha(p_header_a, ancho_a_emu)
        
            run_a_lbl = p_header_a.add_run("ANEXOS")
            run_a_lbl.font.name = 'Times New Roman'
            run_a_lbl.font.size = Pt(12)
            run_a_lbl.font.bold = True
        
            p_header_a.add_run("\t")  # Salto de tabulador explícito para OXML
        
            run_pp_a = p_header_a.add_run("pp.")
            run_pp_a.font.name = 'Times New Roman'
            run_pp_a.font.size = Pt(12)
            run_pp_a.font.bold = True
        
            anexos_lista = getattr(c, 'ANEXOS_LISTA', [])
            for idx, anexo in enumerate(anexos_lista):
                cod, desc = anexo[:2]
                letra = cod.split(" ")[-1]
                pag_est = str(13 + idx)
                bookmark_id = f"bm_anexo{letra.upper()}" if letra else None
                agregar_fila_lista_preliminar_nativa(doc, letra, desc, pag_est, bookmark_id=bookmark_id)
    
        # --- RESUMEN ---
        if hasattr(c, 'RESUMEN_TEXTO') and c.RESUMEN_TEXTO:
            iniciar_seccion_resumen(doc, c, bookmark_id="bm_resumen")
            agregar_parrafo_resumen(doc, c.RESUMEN_TEXTO)
            doc.add_paragraph()
            p_kw = doc.add_paragraph()
            p_kw.paragraph_format.line_spacing = 1.0
            p_kw.paragraph_format.first_line_indent = Cm(1.25)
            p_kw.paragraph_format.space_before = Pt(0)
            p_kw.paragraph_format.space_after = Pt(0)
            run_kw_label = p_kw.add_run("Palabras claves: ")
            run_kw_label.font.name = FUENTE
            run_kw_label.font.size = Pt(TAMANO_BASE)
            run_kw_label.font.bold = True
            run_kw = p_kw.add_run(c.PALABRAS_CLAVE)
            run_kw.font.name = FUENTE
            run_kw.font.size = Pt(TAMANO_BASE)

        # --- REGISTRO DEL INICIO DEL CUERPO ---
        idx_inicio_arabigo = len(doc.sections)
        iniciar_seccion_preliminar(
            doc,
            "INTRODUCCIÓN",
            bookmark_id="bm_introduccion",
            ocultar_primera_pagina=True,
        )
        introduccion = getattr(c, 'INTRODUCCION_TEXTO', 'Texto de introducción no proporcionado.')
        if isinstance(introduccion, (list, tuple)):
            for parrafo in introduccion:
                agregar_parrafo_normado(doc, parrafo)
        else:
            agregar_parrafo_normado(doc, introduccion)

        # Capítulo I continúa la secuencia arábiga iniciada en Introducción.
        idx_cap1 = len(doc.sections)
    else:
        idx_cap1 = len(doc.sections)

# --- CAPÍTULO I: REALIDAD ORGANIZACIONAL ---
    iniciar_capitulo(doc, "I", "REALIDAD ORGANIZACIONAL", bookmark_id="bm_cap1")
    p_ident = doc.add_paragraph(style=ESTILO_SUBTITULO_CENTRADO)
    p_ident.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ident.paragraph_format.space_before = ESP_DOBLE
    p_ident.paragraph_format.space_after = ESP_DOBLE
    p_ident.paragraph_format.line_spacing = INTERLINEADO
    r_ident = p_ident.add_run("Identificación de la empresa")
    r_ident.font.name = FUENTE
    r_ident.font.size = Pt(TAMANO_BASE)
    r_ident.font.bold = True
    _agregar_bookmark(p_ident, "bm_cap1_ident")

    agregar_titulo_nivel2(doc, "Razón social")
    agregar_parrafo_normado(doc, getattr(c, 'RAZON_SOCIAL', 'Razón Social no proporcionada.'), sangria=False)

    agregar_titulo_nivel2(doc, "Reseña histórica", bookmark_id="bm_cap1_resena")
    resena_data = getattr(c, 'RESENA_HISTORICA', [])
    if isinstance(resena_data, str):
        agregar_parrafo_normado(doc, resena_data)
    else:
        for parrafo in resena_data:
            agregar_parrafo_normado(doc, parrafo)

    agregar_titulo_nivel2(doc, "Misión", bookmark_id="bm_cap1_mision")
    agregar_parrafo_normado(doc, getattr(c, 'MISION', 'Misión no proporcionada.'))

    agregar_titulo_nivel2(doc, "Visión", bookmark_id="bm_cap1_vision")
    agregar_parrafo_normado(doc, getattr(c, 'VISION', 'Visión no proporcionada.'))

    agregar_titulo_nivel2(doc, "Valores", bookmark_id="bm_cap1_valores")
    agregar_parrafo_normado(doc, "Los valores que orientan las actividades de la organización destacan:")
    valores_data = getattr(c, 'VALORES', [])
    for i, (valor, descripcion) in enumerate(valores_data, 1):
        agregar_item_lista(doc, i, descripcion, valor)

    agregar_titulo_nivel2(doc, "Objetivos Organizacionales", bookmark_id="bm_cap1_obj")
    agregar_titulo_nivel2(doc, "Objetivo General")
    agregar_parrafo_normado(doc, getattr(c, 'OBJETIVO_GENERAL_EMPRESA', ''))
    objs_espec_emp = getattr(c, 'OBJETIVOS_ESPECIFICOS_EMPRESA', [])
    if objs_espec_emp:
        agregar_titulo_nivel2(doc, "Objetivos Específicos")
        for i, obj in enumerate(objs_espec_emp, 1):
            agregar_item_lista(doc, i, obj)

    carpeta_imagenes = getattr(c, 'CARPETA_IMAGENES', 'imagenes')
    _insertar_logo_empresa(doc, carpeta_imagenes)

    agregar_titulo_nivel2(doc, "Ubicación geográfica", bookmark_id="bm_cap1_ubic")
    agregar_parrafo_normado(doc, getattr(c, 'UBICACION', 'Ubicación no proporcionada.'), sangria=False)

    _insertar_graficos_por_ancla(doc, carpeta_imagenes, "ubicacion")

    agregar_titulo_nivel2(doc, "Población de los trabajadores de la empresa", bookmark_id="bm_cap1_pobla")
    poblacion_tabla = getattr(c, 'POBLACION_TABLA', None)
    if poblacion_tabla:
        cuadro_poblacion = _registro_cuadros()[0]
        agregar_titulo_cuadro(
            doc,
            cuadro_poblacion['titulo'],
            bookmark_id=cuadro_poblacion['bookmark'],
        )
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False
        anchos = [3.5, 5.09, 2, 2, 2]
        aplicar_formato_tabla_xml(table, anchos)
        _repetir_encabezado(table.rows[0])
        headers = ["Departamento / Área", "Cargo", "Femenino", "Masculino", "Total"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(h)
            run.font.name = FUENTE
            run.font.size = Pt(TAMANO_TABLA)
            run.font.bold = True
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), "D9E2F3")
            cell._tc.get_or_add_tcPr().append(shading)
        for dep, cargo, fem, masc, total in poblacion_tabla:
            row = table.add_row()
            for j, val in enumerate([dep, cargo, fem, masc, total]):
                cell = row.cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j >= 2 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(str(val))
                run.font.name = FUENTE
                run.font.size = Pt(TAMANO_TABLA)
                if j == 0 and dep.startswith("**"):
                    run.font.bold = True
        _agregar_fuente(doc, getattr(c, 'POBLACION_FUENTE', None))
    poblacion_data = getattr(c, 'POBLACION', '')
    if isinstance(poblacion_data, str):
        agregar_parrafo_normado(doc, poblacion_data)
    else:
        for parrafo in poblacion_data:
            agregar_parrafo_normado(doc, parrafo)

    agregar_titulo_nivel2(doc, "Estructura Organizativa", bookmark_id="bm_cap1_estruct")
    org_texto = getattr(c, 'ORGANIGRAMA_TEXTO', 'Estructura organizativa.')
    if isinstance(org_texto, list):
        for parrafo in org_texto:
            agregar_parrafo_normado(doc, parrafo)
    else:
        agregar_parrafo_normado(doc, org_texto)

    _insertar_graficos_por_ancla(doc, carpeta_imagenes, "estructura")

    agregar_titulo_nivel2(
        doc,
        "Descripción del departamento donde realizó la pasantía",
        bookmark_id="bm_cap1_departamento",
    )
    descripcion_departamento = getattr(c, 'DESCRIPCION_DEPARTAMENTO', '')
    if isinstance(descripcion_departamento, (list, tuple)):
        for parrafo in descripcion_departamento:
            agregar_parrafo_normado(doc, parrafo)
    elif descripcion_departamento:
        agregar_parrafo_normado(doc, descripcion_departamento)

    # --- CAPÍTULO II: DIAGNÓSTICO SITUACIONAL ---
    if tiene_cap2:
        iniciar_capitulo(doc, "II", "DIAGNÓSTICO SITUACIONAL", bookmark_id="bm_cap2")
        agregar_titulo_nivel2(doc, "Identificación de la Situación Problemática", bookmark_id="bm_cap2_sit")
        situacion_problematica = getattr(c, 'SITUACION_PROBLEMATICA', [])
        if isinstance(situacion_problematica, str):
            agregar_parrafo_normado(doc, situacion_problematica)
        else:
            mostrar_niveles_diagnostico = bool(
                getattr(c, 'MOSTRAR_NIVELES_DIAGNOSTICO', False)
            )
            for bloque in situacion_problematica:
                if isinstance(bloque, dict):
                    titulo_bloque = bloque.get('titulo')
                    titulo_visible = str(titulo_bloque).strip().casefold() if titulo_bloque else ''
                    es_nivel_diagnostico = titulo_visible in {
                        'nivel macro',
                        'nivel meso',
                        'nivel micro',
                    }
                    if titulo_bloque and (mostrar_niveles_diagnostico or not es_nivel_diagnostico):
                        agregar_titulo_nivel2(doc, titulo_bloque)
                    parrafos_bloque = bloque.get('parrafos', [])
                    if isinstance(parrafos_bloque, str):
                        parrafos_bloque = [parrafos_bloque]
                    for parrafo in parrafos_bloque:
                        agregar_parrafo_normado(doc, parrafo)
                else:
                    agregar_parrafo_normado(doc, bloque)

        interrogante = getattr(c, 'INTERROGANTE_PROBLEMA', '')
        if interrogante:
            agregar_titulo_nivel2(doc, getattr(c, 'INTERROGANTE_TITULO', 'Interrogante orientadora'))
            agregar_parrafo_normado(doc, interrogante)

        agregar_titulo_nivel2(doc, "Objetivo General", bookmark_id="bm_cap2_objg")
        agregar_parrafo_normado(doc, getattr(c, 'OBJETIVO_GENERAL', 'Objetivo general no proporcionado.'))

        agregar_titulo_nivel2(doc, "Objetivos Específicos", bookmark_id="bm_cap2_obje")
        objs_especificos = getattr(c, 'OBJETIVOS_ESPECIFICOS', [])
        for i, obj in enumerate(objs_especificos, 1):
            agregar_item_lista(doc, i, obj)

        agregar_titulo_nivel2(doc, "Planificación integral de objetivos", bookmark_id="bm_cap2_planif")
        intro_planif = getattr(c, 'PLANIFICACION_INTRO_TEXTO',
            "La planificación establece la relación entre cada objetivo y las actividades administrativas a ejecutar:")
        agregar_parrafo_normado(doc, intro_planif)
        cuadro_planificacion = _registro_cuadros()[1]
        agregar_tabla_planificacion(
            doc,
            getattr(c, 'PLANIFICACION_DATOS', []),
            titulo_cuadro=cuadro_planificacion['titulo'],
            bookmark_id=cuadro_planificacion['bookmark'],
            fuente=getattr(c, 'CUADRO_PLANIFICACION_FUENTE', None),
        )
        doc.add_paragraph()

        agregar_titulo_nivel2(doc, "Cronograma de actividades", bookmark_id="bm_cap2_crono")
        intro_crono = getattr(c, 'CRONOGRAMA_INTRO_TEXTO',
            "El cronograma estructura temporalmente las tareas administrativas garantizando el cumplimiento del manual documental propuesto:")
        agregar_parrafo_normado(doc, intro_crono)
        cuadro_cronograma = _registro_cuadros()[2]
        agregar_gantt(
            doc,
            getattr(c, 'CRONOGRAMA_DATOS', []),
            titulo_cuadro=cuadro_cronograma['titulo'],
            bookmark_id=cuadro_cronograma['bookmark'],
            fuente=getattr(c, 'CUADRO_CRONOGRAMA_FUENTE', None),
        )

    # --- CAPÍTULO III: MARCO TEÓRICO ---
    if tiene_cap3:
        iniciar_capitulo(doc, "III", "MARCO TEÓRICO", bookmark_id="bm_cap3")
        if hasattr(c, 'BASES_TEORICAS') and isinstance(c.BASES_TEORICAS, list) and c.BASES_TEORICAS and isinstance(c.BASES_TEORICAS[0], dict):
            primer_sub = True
            for sub in c.BASES_TEORICAS:
                bm = "bm_cap3_bases" if primer_sub else None
                agregar_titulo_nivel2(doc, sub.get('titulo', ''), bookmark_id=bm)
                primer_sub = False
                for p in sub.get('parrafos', []):
                    agregar_parrafo_normado(doc, p)
                cita = sub.get('cita_larga')
                if cita and cita.get('texto'):
                    agregar_cita_larga(doc, cita['texto'], cita.get('autor', ''))
                    post_cita = sub.get('post_cita', getattr(c, 'POST_CITA_TEXTO', POST_CITA_TEXTO_DEF))
                    if post_cita:
                        agregar_parrafo_normado(doc, post_cita, sangria=True)
                posicion_autor = sub.get('posicion_autor')
                if posicion_autor:
                    agregar_parrafo_normado(doc, posicion_autor)
        else:
            agregar_titulo_nivel2(doc, "Bases Teóricas Referenciales", bookmark_id="bm_cap3_bases")
            bases_teoricas = getattr(c, 'BASES_TEORICAS_PARRAFOS', ['Bases teóricas referenciales.'])
            for parrafo in bases_teoricas:
                agregar_parrafo_normado(doc, parrafo)
            if hasattr(c, 'CITA_LARGA_TEXTO') and c.CITA_LARGA_TEXTO:
                agregar_cita_larga(doc, c.CITA_LARGA_TEXTO, getattr(c, 'CITA_LARGA_AUTOR', ''))
                post_cita = getattr(c, 'POST_CITA_TEXTO', POST_CITA_TEXTO_DEF)
                agregar_parrafo_normado(doc, post_cita, sangria=True)

    # --- CAPÍTULO IV: ACTIVIDADES REALIZADAS ---
    if tiene_cap4:
        iniciar_capitulo(doc, "IV", "ACTIVIDADES REALIZADAS", bookmark_id="bm_cap4")
        agregar_titulo_nivel2(doc, "Descripción de Actividades Ejecutadas por Semana", bookmark_id="bm_cap4_desc")
        agregar_parrafo_normado(doc, getattr(c, 'ACTIVIDADES_DESCRIPCION', 'Descripción de actividades ejecutadas.'))
        actividades_lista = getattr(c, 'ACTIVIDADES_LISTA', [])
        for i, actividad in enumerate(actividades_lista, 1):
            if isinstance(actividad, dict):
                agregar_titulo_nivel2(doc, f"Semana {actividad.get('semana', i)}")
                if actividad.get('operativa'):
                    agregar_item_lista(doc, 1, actividad['operativa'], negrita_inicio="Actividad operativa")
                if actividad.get('investigacion'):
                    agregar_item_lista(
                        doc,
                        2,
                        actividad['investigacion'],
                        negrita_inicio=getattr(c, 'ETIQUETA_ACTIVIDAD_ANALISIS', 'Actividad de investigación'),
                    )
            else:
                agregar_item_lista(doc, i, actividad)

    # --- CAPÍTULO V: CONCLUSIONES Y RECOMENDACIONES ---
    if tiene_cap5:
        iniciar_capitulo(doc, "V", "CONCLUSIONES Y RECOMENDACIONES", bookmark_id="bm_cap5")
        agregar_titulo_nivel2(doc, "Conclusiones", bookmark_id="bm_cap5_concl")
        conclusiones = getattr(c, 'CONCLUSIONES', [])
        for i, conclusion in enumerate(conclusiones, 1):
            agregar_item_lista(doc, i, conclusion)

        agregar_titulo_nivel2(doc, "Recomendaciones", bookmark_id="bm_cap5_recom")
        recomendaciones = getattr(c, 'RECOMENDACIONES', [])
        for i, recomendacion in enumerate(recomendaciones, 1):
            agregar_item_lista(doc, i, recomendacion)

    # --- REFERENCIAS BIBLIOGRÁFICAS ---
    iniciar_seccion_preliminar(
        doc,
        "REFERENCIAS",
        bookmark_id="bm_referencias",
        ocultar_primera_pagina=True,
    )
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(24)
    p_sep.paragraph_format.space_after = Pt(0)
    referencias_lista = getattr(c, 'REFERENCIAS_LISTA', [])
    for ref in referencias_lista:
        agregar_referencia(doc, ref)

    # --- ANEXOS ---
    if hasattr(c, 'ANEXOS_LISTA') and c.ANEXOS_LISTA is not None:
        # Portadilla de ANEXOS (Art. 26: una hoja sola con la palabra ANEXOS centrada y en negrita)
        sec_portadilla = doc.add_section(WD_SECTION_START.NEW_PAGE)
        _config_seccion(sec_portadilla)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        alto_util_pt = (
            sec_portadilla.page_height
            - sec_portadilla.top_margin
            - sec_portadilla.bottom_margin
        ) / 12700
        p.paragraph_format.space_before = Pt(
            max(0, (alto_util_pt - TAMANO_PORTADILLA_ANEXOS * 1.2) / 2)
        )
        run_anexos_tit = p.add_run("ANEXOS")
        run_anexos_tit.font.name = FUENTE
        _agregar_bookmark(p, "bm_anexos")
        run_anexos_tit.font.size = Pt(TAMANO_PORTADILLA_ANEXOS)
        run_anexos_tit.font.bold = True

        # Anexos individuales (Art. 15: cada uno en página nueva, arriba y centrado, subtítulo entre corchetes)
        for anexo in c.ANEXOS_LISTA:
            cod, desc = anexo[:2]
            numero_imagen = anexo[2] if len(anexo) > 2 else None
            alto_imagen = anexo[3] if len(anexo) > 3 else None
            contenido_anexo = anexo[4] if len(anexo) > 4 else None
            sec_anexo = doc.add_section(WD_SECTION_START.NEW_PAGE)
            _config_seccion(sec_anexo)

            p_anexo = doc.add_paragraph()
            p_anexo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_anexo.paragraph_format.space_before = Pt(0)
            p_anexo.paragraph_format.space_after = ESP_DOBLE
            _mantener_con_siguiente(p_anexo)
            letra_anexo = cod.split(" ")[-1].upper()
            _agregar_bookmark(p_anexo, f"bm_anexo{letra_anexo}")
            
            # Nombre del Anexo (ej: ANEXO A)
            run_cod = p_anexo.add_run(cod.upper())
            run_cod.font.name = FUENTE
            run_cod.font.size = Pt(TAMANO_TABLA)
            run_cod.font.bold = True
            
            # Subtítulo del contenido centrado entre corchetes [ ]
            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.paragraph_format.space_before = ESP_SENCILLO
            p_sub.paragraph_format.space_after = ESP_DOBLE
            _mantener_con_siguiente(p_sub)
            
            run_desc = p_sub.add_run(f"[{desc}]")
            run_desc.font.name = FUENTE
            run_desc.font.size = Pt(TAMANO_TABLA)
            run_desc.font.bold = True
            
            if numero_imagen is not None:
                if isinstance(numero_imagen, (list, tuple)):
                    imagenes = numero_imagen
                else:
                    imagenes = [numero_imagen]

                for indice_imagen, imagen_cfg in enumerate(imagenes):
                    salto_pagina = indice_imagen > 0

                    if isinstance(imagen_cfg, dict):
                        referencia = imagen_cfg.get('archivo', imagen_cfg.get('referencia'))
                        configuracion = imagen_cfg.get('configuracion', imagen_cfg.get('config', alto_imagen))
                        titulo_fotografia = imagen_cfg.get('titulo')
                    else:
                        referencia = imagen_cfg
                        configuracion = alto_imagen
                        titulo_fotografia = None

                    ruta_imagen = buscar_imagen_por_referencia(carpeta_imagenes, referencia)
                    if not ruta_imagen:
                        continue

                    p_imagen = doc.add_paragraph()
                    p_imagen.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if salto_pagina:
                        # Aplicar el salto al párrafo de la imagen evita que
                        # LibreOffice cree una página intermedia vacía.
                        p_imagen.paragraph_format.page_break_before = True
                    _mantener_con_siguiente(p_imagen)
                    try:
                        tamano_anexo = _calcular_tamano_anexo_proporcional(
                            ruta_imagen, configuracion
                        )
                        p_imagen.add_run().add_picture(ruta_imagen, **tamano_anexo)
                    except Exception as e:
                        print(f"⚠ No se pudo insertar el anexo {cod} ({ruta_imagen}): {e}")
                        continue

                    if titulo_fotografia:
                        p_foto = doc.add_paragraph()
                        p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_foto.paragraph_format.space_before = Pt(6)
                        p_foto.paragraph_format.space_after = Pt(12)
                        run_foto = p_foto.add_run(titulo_fotografia)
                        run_foto.font.name = FUENTE
                        run_foto.font.size = Pt(TAMANO_TABLA)
                        run_foto.font.bold = True

            if contenido_anexo:
                bloques = contenido_anexo if isinstance(contenido_anexo, (list, tuple)) else [contenido_anexo]
                for bloque in bloques:
                    _agregar_fuente(doc, bloque, centrada=True)

    return idx_cap1, idx_indice_inicio, idx_indice_fin, idx_inicio_arabigo

# ================================================================
#  EJECUCIÓN PRINCIPAL
# ================================================================

def generar_reporte_completo(modo="completo"):
    modos_validos = {"completo", "borrador1", "borrador2", "borrador3", "borrador4"}
    if modo not in modos_validos:
        raise ValueError(f"Modo de generación inválido: {modo!r}")
    validar_contenido()
    print("» Inicializando documento...")
    doc = setup_iutecp_document()
    
    # 1. Portada (Sección 0) — solo autor, sin tutores
    construir_portada(doc, solo_autor=True, idx_seccion=0)
    
    # 2. Contraportada (Sección 1)
    sec_contra = doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec_contra.top_margin = Cm(3)
    sec_contra.bottom_margin = Cm(3)
    sec_contra.left_margin = Cm(4)
    sec_contra.right_margin = Cm(3)
    sec_contra.different_first_page_header_footer = False
    construir_portada(doc, idx_seccion=1, bookmark_id="bm_contraportada")

    # 3. Cuerpo (Sección 2 o 1 en borrador)
    idx_cap1, idx_indice_inicio, idx_indice_fin, idx_inicio_arabigo = construir_cuerpo_documento(doc, modo=modo)

    # La secuencia arábiga comienza en Introducción, no en el Capítulo I.
    agregar_numeracion_pie(doc, idx_inicio_arabigo=idx_inicio_arabigo)

    if modo == "borrador1":
        sufijo = "_BORRADOR1"
    elif modo == "borrador2":
        sufijo = "_BORRADOR2"
    elif modo == "borrador3":
        sufijo = "_BORRADOR3"
    elif modo == "borrador4":
        sufijo = "_BORRADOR4"
    else:
        sufijo = ""
    docx_output = os.path.join(BASE_DIR, f"Informe_Pasantia_IUTECP{sufijo}.docx")
    doc.save(docx_output)
    print(f"✔ Archivo Word generado: {docx_output}")

    print("» Renderizando PDF usando LibreOffice...")
    pdf_output = os.path.join(BASE_DIR, f"Informe_Pasantia_IUTECP{sufijo}.pdf")
    soffice_cmd = 'libreoffice'
    if platform.system() == 'Windows':
        possible_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for path in possible_paths:
            if os.path.exists(path):
                soffice_cmd = path
                break

    try:
        if os.path.exists(pdf_output):
            os.remove(pdf_output)
        subprocess.run(
            [soffice_cmd, '--headless', '--convert-to', 'pdf', '--outdir', BASE_DIR, docx_output],
            check=True,
            stdout=subprocess.DEVNULL,
        )
        if not os.path.isfile(pdf_output) or os.path.getsize(pdf_output) == 0:
            raise RuntimeError(f"LibreOffice no produjo el PDF esperado: {pdf_output}")
        print(f"✔ PDF generado y verificado: {pdf_output}")
    except FileNotFoundError:
        raise RuntimeError("LibreOffice no está instalado o no se encontró en el PATH.")
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"Error en la conversión a PDF: {e}") from e

if __name__ == "__main__":
    import sys
    modo = "completo"
    if "--modo" in sys.argv:
        idx = sys.argv.index("--modo")
        if idx + 1 < len(sys.argv):
            modo = sys.argv[idx + 1]
    generar_reporte_completo(modo=modo)
