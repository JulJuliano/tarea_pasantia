import os
import subprocess
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ================================================================
#  FUNCIONES DE CONFIGURACIÓN Y FORMATO (ESTÁTICAS - NORMA IUTECP)
# ================================================================

def setup_page(doc):
    """Márgenes Art. 6 (4cm Izquierdo para encuadernación, 3cm los demás)"""
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

def add_paragraph_normado(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)

    run_l = p.add_run(label)
    run_l.font.name = 'Times New Roman'
    run_l.font.size = Pt(12)
    run_l.font.bold = True

    run_v = p.add_run(value)
    run_v.font.name = 'Times New Roman'
    run_v.font.size = Pt(12)
    return p

def set_cell_format(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        borders.append(border)
    tblPr.append(borders)

# ================================================================
#  PROCESADOR CORE: COMPILA UNA SEMANA ESPECÍFICA
# ================================================================

def generar_documento_semana(datos):
    doc = Document()
    setup_page(doc)

    # 1. Título Centralizado
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_after = Pt(12)
    run_t = p_titulo.add_run("PLAN SEMANAL DE PASANTÍAS")
    run_t.font.name = 'Times New Roman'
    run_t.font.size = Pt(14)
    run_t.font.bold = True

    # 2. Datos Informativos de la Empresa y la Pasante
    add_label_value(doc, "EMPRESA: ", "Lubricantes y Equipos Varyna, C.A.")
    add_label_value(doc, "DEPARTAMENTO: ", "Departamento Administrativo / Área de Procura")
    add_label_value(doc, "ESPECIALIDAD: ", "Administración / Ciencias Comerciales")
    add_label_value(doc, "SEMANA N°: ", datos["num_semana"])
    add_label_value(doc, "PERÍODO: ", datos["periodo"])
    add_label_value(doc, "PASANTE: ", "Keidy Guzmán |  C.I.: 28.706.352")
#    add_label_value(doc, "TUTOR INDUSTRIAL: ", "[Nombre del Tutor]  |  C.I.: [Cédula]")

    doc.add_paragraph()

    # 3. Objetivo de la semana
    p_obj_t = doc.add_paragraph()
    p_obj_t.paragraph_format.space_after = Pt(4)
    run_ot = p_obj_t.add_run("OBJETIVO DE LA SEMANA:")
    run_ot.font.name = 'Times New Roman'
    run_ot.font.size = Pt(12)
    run_ot.font.bold = True

    add_paragraph_normado(doc, datos["objetivo"])

    doc.add_paragraph()

    # 4. Tabla de Actividades Planificadas (Dimensiones fijas para evitar desbordes)
    p_cuarto_t = doc.add_paragraph()
    p_cuarto_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cuarto_t.paragraph_format.space_after = Pt(6)
    run_ct = p_cuarto_t.add_run(f"CUADRO 1. Actividades Planificadas de la Semana {datos['num_semana']}")
    run_ct.font.name = 'Times New Roman'
    run_ct.font.size = Pt(10)
    run_ct.font.bold = True

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    table._tbl.tblPr.append(tblLayout)

    # Encabezados con Sombreado Azul Pastel
    headers = ["Nro", "Actividad", "Descripción", "Recursos"]
    for i, h in enumerate(headers):
        set_cell_format(table.rows[0].cells[i], h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), "D9E2F3")
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_elm)

    # Inyección de actividades semanales
    for nro, act, desc, rec in datos["actividades"]:
        row = table.add_row()
        set_cell_format(row.cells[0], nro, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_format(row.cells[1], act, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_format(row.cells[2], desc, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_cell_format(row.cells[3], rec, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Ancho total útil: 14.59 cm
    widths = [Cm(1.2), Cm(3.5), Cm(6.5), Cm(3.39)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.add_paragraph()

    # 5. Entregables
    p_ent_t = doc.add_paragraph()
    p_ent_t.paragraph_format.space_after = Pt(4)
    run_et = p_ent_t.add_run("ENTREGABLES DE LA SEMANA:")
    run_et.font.name = 'Times New Roman'
    run_et.font.size = Pt(12)
    run_et.font.bold = True

    add_paragraph_normado(doc, datos["entregables"])

    doc.add_paragraph()

    # 6. Observaciones del Tutor Industrial (Inyección cualitativa real)
    p_obs_t = doc.add_paragraph()
    p_obs_t.paragraph_format.space_after = Pt(4)
    run_ot2 = p_obs_t.add_run("OBSERVACIONES DEL TUTOR INDUSTRIAL:")
    run_ot2.font.name = 'Times New Roman'
    run_ot2.font.size = Pt(12)
    run_ot2.font.bold = True

    add_paragraph_normado(doc, datos["observaciones_tutor"])

#    p_obs = doc.add_paragraph()
#    p_obs.paragraph_format.line_spacing = 1.5
#    run_obs = p_obs.add_run("(Espacio reservado para la evaluación cualitativa del tutor sobre control de procesos, puntualidad, manejo de solicitudes y cumplimiento de directrices contables)")
#    run_obs.font.name = 'Times New Roman'
#    run_obs.font.size = Pt(11)
#    run_obs.font.italic = True

#    for _ in range(4):
#        p_line = doc.add_paragraph()
#        p_line.paragraph_format.space_before = Pt(8)
#        p_line.paragraph_format.space_after = Pt(0)
#        run_line = p_line.add_run("_" * 75)
#        run_line.font.name = 'Times New Roman'
#        run_line.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # 7. Firmas (Estructura de Tabla Invisible)
    sign_table = doc.add_table(rows=2, cols=2)
    remove_table_borders(sign_table)
    sign_table.allow_autofit = False
    for row in sign_table.rows:
        row.cells[0].width = Cm(7.29)
        row.cells[1].width = Cm(7.30)

    set_cell_format(sign_table.rows[0].cells[0], "_________________________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
#    set_cell_format(sign_table.rows[0].cells[1], "_________________________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_format(sign_table.rows[1].cells[0], "Firma del Pasante:\nKeidy Guzmán |  C.I.: 28.706.352", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
#    set_cell_format(sign_table.rows[1].cells[1], "Firma del Tutor Industrial:\n[Nombre del Tutor] | C.I.: [Cédula]", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # 8. Bloque Sello Corporativo
#    p_sello = doc.add_paragraph()
#    p_sello.alignment = WD_ALIGN_PARAGRAPH.CENTER
#    p_sello.paragraph_format.space_before = Pt(12)
#    run_s = p_sello.add_run("SELLO DE LA EMPRESA")
#    run_s.font.name = 'Times New Roman'
#    run_s.font.size = Pt(10)
#    run_s.font.bold = True

    output_dir = "cronogramas generados"
    os.makedirs(output_dir, exist_ok=True)
    fn_word = os.path.join(output_dir, f"Cronograma_Procura_Semana{datos['num_semana']}_IUTECP.docx")
    doc.save(fn_word)
    print(f"✔ Generado Word: {fn_word}")

    try:
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, fn_word], check=True, stdout=subprocess.DEVNULL)
        print(f"✔ Convertido PDF: {fn_word.replace('.docx', '.pdf')}")
    except Exception:
        print(f"⚠️ Word creado. LibreOffice headless no se ejecutó.")

# ================================================================
#  COLECCIÓN DE DATOS MAESTRA (CRONOGRAMA ADAPTADO A TSU - 10 SEMANAS)
# ================================================================

reportes_pasantias = [
    # ── SEMANA 1 ────────────────────────────────────────────────
    {
        "num_semana": "1",
        "periodo": "Del 01 al 05 de junio de 2026",
        "objetivo": "Formalizar el inicio de las pasantías dentro de la modalidad institucional en el puesto habitual de trabajo y delimitar los tiempos de investigación.",
        "entregables": "Acta de inicio de pasantías visada por la empresa y cronograma operativo validado por el tutor industrial.",
        "observaciones_tutor": "La trabajadora y ahora pasante Keidy Guzmán formalizó el inicio de su período académico en su puesto habitual del área de procura. Mantiene su excelente desempeño y puntualidad, coordinando eficientemente sus obligaciones rutinarias con el levantamiento del proyecto.",
        "actividades": [
            ("1", "Formalización del inicio", "Reunión con la gerencia general y el tutor industrial para consignar los formatos institucionales del IUTECP.", "Cartas de postulación, plan institucional"),
            ("2", "Alineación de objetivos", "Vincular las metas del proyecto técnico con las funciones contables y de compra que ya ejecuta en su día a día.", "Manual de procesos, PC de oficina"),
            ("3", "Establecimiento de tiempos", "Definir los bloques horarios de la jornada diaria dedicados exclusivamente a la recolección de datos y revisión de expedientes.", "Calendario laboral, agenda"),
            ("4", "Revisión metodológica", "Revisar los lineamientos del instituto con el tutor para asegurar el correcto resguardo confidencial de la información empresarial.", "Normativa IUTECP, documentos internos"),
            ("5", "Apertura del cuaderno de campo", "Organizar la libreta de registro técnico para documentar los eventos específicos observados durante la jornada.", "Cuaderno de apuntes, papelería")
        ]
    },
    # ── SEMANA 2 ────────────────────────────────────────────────
    {
        "num_semana": "2",
        "periodo": "Del 08 al 12 de junio de 2026",
        "objetivo": "Observar y registrar cómo se reciben y procesan actualmente las solicitudes de compra que llegan al departamento.",
        "entregables": "Lista de pasos detallada sobre el recorrido físico que hace una solicitud de compra desde que llega.",
        "observaciones_tutor": "La pasante recolectó la información solicitada con mucho orden. Identificó con facilidad la manera en que se archivan los papeles y se comunicó bien con el resto del personal.",
        "actividades": [
            ("1", "Seguimiento de solicitudes", "Observar el momento en que llegan las hojas de solicitud de materiales desde el área de operaciones.", "Formatos de requisición, block de notas"),
            ("2", "Registro de documentos", "Anotar la correspondencia de compras entrante en el libro de control diario del departamento.", "Libro diario, PC de oficina"),
            ("3", "Clasificación de papeles", "Separar los documentos entre compras de oficina comunes y pedidos de repuestos para maquinaria.", "Hojas de cálculo, archivos físicos"),
            ("4", "Revisión de carpetas", "Verificar el estado y orden de las carpetas físicas donde se guardan los datos de los proveedores.", "Archiveros del departamento"),
            ("5", "Mapeo preliminar de ruta", "Dibujar un esquema simple con el trayecto físico que sigue cada solicitud dentro de la oficina.", "Papel, lápiz, regla")
        ]
    },
    # ── SEMANA 3 ────────────────────────────────────────────────
    {
        "num_semana": "3",
        "periodo": "Del 15 al 19 de junio de 2026",
        "objetivo": "Preparar una lista de preguntas sencillas para conversar con el personal sobre los retrasos en las compras.",
        "entregables": "Guía escrita con las preguntas básicas para las entrevistas, revisada por el tutor industrial.",
        "observaciones_tutor": "Keidy preparó un cuestionario bastante claro y directo. Las preguntas están enfocadas en entender por qué se tardan en autorizar los presupuestos de compras.",
        "actividades": [
            ("1", "Identificación de fallas", "Anotar cuáles son los retrasos más comunes comentados por los compañeros de oficina.", "Cuaderno de apuntes"),
            ("2", "Redacción de preguntas", "Escribir un cuestionario corto sobre los problemas diarios con las firmas y los formatos.", "Papelería, procesador de palabras"),
            ("3", "Estructuración de la guía", "Organizar el orden de las preguntas para asegurar que el diálogo sea rápido y fluido.", "Borrador de guía, software de oficina"),
            ("4", "Aprobación de la guía", "Mostrar el cuestionario al tutor para confirmar que las preguntas sean las adecuadas para el personal.", "Formato borrador"),
            ("5", "Impresión de instrumentos", "Preparar las copias físicas definitivas del cuestionario para su aplicación en el área.", "Impresora, hojas base")
        ]
    },
    # ── SEMANA 4 ────────────────────────────────────────────────
    {
        "num_semana": "4",
        "periodo": "Del 22 al 26 de junio de 2026",
        "objetivo": "Conversar con el personal administrativo usando el cuestionario para anotar los inconvenientes del proceso.",
        "entregables": "Hojas de respuestas llenas con los comentarios y notas tomadas del personal de compras.",
        "observaciones_tutor": "La pasante realizó las actividades con mucho respeto y organización. Logró obtener información valiosa sobre las demoras en la recolección de firmas de la junta directiva.",
        "actividades": [
            ("1", "Coordinación de tiempos", "Acordar los momentos disponibles con los compañeros para realizar las entrevistas sin interrumpir.", "Calendario de oficina"),
            ("2", "Entrevistas al personal", "Hacer las preguntas preparadas al asistente de compras y al personal encargado de la administración.", "Cuestionarios impresos"),
            ("3", "Registro de quejas", "Tomar nota sobre los días que se tarda un papel en ser firmado o revisado.", "Libreta de notas, lápiz"),
            ("4", "Observación de firmas", "Anotar el camino que sigue una carpeta de compra cuando sale a aprobación final.", "Expedientes en trámite"),
            ("5", "Cotejo de respuestas", "Revisar los puntos en común donde los entrevistados coinciden sobre los cuellos de botella.", "Cuestionarios llenos")
        ]
    },
    # ── SEMANA 5 ────────────────────────────────────────────────
    {
        "num_semana": "5",
        "periodo": "Del 29 de junio al 03 de julio de 2026",
        "objetivo": "Transcribir las entrevistas y analizar las órdenes de compra anteriores para identificar los problemas principales.",
        "entregables": "Resumen de entrevistas y tabla de tiempos históricos de órdenes de compra del último trimestre.",
        "observaciones_tutor": "Buen trabajo de oficina. La pasante organizó las respuestas y el análisis de tiempos históricos, dejando en evidencia los cuellos de botella con ciertos proveedores.",
        "actividades": [
            ("1", "Transcripción de notas", "Pasar a la computadora las respuestas obtenidas en las hojas de las entrevistas.", "Computadora, Word"),
            ("2", "Acceso al archivo histórico", "Buscar e identificar las carpetas de órdenes de compra del último trimestre en el archivo físico.", "Archivo de documentos"),
            ("3", "Cálculo de tiempos", "Calcular los días exactos transcurridos entre la solicitud del material y su compra definitiva.", "Excel, cuaderno de notas"),
            ("4", "Agrupación de fallas", "Juntar en una lista las fallas similares como falta de cotizaciones o retraso de firmas.", "Software de oficina"),
            ("5", "Consolidación de resultados", "Estructurar una tabla de tiempos de respuesta históricos del proceso de compras.", "PC de oficina, Excel")
        ]
    },
    # ── SEMANA 6 ────────────────────────────────────────────────
    {
        "num_semana": "6",
        "periodo": "Del 06 al 10 de julio de 2026",
        "objetivo": "Estructurar las causas de las fallas del departamento en un diagrama de Ishikawa y fundamentar conceptualmente el proyecto.",
        "entregables": "Diagrama de Ishikawa digital y primer borrador del marco teórico del informe.",
        "observaciones_tutor": "El diagrama causa-efecto quedó muy claro y adaptado a nuestra oficina. El soporte teórico recopilado es pertinente con lo exigido para el nivel técnico.",
        "actividades": [
            ("1", "Clasificación de causas", "Dividir las causas de retrasos en las categorías de métodos, herramientas o factor humano.", "Block de notas"),
            ("2", "Estructuración del Ishikawa", "Diseñar en papel la estructura del diagrama de espina de pescado (causas y efectos).", "Lápiz, papel, regla"),
            ("3", "Digitalización del gráfico", "Pasar a limpio el diagrama de espina de pescado usando herramientas de la computadora.", "PC de oficina, Word"),
            ("4", "Lectura de manuales", "Revisar libros de administración y control de compras para fundamentar conceptualmente la propuesta.", "Manuales de administración, guías"),
            ("5", "Cierre de diagnóstico", "Redactar la conclusión final del diagnóstico situacional del departamento para el informe.", "Procesador de textos")
        ]
    },
    # ── SEMANA 7 ────────────────────────────────────────────────
    {
        "num_semana": "7",
        "periodo": "Del 13 al 17 de julio de 2026",
        "objetivo": "Diseñar un nuevo protocolo procedimental y recorrido optimizado para las solicitudes de compras.",
        "entregables": "Esquema preliminar del nuevo circuito de compras y asignación de tareas del área.",
        "observaciones_tutor": "La ruta propuesta por la pasante es sencilla. La propuesta de utilizar un buzón de correo único evitará que se pierdan las cotizaciones recibidas.",
        "actividades": [
            ("1", "Lluvia de ideas", "Esbozar alternativas para simplificar los pasos de entrega y firma de las requisiciones.", "Papel de reciclaje, lápiz"),
            ("2", "Diseño del circuito", "Escribir el paso a paso ideal desde la recepción de la requisición hasta la orden de compra final.", "Hojas blancas, lápiz"),
            ("3", "Asignación de roles", "Definir las responsabilidades del personal en cada fase del nuevo procedimiento de compras.", "Procesador de palabras"),
            ("4", "Propuesta de correo único", "Definir las reglas para centralizar el flujo de cotizaciones en una sola cuenta de correo.", "Computadora"),
            ("5", "Elaboración de flujograma", "Representar de forma gráfica la nueva secuencia de actividades administrativas diseñadas.", "PC de oficina, Word")
        ]
    },
    # ── SEMANA 8 ────────────────────────────────────────────────
    {
        "num_semana": "8",
        "periodo": "Del 20 al 24 de julio de 2026",
        "objetivo": "Crear planillas automatizadas en Excel para registrar y comparar de manera objetiva los precios de proveedores.",
        "entregables": "Plantilla de Excel para la comparación de tres cotizaciones simultáneas con cálculo automático de totales.",
        "observaciones_tutor": "Keidy demostró un excelente dominio de Excel. La planilla diseñada permite visualizar de manera rápida la opción más económica e idónea.",
        "actividades": [
            ("1", "Planificación del formato", "Definir los campos esenciales (proveedor, RIF, subtotal, IVA, total, entrega) a incluir en la planilla.", "Cuaderno de notas"),
            ("2", "Diseño de la estructura", "Crear las columnas y formato visual balanceado para comparar tres proveedores al mismo tiempo.", "Excel / Calc de oficina"),
            ("3", "Prueba de fórmulas", "Colocar fórmulas de sumatoria y cálculo de impuestos para automatizar los resultados financieros.", "Excel, funciones lógicas"),
            ("4", "Llenado de muestra", "Probar el funcionamiento de la tabla ingresando datos históricos de cotizaciones pasadas.", "Datos de prueba, PC"),
            ("5", "Ajustes de visualización", "Aplicar colores de destaque y bordes legibles para facilitar la comprensión de la tabla por el personal.", "Estilos de celda")
        ]
    },
    # ── SEMANA 9 ────────────────────────────────────────────────
    {
        "num_semana": "9",
        "periodo": "Del 27 al 31 de julio de 2026",
        "objetivo": "Probar y validar el uso de la planilla de Excel y el nuevo flujo operativo con los compañeros del departamento.",
        "entregables": "Registro de retroalimentación de los usuarios y formato Excel definitivo con correcciones aplicadas.",
        "observaciones_tutor": "Pusimos a prueba la planilla y el nuevo flujo. El personal se adaptó muy bien, y logramos reducir significativamente el tiempo de análisis de ofertas.",
        "actividades": [
            ("1", "Explicación del recurso", "Presentar y explicar el funcionamiento del archivo de Excel al asistente administrativo de compras.", "Estación de trabajo"),
            ("2", "Prueba en vivo", "Acompañar al asistente en la carga de una cotización real usando la nueva planilla.", "Planilla de Excel, PC de oficina"),
            ("3", "Registro de sugerencias", "Anotar sugerencias del personal para mejorar la comodidad de uso de la herramienta.", "Papel y lápiz"),
            ("4", "Ajuste de celdas", "Realizar las correcciones y bloqueos de celdas sugeridos para evitar errores de tipeo accidentales.", "Excel, protección de hojas"),
            ("5", "Validación del proceso", "Confirmar que la nueva metodología agiliza la toma de decisiones sobre las adquisiciones de la semana.", "Formato Excel final")
        ]
    },
    # ── SEMANA 10 ───────────────────────────────────────────────
    {
        "num_semana": "10",
        "periodo": "Del 03 al 07 de agosto de 2026",
        "objetivo": "Consolidar el informe final del proyecto de pasantías y presentarlo ante la gerencia de la empresa.",
        "entregables": "Informe final de pasantías con anexos listos para consignar ante el IUTECP.",
        "observaciones_tutor": "La pasante Keidy Guzmán culminó exitosamente su proyecto. La propuesta de ordenamiento y formatos de compras constituye un gran aporte para el control interno. Proyecto aprobado.",
        "actividades": [
            ("1", "Integración del informe", "Compilar los capítulos, diagnósticos, diagramas y anexos de formatos en el documento definitivo.", "Word, Normas IUTECP"),
            ("2", "Revisión de estilo", "Verificar la ortografía, referencias, sangrías y márgenes del informe final.", "PC de oficina"),
            ("3", "Firma de avales", "Coordinar con el tutor industrial la firma y sellado de las planillas de evaluación y culminación.", "Documentos impresos, bolígrafos"),
            ("4", "Resguardo de archivos", "Guardar copias digitales de los informes y del formato de Excel en un dispositivo de almacenamiento.", "Pendrive, carpetas compartidas"),
            ("5", "Exposición de cierre", "Explicar a la gerencia de la empresa cómo quedan instaurados los nuevos formatos y el protocolo del departamento.", "Formatos finales, oficina de gerencia")
        ]
    }
]

# ================================================================
#  EJECUCIÓN GENERAL
# ================================================================

if __name__ == "__main__":
    print("🚀 Iniciando motor de automatización de cronogramas para Procura (Keidy)...")
    for idx, reporte in enumerate(reportes_pasantias):
        print(f"\n[Compilando {idx+1}/{len(reportes_pasantias)}] Semana {reporte['num_semana']}...")
        generar_documento_semana(reporte)
    print("\n✨ ¡Proceso culminado con éxito! Se han generado los 10 archivos independientes en Word y PDF.")
