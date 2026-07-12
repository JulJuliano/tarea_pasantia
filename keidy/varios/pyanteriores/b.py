import os
import subprocess
import platform
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ================================================================
#  FUNCIONES AUXILIARES DE FORMATO (NORMATIVA IUTECP 2025)
# ================================================================

def setup_iutecp_document():
    """Configura el documento base: Tamaño Carta, Márgenes Art. 6, Fuente Art. 5"""
    doc = Document()

    # Configuración inicial de la primera sección
    section = doc.sections[0]
    section.page_width = Cm(21.59)   # Tamaño Carta
    section.page_height = Cm(27.94)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)      # Encuadernación
    section.right_margin = Cm(3)

    # Estilo base Normal
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)

    return doc

def agregar_parrafo_normado(doc, texto, cursiva=False, sangria=True):
    """Párrafo justificado, 1.5 interlineado, sangría 1.25cm (Art. 7, 8)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.first_line_indent = Cm(1.25) if sangria else Cm(0)
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)

    run = p.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = cursiva
    return p

def agregar_item_lista(doc, numero, texto, negrita_inicio=""):
    """Agrega un elemento enumerado (1., 2.) con sangría, sin viñetas (bullets)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.first_line_indent = Cm(1.25)

    run_num = p.add_run(f"{numero}. ")
    run_num.font.name = 'Times New Roman'
    run_num.font.size = Pt(12)

    if negrita_inicio:
        run_bold = p.add_run(f"{negrita_inicio}: ")
        run_bold.font.name = 'Times New Roman'
        run_bold.font.size = Pt(12)
        run_bold.font.bold = True

    run_text = p.add_run(texto)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)
    return p

def agregar_titulo_nivel2(doc, texto):
    """Nivel 2: Centrado, Negrita, sin numeración (Art. 9)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    p.paragraph_format.line_spacing = 1.5

    run = p.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    return p

def agregar_titulo_nivel3(doc, texto):
    """Nivel 3: Justificado, Negrita, sin centrar"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.25)

    run = p.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    return p

def iniciar_capitulo(doc, numero_romano, titulo):
    """
    Crea un nuevo capítulo con margen superior de 5cm en la primera página (Art. 6),
    suprime el número de página y restaura el margen a 3cm.
    """
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec.top_margin = Cm(5)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(4)
    sec.right_margin = Cm(3)
    sec.different_first_page_header_footer = True

    # Título CAPÍTULO X
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(24)
    r1 = p1.add_run(f"CAPÍTULO {numero_romano}")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.font.bold = True

    # Título del Capítulo
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(24)
    r2 = p2.add_run(titulo.upper())
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2.font.bold = True

    # Salto continuo para restaurar margen a 3cm en el resto del capítulo
    sec2 = doc.add_section(WD_SECTION_START.CONTINUOUS)
    sec2.top_margin = Cm(3)
    sec2.bottom_margin = Cm(3)
    sec2.left_margin = Cm(4)
    sec2.right_margin = Cm(3)

def iniciar_seccion_preliminar(doc, titulo):
    """Para páginas preliminares que inician en página nueva (Art. 12)"""
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec.top_margin = Cm(5)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(4)
    sec.right_margin = Cm(3)
    sec.different_first_page_header_footer = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(titulo.upper())
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = True

    # Restaurar margen
    sec2 = doc.add_section(WD_SECTION_START.CONTINUOUS)
    sec2.top_margin = Cm(3)
    sec2.bottom_margin = Cm(3)
    sec2.left_margin = Cm(4)
    sec2.right_margin = Cm(3)

def agregar_cita_larga(doc, texto, cita):
    """Cita > 40 palabras: sangría 1.25cm ambos lados, espacio sencillo (Art. 22)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.right_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = p.add_run(f"{texto} {cita}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def agregar_referencia(doc, texto):
    """Referencia con sangría francesa 0.75cm, espacio sencillo (Art. 7, 25)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(24)

    run = p.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def agregar_numeracion_pie(doc):
    """Añade el número de página centrado en el pie de página (Art. 11, 12)"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()

        # Estructura XML correcta para campos complejos (PAGE)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        # CORRECCIÓN CRÍTICA: Falta el marcador 'separate'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# ================================================================
#  FUNCIONES MODULARES PARA IMÁGENES
# ================================================================

def buscar_imagen_por_numero(carpeta, numero, extensiones=None):
    """
    Busca una imagen por número en una carpeta específica.
    Busca archivos como: "1.png", "1.jpg", "1.jpeg", "2.png", etc.
    """
    if extensiones is None:
        extensiones = ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']

    if not os.path.exists(carpeta):
        print(f"⚠ Advertencia: La carpeta '{carpeta}' no existe")
        return None

    numero_str = str(numero)

    for archivo in os.listdir(carpeta):
        nombre, ext = os.path.splitext(archivo)
        if nombre == numero_str and ext.lower()[1:] in extensiones:
            ruta_completa = os.path.join(carpeta, archivo)
            print(f"✓ Imagen encontrada: {ruta_completa}")
            return ruta_completa

    print(f"⚠ No se encontró imagen con número {numero} en {carpeta}")
    return None

def agregar_imagen(doc, ruta_imagen, titulo, ancho=Cm(15)):
    """
    Agrega una imagen al documento con su título/crédito centrado.
    """
    if not ruta_imagen or not os.path.exists(ruta_imagen):
        print(f"⚠ No se puede agregar imagen: {ruta_imagen}")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(f"[IMAGEN NO ENCONTRADA: {titulo}]")
        run.font.italic = True
        return

    # CORRECCIÓN: Evitar que la imagen rompa los márgenes
    section = doc.sections[-1]
    max_width = section.page_width - section.left_margin - section.right_margin
    if ancho > max_width:
        print(f"⚠ Ajustando ancho de imagen '{titulo}' para que no exceda los márgenes.")
        ancho = max_width

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(ruta_imagen, width=ancho)
        print(f"✓ Imagen agregada: {os.path.basename(ruta_imagen)}")
    except Exception as e:
        print(f"❌ Error al agregar imagen {ruta_imagen}: {e}")
        return

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_before = Pt(12)
    p_titulo.paragraph_format.space_after = Pt(24)
    run_titulo = p_titulo.add_run(titulo)
    run_titulo.font.name = 'Times New Roman'
    run_titulo.font.size = Pt(12)
    run_titulo.font.bold = True

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_after = Pt(24)

# ================================================================
#  CONSTRUCCIÓN DEL DOCUMENTO (SIN CAMBIOS EN EL CONTENIDO)
# ================================================================

def construir_portada(doc):
    """Anexo I: Portada - Con espacios calculados por porcentaje"""
    # Calcular altura útil de la página
    section = doc.sections[0]
    altura_pagina_pt = section.page_height.pt  # 794.64 Pt
    margen_sup_pt = section.top_margin.pt      # 85.04 Pt
    margen_inf_pt = section.bottom_margin.pt   # 85.04 Pt
    altura_util_pt = altura_pagina_pt - margen_sup_pt - margen_inf_pt  # ~624.56 Pt

    # Distribución porcentual del contenido
    porcentaje_membrete = 0.12   # 12% para membretes
    porcentaje_espacio1 = 0.18   # 18% espacio antes del título (IGUAL)
    porcentaje_titulo = 0.10     # 10% para el título
    porcentaje_espacio2 = 0.18   # 18% espacio antes de datos (IGUAL)
    porcentaje_datos = 0.18      # 18% para datos del autor
    porcentaje_espacio3 = 0.08   # 8% espacio antes de fecha

    # Membretes superiores
    p_memb = doc.add_paragraph()
    p_memb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_memb.paragraph_format.line_spacing = 1.15
    p_memb.paragraph_format.space_after = Pt(0)
    p_memb.paragraph_format.space_before = Pt(altura_util_pt * porcentaje_membrete * 0.3)

    lineas_membrete = [
        "REPÚBLICA BOLIVARIANA DE VENEZUELA",
        "MINISTERIO DEL PODER POPULAR PARA LA EDUCACIÓN UNIVERSITARIA",
        "INSTITUTO UNIVERSITARIO DE TECNOLOGÍA",
        "\"ELÍAS CALIXTO POMPA\" (IUTECP)",
        "EL TIGRE, ESTADO ANZOÁTEGUI"
    ]

    for i, linea in enumerate(lineas_membrete):
        r = p_memb.add_run(linea)
        r.font.bold = True
        r.font.size = Pt(12)
        if i < len(lineas_membrete) - 1:
            r.add_break()

    # Espacio antes del título
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_before = Pt(altura_util_pt * porcentaje_espacio1)
    p_titulo.paragraph_format.space_after = Pt(0)

    r_title = p_titulo.add_run("DISEÑO DE UN SISTEMA PARA EL CONTROL, TRAZABILIDAD Y REPORTE DE MOVIMIENTOS DOCUMENTALES EN LA PRESIDENCIA DE LA EMPRESA MIXTA PETROLERA VENANGOCUPET, S.A.")
    r_title.font.bold = True
    r_title.font.size = Pt(12)

    # Espacio antes de los datos
    p_datos = doc.add_paragraph()
    p_datos.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_datos.paragraph_format.space_before = Pt(altura_util_pt * porcentaje_espacio2)
    p_datos.paragraph_format.space_after = Pt(0)

    r_datos = p_datos.add_run("Autor: Cardona, Juliano\nC.I.: 32.281.199\n\nTutor Industrial: Ing. Yasmin Sabaneta\nC.I.: 14.187.924\n\nTutor Académico: Lic. Carlos Mendoza")
    r_datos.font.bold = True

    # Espacio antes de la fecha
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie.paragraph_format.space_before = Pt(altura_util_pt * porcentaje_espacio3)
    p_pie.paragraph_format.space_after = Pt(0)

    r_pie = p_pie.add_run("El Tigre, julio de 2026")
    r_pie.font.size = Pt(12)

def construir_contraportada(doc):
    """Contraportada - IDÉNTICA a la portada con espacios por porcentaje"""
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    # Calcular altura útil de la página
    section = doc.sections[-1]  # Última sección (la nueva)
    altura_pagina_pt = section.page_height.pt
    margen_sup_pt = section.top_margin.pt
    margen_inf_pt = section.bottom_margin.pt
    altura_util_pt = altura_pagina_pt - margen_sup_pt - margen_inf_pt

    # Distribución porcentual del contenido
    porcentaje_membrete = 0.12
    porcentaje_espacio1 = 0.18   # IGUAL
    porcentaje_titulo = 0.10
    porcentaje_espacio2 = 0.18   # IGUAL
    porcentaje_datos = 0.18
    porcentaje_espacio3 = 0.08

    # Membretes superiores
    p_memb = doc.add_paragraph()
    p_memb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_memb.paragraph_format.line_spacing = 1.15
    p_memb.paragraph_format.space_after = Pt(0)
    p_memb.paragraph_format.space_before = Pt(altura_util_pt * porcentaje_membrete * 0.3)

    lineas_membrete = [
        "REPÚBLICA BOLIVARIANA DE VENEZUELA",
        "MINISTERIO DEL PODER POPULAR PARA LA EDUCACIÓN UNIVERSITARIA",
        "INSTITUTO UNIVERSITARIO DE TECNOLOGÍA",
        "\"ELÍAS CALIXTO POMPA\" (IUTECP)",
        "EL TIGRE, ESTADO ANZOÁTEGUI"
    ]

    for i, linea in enumerate(lineas_membrete):
        r = p_memb.add_run(linea)
        r.font.bold = True
        r.font.size = Pt(12)
        if i < len(lineas_membrete) - 1:
            r.add_break()

    # Espacio antes del título
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_before = Pt(altura_util_pt * porcentaje_espacio1)
    p_titulo.paragraph_format.space_after = Pt(0)

    r_title = p_titulo.add_run("DISEÑO DE UN SISTEMA PARA EL CONTROL, TRAZABILIDAD Y REPORTE DE MOVIMIENTOS DOCUMENTALES EN LA PRESIDENCIA DE LA EMPRESA MIXTA PETROLERA VENANGOCUPET, S.A.")
    r_title.font.bold = True
    r_title.font.size = Pt(12)

    # Espacio antes de los datos
    p_datos = doc.add_paragraph()
    p_datos.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_datos.paragraph_format.space_before = Pt(altura_util_pt * porcentaje_espacio2)
    p_datos.paragraph_format.space_after = Pt(0)

    r_datos = p_datos.add_run("Autor: Cardona, Juliano\nC.I.: 32.281.199\n\nTutor Industrial: Ing. Yasmin Sabaneta\nC.I.: 14.187.924\n\nTutor Académico: Lic. Carlos Mendoza")
    r_datos.font.bold = True

    # Espacio antes de la fecha
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie.paragraph_format.space_before = Pt(altura_util_pt * porcentaje_espacio3)
    p_pie.paragraph_format.space_after = Pt(0)

    r_pie = p_pie.add_run("El Tigre, julio de 2026")
    r_pie.font.size = Pt(12)

def construir_paginas_preliminares(doc):
    """Anexos J, K, L, M, N, O, P, Q, R"""
    # Contraportada (ahora usa la función dedicada)
    construir_contraportada(doc)

    # Actas de Aprobación
    for titulo in ["APROBACIÓN DEL TUTOR INDUSTRIAL (Anexo K)", "APROBACIÓN DEL TUTOR ACADÉMICO (Anexo L)"]:
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(titulo).font.bold = True
        agregar_parrafo_normado(doc, "En mi carácter de tutor... manifiesto que cumple con los requisitos exigidos por el IUTECP...")

    # Agradecimientos y Dedicatoria
    for titulo in ["AGRADECIMIENTOS", "DEDICATORIA"]:
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(titulo).font.bold = True
        agregar_parrafo_normado(doc, "Texto opcional de máximo una hoja...")

    # Índices y Listas
    for titulo in ["ÍNDICE GENERAL", "LISTA DE CUADROS", "LISTA DE FIGURAS", "LISTA DE GRÁFICOS", "LISTA DE ANEXOS"]:
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(titulo).font.bold = True
        agregar_parrafo_normado(doc, "(Insertar tabla de índice según formato de los Anexos M-Q)")

    # Resumen
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("RESUMEN").font.bold = True
    p_res = doc.add_paragraph()
    p_res.paragraph_format.line_spacing = 1.0
    p_res.add_run("Texto del resumen en espacio sencillo, máximo 300 palabras...")
    p_kw = doc.add_paragraph()
    p_kw.paragraph_format.space_before = Pt(12)
    r_kw = p_kw.add_run("Palabras claves: ")
    r_kw.font.bold = True
    p_kw.add_run("sistema, trazabilidad, petrolera, IUTECP, pasantías.")

def construir_cuerpo_documento(doc, carpeta_imagenes="imagenes"):
    """Desde Introducción hasta Anexos - CONTENIDO ACTUALIZADO"""

    # INTRODUCCIÓN
    iniciar_seccion_preliminar(doc, "INTRODUCCIÓN")
    agregar_parrafo_normado(doc, "La presente pasantía realizada en Lubricantes y Equipos Varyna, C.A., tiene como propósito desarrollar un sistema para el control, trazabilidad y reporte de movimientos documentales en la presidencia de la empresa, contribuyendo al fortalecimiento de los procesos administrativos y operativos de la organización.")

    # CAPÍTULO I - FASE I: REALIDAD ORGANIZACIONAL
    iniciar_capitulo(doc, "I", "REALIDAD ORGANIZACIONAL")

    agregar_titulo_nivel2(doc, "IDENTIFICACIÓN DE LA EMPRESA")

    # 1.1.1 Razón social
    agregar_titulo_nivel3(doc, "1.1.1 Razón social")
    agregar_parrafo_normado(doc, "Lubricantes y Equipos Varyna, C.A.", sangria=False)

    # 1.1.2 Reseña histórica
    agregar_titulo_nivel3(doc, "1.1.2 Reseña histórica")
    agregar_parrafo_normado(doc, "Lubricantes y Equipos Varyná C.A. es una empresa venezolana con más de treinta y seis (36) años de trayectoria en el sector petrolero, industrial y de construcción, consolidándose como una organización de amplia experiencia y reconocimiento dentro del mercado nacional. Desde sus inicios, la empresa ha orientando sus esfuerzos al desarrollo de soluciones integrales que contribuyan al fortalecimiento de las actividades productivas del país, especialmente en áreas estratégicas vinculadas con la industria energética.")
    agregar_parrafo_normado(doc, "Como empresa integrante y pilar fundamental del Grupo Corporativo VTC, Lubricantes y Equipos Varyná C.A. ha mantenido un crecimiento sostenido basado en la innovación, la calidad de sus servicios y el compromiso con sus clientes. Su actividad principal se enfoca en el procesamiento y suministro de productos químicos especializados, el tratamiento de crudo y la provisión de maquinaria pesada para operaciones industriales y petroleras.")
    agregar_parrafo_normado(doc, "A lo largo de su trayectoria, la organización ha participado activamente en proyectos destinados al crecimiento, recuperación y mantenimiento de la producción de pozos petroleros, aportando recursos técnicos y operativos que contribuyen al desarrollo del sector energético venezolano. Asimismo, ha fortalecido su presencia en la región oriental del país, donde mantiene una sólida base de operaciones y una participación constante en ferias y eventos empresariales de carácter regional y nacional.")
    agregar_parrafo_normado(doc, "En la actualidad, Lubricantes y Equipos Varyná C.A. continúa posicionándose como una empresa comprometida con la excelencia, la eficiencia operativa y el desarrollo sostenible, ofreciendo productos y servicios que responden a las necesidades de los sectores petrolero, industrial y de construcción, contribuyendo de manera significativa al progreso económico de Venezuela.")

    # 1.1.3 Misión
    agregar_titulo_nivel3(doc, "1.1.3 Misión")
    agregar_parrafo_normado(doc, "Brindar a nuestros clientes objetivos, soluciones de calidad en las áreas en las cuales nos desempeñamos, para contribuir de manera significativa en sus resultados. Aportando valor con nuestras respuestas a sus requerimientos. En la búsqueda de un mejor país y de una mejor humanidad.", cursiva=True)

    # 1.1.4 Visión
    agregar_titulo_nivel3(doc, "1.1.4 Visión")
    agregar_parrafo_normado(doc, "Ser el conglomerado de empresas líderes en cada una de las categorías en las que participamos, generando modelos de negocios altamente competitivos, atendiendo a nuestros distintos beneficiarios con productos y servicios de calidad, a la vez que contribuimos con el desarrollo social y ambiental del planeta. Buscando cubrir y trascender los mercados locales, teniendo presencia importante en el mercado global.", cursiva=True)

    # 1.1.5 Valores
    agregar_titulo_nivel3(doc, "1.1.5 Valores")
    agregar_parrafo_normado(doc, "Los valores que orientan las actividades de Lubricantes y Equipos Varyná C.A. están enfocados en el compromiso con la calidad, la excelencia y la satisfacción de sus clientes. Entre los principales valores de la organización destacan:")

    valores = [
        ("Responsabilidad", "Cumplimiento eficiente de los compromisos adquiridos con clientes, proveedores y trabajadores."),
        ("Integridad", "Actuar con honestidad, ética y transparencia en todas las operaciones de la empresa."),
        ("Calidad", "Garantizar productos y servicios que satisfagan las necesidades y expectativas de los clientes."),
        ("Compromiso", "Mantener una actitud de dedicación y esfuerzo para alcanzar los objetivos organizacionales."),
        ("Trabajo en equipo", "Fomentar la cooperación y el respeto entre los trabajadores para lograr resultados exitosos."),
        ("Innovación", "Promover la mejora continua de los procesos, productos y servicios ofrecidos."),
        ("Seguridad", "Velar por la protección de los trabajadores, las instalaciones y el medio ambiente en cada actividad realizada.")
    ]

    for i, (valor, descripcion) in enumerate(valores, 1):
        agregar_item_lista(doc, i, descripcion, valor)

    # 1.1.6 Objetivos Organizacionales
    agregar_titulo_nivel3(doc, "1.1.6 Objetivos Organizacionales")
    agregar_parrafo_normado(doc, "Lubricantes y Equipos Varyná C.A. tiene como objetivo principal contribuir al desarrollo de los sectores petrolero, industrial y de construcción mediante la prestación de servicios especializados y el suministro de productos de alta calidad. Entre sus objetivos organizacionales se encuentran:")

    objetivos = [
        "Proporcionar soluciones eficientes y oportunas a las necesidades de sus clientes.",
        "Mantener altos estándares de calidad en los productos y servicios ofrecidos.",
        "Impulsar el crecimiento y fortalecimiento de la industria petrolera nacional.",
        "Garantizar la satisfacción de los clientes mediante la mejora continua de los procesos.",
        "Promover el desarrollo profesional y personal de sus trabajadores.",
        "Contribuir al desarrollo económico, social y ambiental del país mediante una gestión responsable y sostenible."
    ]

    for i, objetivo in enumerate(objetivos, 1):
        agregar_item_lista(doc, i, objetivo)

    # 1.1.7 Ubicación geográfica
    agregar_titulo_nivel3(doc, "1.1.7 Ubicación geográfica")
    agregar_parrafo_normado(doc, "Calle 23 de enero entre calle principal el palomar y calle la paz sector vista al sol, San José de Guanipa Edo. Anzoátegui.", sangria=False)

    # Agregar Gráfico 1
    ruta_imagen_1 = buscar_imagen_por_numero(carpeta_imagenes, 1)
    agregar_imagen(doc, ruta_imagen_1, "Gráfico 1. Representación cartográfica y ubicación espacial de la empresa.")

    # 1.1.8 Población de los trabajadores de la empresa
    agregar_titulo_nivel3(doc, "1.1.8 Población de los trabajadores de la empresa")
    agregar_parrafo_normado(doc, "La población de trabajadores de Lubricantes y Equipos Varyná C.A. está conformada por personal administrativo, técnico y operativo, quienes desempeñan funciones esenciales para el cumplimiento de las actividades de la organización. El talento humano constituye uno de los recursos más importantes de la empresa, ya que contribuye al desarrollo eficiente de los procesos relacionados con el sector petrolero, industrial y de construcción.")
    agregar_parrafo_normado(doc, "La estructura laboral se encuentra distribuida en diferentes áreas funcionales, entre las que destacan administración, operaciones, logística, mantenimiento, seguridad industrial y atención al cliente. Todo el personal trabaja de manera coordinada para garantizar la calidad de los servicios prestados y el logro de los objetivos organizacionales.")

    # 1.1.9 Estructura organizacional de la empresa
    agregar_titulo_nivel3(doc, "1.1.9 Estructura organizacional de la empresa")

    # Agregar Gráfico 2
    ruta_imagen_2 = buscar_imagen_por_numero(carpeta_imagenes, 2)
    agregar_imagen(doc, ruta_imagen_2, "Gráfico 2. Organigrama estructural y niveles jerárquicos de la organización.")

    # Resto del documento (placeholders)
    iniciar_capitulo(doc, "II", "DIAGNÓSTICO SITUACIONAL")
    agregar_titulo_nivel2(doc, "Identificación de la Situación Problemática")
    agregar_parrafo_normado(doc, "Durante las actividades en la Presidencia de Lubricantes y Equipos Varyna, C.A., se detectó la necesidad de implementar un sistema que permita el control, trazabilidad y reporte de movimientos documentales de manera eficiente...")

    iniciar_capitulo(doc, "III", "MARCO TEÓRICO")
    agregar_titulo_nivel2(doc, "Bases Teóricas Referenciales")
    agregar_parrafo_normado(doc, "Según Pérez (2020), la trazabilidad documental es fundamental para el control eficiente de los procesos administrativos en las organizaciones...")

    iniciar_capitulo(doc, "IV", "ACTIVIDADES REALIZADAS")
    agregar_titulo_nivel2(doc, "Descripción de Actividades Ejecutadas por Semana")
    agregar_parrafo_normado(doc, "Semana 1: Inducción y levantamiento de requerimientos en Lubricantes y Equipos Varyna, C.A....")

    iniciar_capitulo(doc, "V", "CONCLUSIONES Y RECOMENDACIONES")
    agregar_titulo_nivel2(doc, "Conclusiones")
    agregar_parrafo_normado(doc, "Se logró diseñar el sistema cumpliendo con los objetivos planteados para Lubricantes y Equipos Varyna, C.A....")
    agregar_titulo_nivel2(doc, "Recomendaciones")
    agregar_parrafo_normado(doc, "Se recomienda a la empresa capacitar al personal en el nuevo sistema de control documental...")

    # REFERENCIAS
    iniciar_seccion_preliminar(doc, "REFERENCIAS")
    agregar_referencia(doc, "Asamblea Nacional. (1999). Constitución de la República Bolivariana de Venezuela. Caracas, La Torre.")
    agregar_referencia(doc, "Lubricantes y Equipos Varyná C.A. (2026). Manual de Normas y Procedimientos Internos. San José de Guanipa, Anzoátegui.")
    agregar_referencia(doc, "Pérez, L. (2020). Trazabilidad de Movimientos Documentales. Universidad Pedagógica Experimental Libertador.")

    # ANEXOS
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ANEXOS").font.bold = True
    agregar_parrafo_normado(doc, "ANEXO A\n[Definición de Términos Básicos]", sangria=False)
    agregar_parrafo_normado(doc, "ANEXO B\n[Planes de Trabajo]", sangria=False)
    agregar_parrafo_normado(doc, "ANEXO C\n[Memoria Fotográfica]", sangria=False)


# ================================================================
#  EJECUCIÓN PRINCIPAL
# ================================================================

def generar_reporte_completo(carpeta_imagenes="imagenes"):
    print("» Inicializando documento bajo Normas IUTECP 2025...")
    doc = setup_iutecp_document()

    # 1. Portada
    construir_portada(doc)

    # 2. Preliminares (incluye contraportada)
    construir_paginas_preliminares(doc)

    # 3. Cuerpo del documento
    construir_cuerpo_documento(doc, carpeta_imagenes=carpeta_imagenes)

    # 4. Añadir numeración de página al pie
    agregar_numeracion_pie(doc)

    # Guardar Word
    docx_output = "Informe_Pasantia_Varyna_IUTECP.docx"
    doc.save(docx_output)
    print(f"✔ Archivo Word generado: {docx_output}")

    # 5. Conversión a PDF
    print("» Renderizando PDF usando LibreOffice...")

    # MEJORA: Detección automática del ejecutable de LibreOffice (soffice) en Windows
    soffice_cmd = 'libreoffice'
    if platform.system() == 'Windows':
        possible_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for path in possible_paths:
            if os.path.exists(path):
                soffice_cmd = path
                break
        else:
            soffice_cmd = 'soffice' # Fallback al PATH

    try:
        subprocess.run([soffice_cmd, '--headless', '--convert-to', 'pdf', docx_output], check=True, stdout=subprocess.DEVNULL)
        print("✔ ¡PDF académico generado con éxito!")
    except FileNotFoundError:
        print("❌ Error: LibreOffice no está instalado o no se encontró la ruta. Abre el Word y guárdalo como PDF manualmente.")
    except subprocess.CalledProcessError as e:
        print(f"❌ Error en la conversión a PDF: {e}")

if __name__ == "__main__":
    CARPETA_IMAGENES = "imagenes"
    generar_reporte_completo(carpeta_imagenes=CARPETA_IMAGENES)
