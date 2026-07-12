import os
import subprocess
import platform
from docx import Document
from docx.shared import Cm, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Importamos el contenido
import contenido as c

# ================================================================
# FUNCIONES AUXILIARES DE FORMATO
# ================================================================
def setup_iutecp_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    return doc

def agregar_parrafo_normado(doc, texto, cursiva=False, sangria=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.first_line_indent = Cm(1.25) if sangria else Cm(0)
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)

    run = p.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = cursiva
    return p

def agregar_item_lista(doc, numero, texto, negrita_inicio=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.first_line_indent = Cm(1.25)

    run_num = p.add_run(f"{numero}. ")
    run_num.font.name = 'Times New Roman'
    run_num.font.size = Pt(12)

    if negrita_inicio:
        run_bold = p.add_run(f"{negrita_inicio}: ")
        run_bold.font.name = 'Times New Roman'
        run_bold.font.size = Pt(12)
        run_bold.font.bold = True

    run_text = p.add_run(texto)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)
    return p

def agregar_titulo_nivel2(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    p.paragraph_format.line_spacing = 1.5

    run = p.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    return p

def agregar_titulo_nivel3(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.25)

    run = p.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    return p

def iniciar_capitulo(doc, numero_romano, titulo):
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec.top_margin = Cm(5)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(4)
    sec.right_margin = Cm(3)
    sec.different_first_page_header_footer = True

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(24)
    r1 = p1.add_run(f"CAPÍTULO {numero_romano}")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.font.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(24)
    r2 = p2.add_run(titulo.upper())
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2.font.bold = True

    sec2 = doc.add_section(WD_SECTION_START.CONTINUOUS)
    sec2.top_margin = Cm(3)
    sec2.bottom_margin = Cm(3)
    sec2.left_margin = Cm(4)
    sec2.right_margin = Cm(3)

# ================================================================
# IMÁGENES
# ================================================================
def buscar_imagen_por_numero(carpeta, numero, extensiones=None):
    if extensiones is None:
        extensiones = ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']

    if not os.path.exists(carpeta):
        print(f"⚠ Advertencia: La carpeta '{carpeta}' no existe")
        return None

    numero_str = str(numero)
    for archivo in os.listdir(carpeta):
        nombre, ext = os.path.splitext(archivo)
        if nombre == numero_str and ext.lower()[1:] in extensiones:
            ruta_completa = os.path.join(carpeta, archivo)
            print(f"✓ Imagen encontrada: {ruta_completa}")
            return ruta_completa

    print(f"⚠ No se encontró imagen con número {numero} en {carpeta}")
    return None

def agregar_imagen(doc, ruta_imagen, titulo, ancho=Cm(12)):
    if not ruta_imagen or not os.path.exists(ruta_imagen):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[IMAGEN NO ENCONTRADA: {titulo}]")
        run.font.italic = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        return

    section = doc.sections[-1]
    max_width = Emu(section.page_width - section.left_margin - section.right_margin)
    if ancho > max_width:
        ancho = max_width

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run()

    try:
        run.add_picture(ruta_imagen, width=ancho)
    except Exception as e:
        print(f"❌ Error al agregar imagen {ruta_imagen}: {e}")
        return

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_before = Pt(6)
    p_titulo.paragraph_format.space_after = Pt(24)
    run_titulo = p_titulo.add_run(titulo)
    run_titulo.font.name = 'Times New Roman'
    run_titulo.font.size = Pt(12)
    run_titulo.font.bold = True

# ================================================================
# TABLAS Y CELDAS (SOLUCIÓN DEFINITIVA POR XML)
# ================================================================
def set_cell_shading(cell, fill='4472C4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)

def set_table_fixed_layout(tabla):
    tbl = tabla._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcW')):
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    twips = int(width_cm * 567)
    tcW.set(qn('w:w'), str(twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_table_grid_widths_xml(tabla, lista_anchos_cm):
    tbl = tabla._tbl
    tblGrid = tbl.tblGrid

    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.append(tblGrid)
    else:
        for existing in tblGrid.findall(qn('w:gridCol')):
            tblGrid.remove(existing)

    for ancho_cm in lista_anchos_cm:
        gridCol = OxmlElement('w:gridCol')
        twips = int(ancho_cm * 567)
        gridCol.set(qn('w:w'), str(twips))
        tblGrid.append(gridCol)

def aplicar_formato_tabla_xml(tabla, lista_anchos_cm):
    set_table_fixed_layout(tabla)
    set_table_grid_widths_xml(tabla, lista_anchos_cm)

    for fila in tabla.rows:
        for col_idx, ancho in enumerate(lista_anchos_cm):
            set_cell_width(fila.cells[col_idx], ancho)

def _celda(cell, texto, negrita=False, centrado=False, tamaño=Pt(10), color_texto='000000'):
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centrado else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = 1.0

    run = p.add_run(texto)
    run.font.name  = 'Times New Roman'
    run.font.size  = tamaño
    run.font.bold  = negrita
    run.font.color.rgb = RGBColor(
        int(color_texto[0:2], 16), int(color_texto[2:4], 16), int(color_texto[4:6], 16)
    )

def agregar_tabla_planificacion(doc, datos):
    ENCABEZADOS = ['Objetivo Específico', 'Actividades', 'Recursos', 'Indicadores de Logro']
    ANCHOS_CM   = [4.5, 4.5, 3.0, 3.5]
    AZUL, GRIS  = '4472C4', 'D9E1F2'

    tabla = doc.add_table(rows=1 + len(datos), cols=4)
    tabla.style = 'Table Grid'
    aplicar_formato_tabla_xml(tabla, ANCHOS_CM)

    for col, enc in enumerate(ENCABEZADOS):
        cell = tabla.cell(0, col)
        set_cell_shading(cell, AZUL)
        _celda(cell, enc, negrita=True, centrado=True, tamaño=Pt(10), color_texto='FFFFFF')

    for fila, (obj, act, rec, ind) in enumerate(datos, start=1):
        contenidos = [obj, act, rec, ind]
        fondo = GRIS if fila % 2 == 0 else 'FFFFFF'
        for col, texto in enumerate(contenidos):
            cell = tabla.cell(fila, col)
            set_cell_shading(cell, fondo)
            _celda(cell, texto, tamaño=Pt(10))

def agregar_gantt(doc, semanas):
    num_sem = max(len(s[1]) for s in semanas)
    AZUL, VERDE, GRIS = '4472C4', '70AD47', 'D9E1F2'

    COL_ACT_CM = 9.5
    section = doc.sections[-1]
    ancho_util_cm = Emu(section.page_width - section.left_margin - section.right_margin).cm
    espacio_restante = max(ancho_util_cm - COL_ACT_CM, 1.0)
    COL_SEM_CM = espacio_restante / num_sem if num_sem > 0 else 1.0

    tabla = doc.add_table(rows=1 + len(semanas), cols=1 + num_sem)
    tabla.style = 'Table Grid'

    anchos_gantt = [COL_ACT_CM] + [COL_SEM_CM] * num_sem
    aplicar_formato_tabla_xml(tabla, anchos_gantt)

    cell_h = tabla.cell(0, 0)
    set_cell_shading(cell_h, AZUL)
    _celda(cell_h, 'Actividades Administrativas', negrita=True, centrado=True, tamaño=Pt(10), color_texto='FFFFFF')

    for s in range(num_sem):
        cell_s = tabla.cell(0, s + 1)
        set_cell_shading(cell_s, AZUL)
        _celda(cell_s, f'S{s+1}', negrita=True, centrado=True, tamaño=Pt(9), color_texto='FFFFFF')

    for fila, (desc, activas) in enumerate(semanas, start=1):
        fondo_fila = GRIS if fila % 2 == 0 else 'FFFFFF'
        cell_a = tabla.cell(fila, 0)
        set_cell_shading(cell_a, fondo_fila)
        _celda(cell_a, desc, tamaño=Pt(10))

        for s in range(num_sem):
            cell_s = tabla.cell(fila, s + 1)
            activa = activas[s] if s < len(activas) else False
            set_cell_shading(cell_s, VERDE if activa else fondo_fila)
            _celda(cell_s, '✓' if activa else '', centrado=True, tamaño=Pt(9), color_texto='FFFFFF' if activa else '000000')

def agregar_numeracion_pie(doc):
    # SOLUCIÓN PORTADA: Iteramos a partir de la sección index 1 para dejar la portada limpia
    for section in doc.sections[1:]:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.extend([fldChar1, instrText, fldChar2, fldChar3])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# ================================================================
# CONSTRUCCIÓN DEL DOCUMENTO (PORTADA COMPATIBLE LIBREOFFICE)
# ================================================================
def construir_portada(doc):
    # ================================================================
    # BLOQUE 1: MEMBRETE (Alineado estrictamente al margen superior)
    # ================================================================
    p_memb = doc.add_paragraph()
    p_memb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_memb.paragraph_format.line_spacing = 1.15
    p_memb.paragraph_format.space_before = Pt(0)
    p_memb.paragraph_format.space_after = Pt(0)

    for i, linea in enumerate(c.MEMBRETE):
        r = p_memb.add_run(linea)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.bold = True
        if i < len(c.MEMBRETE) - 1:
            r.add_break()

    # ================================================================
    # BLOQUE 2: TÍTULO DEL PROYECTO (Exactamente al centro de la página)
    # ================================================================
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.line_spacing = 1.15
    # 150pt de separación empujan el título de forma limpia hacia la mitad de la hoja
    p_titulo.paragraph_format.space_before = Pt(150)
    p_titulo.paragraph_format.space_after = Pt(0)

    r_title = p_titulo.add_run(c.TITULO_PROYECTO)
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(12)
    r_title.font.bold = True

    # ================================================================
    # BLOQUE 3: DATOS DEL AUTOR (Ubicados en el tercio inferior derecho)
    # ================================================================
    p_datos = doc.add_paragraph()
    p_datos.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_datos.paragraph_format.line_spacing = 1.15
    # 160pt bajan el bloque del autor justo debajo del centro
    p_datos.paragraph_format.space_before = Pt(160)
    p_datos.paragraph_format.space_after = Pt(0)

    r_datos = p_datos.add_run(c.AUTOR_DATOS)
    r_datos.font.name = 'Times New Roman'
    r_datos.font.size = Pt(12)
    r_datos.font.bold = True

    # ================================================================
    # BLOQUE 4: CIUDAD Y FECHA (Asentado sobre el margen inferior)
    # ================================================================
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie.paragraph_format.line_spacing = 1.15
    # 75pt cierran el espacio y posicionan la fecha en el borde inferior aceptado
    p_pie.paragraph_format.space_before = Pt(75)
    p_pie.paragraph_format.space_after = Pt(0)

    r_pie = p_pie.add_run(c.FECHA_LUGAR)
    r_pie.font.name = 'Times New Roman'
    r_pie.font.size = Pt(12)

def construir_cuerpo_documento(doc):
    iniciar_capitulo(doc, "I", "REALIDAD ORGANIZACIONAL")
    agregar_titulo_nivel2(doc, "IDENTIFICACIÓN DE LA EMPRESA")

    agregar_titulo_nivel3(doc, "1.1.1 Razón social")
    agregar_parrafo_normado(doc, c.RAZON_SOCIAL, sangria=False)

    agregar_titulo_nivel3(doc, "1.1.2 Reseña histórica")
    for parrafo in c.RESENA_HISTORICA:
        agregar_parrafo_normado(doc, parrafo)

    agregar_titulo_nivel3(doc, "1.1.3 Misión")
    agregar_parrafo_normado(doc, c.MISION, cursiva=True)

    agregar_titulo_nivel3(doc, "1.1.4 Visión")
    agregar_parrafo_normado(doc, c.VISION, cursiva=True)

    agregar_titulo_nivel3(doc, "1.1.5 Valores")
    agregar_parrafo_normado(doc, "Los valores que orientan las actividades de la organización destacan:")
    for i, (valor, descripcion) in enumerate(c.VALORES, 1):
        agregar_item_lista(doc, i, descripcion, valor)

    agregar_titulo_nivel3(doc, "1.1.6 Objetivos Organizacionales")
    agregar_parrafo_normado(doc, "Entre sus objetivos organizacionales se encuentran:")
    for i, objetivo in enumerate(c.OBJETIVOS_ORG, 1):
        agregar_item_lista(doc, i, objetivo)

    agregar_titulo_nivel3(doc, "1.1.7 Ubicación geográfica")
    agregar_parrafo_normado(doc, c.UBICACION, sangria=False)

    ruta_imagen_1 = buscar_imagen_por_numero(c.CARPETA_IMAGENES, 1)
    agregar_imagen(doc, ruta_imagen_1, "Gráfico 1. Representación cartográfica y ubicación espacial de la empresa.", ancho=Cm(5))

    agregar_titulo_nivel3(doc, "1.1.8 Población de los trabajadores de la empresa")
    agregar_parrafo_normado(doc, c.POBLACION)

    agregar_titulo_nivel3(doc, "1.1.9 Estructura Organizativa")
    agregar_parrafo_normado(doc, c.ORGANIGRAMA_TEXTO)

    ruta_imagen_2 = buscar_imagen_por_numero(c.CARPETA_IMAGENES, 2)
    agregar_imagen(doc, ruta_imagen_2, "Gráfico 2. Organigrama estructural y niveles jerárquicos de la organización.", ancho=Cm(12))

    iniciar_capitulo(doc, "II", "DIAGNÓSTICO SITUACIONAL")
    agregar_titulo_nivel2(doc, "Identificación de la Situación Problemática")
    for parrafo in c.SITUACION_PROBLEMATICA:
        agregar_parrafo_normado(doc, parrafo)

    agregar_titulo_nivel2(doc, "Objetivo General")
    agregar_parrafo_normado(doc, c.OBJETIVO_GENERAL)

    agregar_titulo_nivel2(doc, "Objetivos Específicos")
    for i, obj in enumerate(c.OBJETIVOS_ESPECIFICOS, 1):
        agregar_item_lista(doc, i, obj)

    agregar_titulo_nivel2(doc, "Planificación integral de objetivos")
    agregar_parrafo_normado(doc, "La planificación establece la relación entre cada objetivo y las actividades administrativas a ejecutar:")
    agregar_tabla_planificacion(doc, c.PLANIFICACION_DATOS)
    doc.add_paragraph()

    agregar_titulo_nivel2(doc, "Cronograma de actividades")
    agregar_parrafo_normado(doc, "El cronograma estructura temporalmente las tareas administrativas garantizando el cumplimiento del manual documental propuesto:")
    agregar_gantt(doc, c.CRONOGRAMA_DATOS)
    doc.add_paragraph()

# ================================================================
# EJECUCIÓN PRINCIPAL
# ================================================================
def generar_reporte_completo():
    print("» Inicializando documento...")
    doc = setup_iutecp_document()
    construir_portada(doc)
    construir_cuerpo_documento(doc)
    agregar_numeracion_pie(doc)

    docx_output = "Informe_Pasantia_Administracion_Varyna.docx"
    doc.save(docx_output)
    print(f"✔ Archivo Word generado: {docx_output}")

    print("» Renderizando PDF usando LibreOffice...")
    soffice_cmd = 'libreoffice'
    if platform.system() == 'Windows':
        possible_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for path in possible_paths:
            if os.path.exists(path):
                soffice_cmd = path
                break

    try:
        subprocess.run([soffice_cmd, '--headless', '--convert-to', 'pdf', docx_output], check=True, stdout=subprocess.DEVNULL)
        print("✔ ¡PDF generado con éxito!")
    except FileNotFoundError:
        print("❌ Error: LibreOffice no está instalado o no está en el PATH.")
    except subprocess.CalledProcessError as e:
        print(f"❌ Error en la conversión a PDF: {e}")

if __name__ == "__main__":
    generar_reporte_completo()
