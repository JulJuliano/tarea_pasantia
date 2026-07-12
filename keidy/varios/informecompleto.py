import os
import subprocess
import platform
from docx import Document
from docx.shared import Cm, Pt, Inches, Emu  # <-- Se añadió 'Emu'
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ================================================================
#  FUNCIONES AUXILIARES DE FORMATO (NORMATIVA IUTECP 2025)
# ================================================================

def setup_iutecp_document():
    """Configura el documento base: Tamaño Carta, Márgenes Art. 6, Fuente Art. 5"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)

    return doc

def agregar_parrafo_normado(doc, texto, cursiva=False, sangria=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.first_line_indent = Cm(1.25) if sangria else Cm(0)
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)

    run = p.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = cursiva
    return p

def agregar_item_lista(doc, numero, texto, negrita_inicio=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format = p.paragraph_format
    p_format.line_spacing = 1.5
    p_format.first_line_indent = Cm(1.25)

    run_num = p.add_run(f"{numero}. ")
    run_num.font.name = 'Times New Roman'
    run_num.font.size = Pt(12)

    if negrita_inicio:
        run_bold = p.add_run(f"{negrita_inicio}: ")
        run_bold.font.name = 'Times New Roman'
        run_bold.font.size = Pt(12)
        run_bold.font.bold = True

    run_text = p.add_run(texto)
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(12)
    return p

def agregar_titulo_nivel2(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    p.paragraph_format.line_spacing = 1.5

    run = p.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    return p

def agregar_titulo_nivel3(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.25)

    run = p.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    return p

def iniciar_capitulo(doc, numero_romano, titulo):
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec.top_margin = Cm(5)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(4)
    sec.right_margin = Cm(3)
    sec.different_first_page_header_footer = True

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(24)
    r1 = p1.add_run(f"CAPÍTULO {numero_romano}")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.font.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(24)
    r2 = p2.add_run(titulo.upper())
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2.font.bold = True

    sec2 = doc.add_section(WD_SECTION_START.CONTINUOUS)
    sec2.top_margin = Cm(3)
    sec2.bottom_margin = Cm(3)
    sec2.left_margin = Cm(4)
    sec2.right_margin = Cm(3)

def iniciar_seccion_preliminar(doc, titulo):
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec.top_margin = Cm(5)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(4)
    sec.right_margin = Cm(3)
    sec.different_first_page_header_footer = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(titulo.upper())
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.bold = True

    sec2 = doc.add_section(WD_SECTION_START.CONTINUOUS)
    sec2.top_margin = Cm(3)
    sec2.bottom_margin = Cm(3)
    sec2.left_margin = Cm(4)
    sec2.right_margin = Cm(3)

def agregar_cita_larga(doc, texto, cita):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.right_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run = p.add_run(f"{texto} {cita}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def agregar_referencia(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(24)

    run = p.add_run(texto)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        val = kwargs.get(side, {'val': 'single', 'sz': '4', 'color': '000000'})
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   val.get('val', 'single'))
        el.set(qn('w:sz'),    val.get('sz', '4'))
        el.set(qn('w:color'), val.get('color', '000000'))
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_cell_shading(cell, fill='4472C4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    twips = int(width_cm * 567)
    tcW.set(qn('w:w'),    str(twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def _celda(cell, texto, negrita=False, centrado=False,
           tamaño=Pt(10), color_texto='000000'):
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if centrado
                   else WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(texto)
    run.font.name  = 'Times New Roman'
    run.font.size  = tamaño
    run.font.bold  = negrita
    run.font.color.rgb = __import__('docx.shared', fromlist=['RGBColor']).RGBColor(
        int(color_texto[0:2], 16),
        int(color_texto[2:4], 16),
        int(color_texto[4:6], 16)
    )

def agregar_tabla_planificacion(doc, datos):
    ENCABEZADOS = ['Objetivo Específico', 'Actividades', 'Recursos', 'Indicadores de Logro']
    ANCHOS_CM   = [4.5, 4.5, 3.0, 3.5]
    AZUL        = '4472C4'
    GRIS        = 'D9E1F2'

    tabla = doc.add_table(rows=1 + len(datos), cols=4)
    tabla.style = 'Table Grid'

    for col, (enc, ancho) in enumerate(zip(ENCABEZADOS, ANCHOS_CM)):
        cell = tabla.cell(0, col)
        set_cell_shading(cell, AZUL)
        set_cell_width(cell, ancho)
        _celda(cell, enc, negrita=True, centrado=True,
               tamaño=Pt(10), color_texto='FFFFFF')

    for fila, (obj, act, rec, ind) in enumerate(datos, start=1):
        contenidos = [obj, act, rec, ind]
        fondo = GRIS if fila % 2 == 0 else 'FFFFFF'
        for col, texto in enumerate(contenidos):
            cell = tabla.cell(fila, col)
            set_cell_shading(cell, fondo)
            set_cell_width(cell, ANCHOS_CM[col])
            _celda(cell, texto, tamaño=Pt(10))

def agregar_gantt(doc, semanas):
    num_sem    = max(len(s[1]) for s in semanas)
    AZUL       = '4472C4'
    VERDE      = '70AD47'
    GRIS       = 'D9E1F2'

    # --- CORRECCIÓN DEL ERROR DE ATRIBUTO .cm ---
    section = doc.sections[-1]
    # La resta devuelve EMUs (int). Lo pasamos a Emu() para poder leer .cm
    ancho_util_emu = section.page_width - section.left_margin - section.right_margin
    ancho_util_cm = Emu(ancho_util_emu).cm

    COL_ACT_CM = 8.5
    espacio_restante = max(ancho_util_cm - COL_ACT_CM, 1.0)
    COL_SEM_CM = espacio_restante / num_sem if num_sem > 0 else 1.0
    # ---------------------------------------------

    tabla = doc.add_table(rows=1 + len(semanas), cols=1 + num_sem)
    tabla.style = 'Table Grid'

    cell_h = tabla.cell(0, 0)
    set_cell_shading(cell_h, AZUL)
    set_cell_width(cell_h, COL_ACT_CM)
    _celda(cell_h, 'Actividad', negrita=True, centrado=True,
           tamaño=Pt(10), color_texto='FFFFFF')

    for s in range(num_sem):
        cell_s = tabla.cell(0, s + 1)
        set_cell_shading(cell_s, AZUL)
        set_cell_width(cell_s, COL_SEM_CM)
        _celda(cell_s, f'S{s+1}', negrita=True, centrado=True,
               tamaño=Pt(9), color_texto='FFFFFF')

    for fila, (desc, activas) in enumerate(semanas, start=1):
        fondo_fila = GRIS if fila % 2 == 0 else 'FFFFFF'

        cell_a = tabla.cell(fila, 0)
        set_cell_shading(cell_a, fondo_fila)
        set_cell_width(cell_a, COL_ACT_CM)
        _celda(cell_a, desc, tamaño=Pt(10))

        for s in range(num_sem):
            cell_s = tabla.cell(fila, s + 1)
            activa = activas[s] if s < len(activas) else False
            set_cell_shading(cell_s, VERDE if activa else fondo_fila)
            set_cell_width(cell_s, COL_SEM_CM)
            _celda(cell_s, '✓' if activa else '', centrado=True,
                   tamaño=Pt(9), color_texto='FFFFFF' if activa else '000000')

def agregar_numeracion_pie(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# ================================================================
#  FUNCIONES MODULARES PARA IMÁGENES
# ================================================================

def buscar_imagen_por_numero(carpeta, numero, extensiones=None):
    if extensiones is None:
        extensiones = ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']

    if not os.path.exists(carpeta):
        print(f"⚠ Advertencia: La carpeta '{carpeta}' no existe")
        return None

    numero_str = str(numero)

    for archivo in os.listdir(carpeta):
        nombre, ext = os.path.splitext(archivo)
        if nombre == numero_str and ext.lower()[1:] in extensiones:
            ruta_completa = os.path.join(carpeta, archivo)
            print(f"✓ Imagen encontrada: {ruta_completa}")
            return ruta_completa

    print(f"⚠ No se encontró imagen con número {numero} en {carpeta}")
    return None

def agregar_imagen(doc, ruta_imagen, titulo, ancho=Cm(15)):
    if not ruta_imagen or not os.path.exists(ruta_imagen):
        print(f"⚠ No se puede agregar imagen: {ruta_imagen}")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(f"[IMAGEN NO ENCONTRADA: {titulo}]")
        run.font.italic = True
        return

    section = doc.sections[-1]
    max_width = section.page_width - section.left_margin - section.right_margin
    if ancho > max_width:
        print(f"⚠ Ajustando ancho de imagen '{titulo}' para que no exceda los márgenes.")
        ancho = max_width

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(ruta_imagen, width=ancho)
        print(f"✓ Imagen agregada: {os.path.basename(ruta_imagen)}")
    except Exception as e:
        print(f"❌ Error al agregar imagen {ruta_imagen}: {e}")
        return

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_before = Pt(12)
    p_titulo.paragraph_format.space_after = Pt(24)
    run_titulo = p_titulo.add_run(titulo)
    run_titulo.font.name = 'Times New Roman'
    run_titulo.font.size = Pt(12)
    run_titulo.font.bold = True

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_after = Pt(24)

# ================================================================
#  CONSTRUCCIÓN DEL DOCUMENTO
# ================================================================

def construir_portada(doc):
    section = doc.sections[0]
    altura_pagina_pt = section.page_height.pt
    margen_sup_pt = section.top_margin.pt
    margen_inf_pt = section.bottom_margin.pt
    altura_util_pt = altura_pagina_pt - margen_sup_pt - margen_inf_pt

    porcentaje_membrete = 0.12
    porcentaje_espacio1 = 0.18
    porcentaje_titulo = 0.10
    porcentaje_espacio2 = 0.18
    porcentaje_datos = 0.18
    porcentaje_espacio3 = 0.08

    p_memb = doc.add_paragraph()
    p_memb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_memb.paragraph_format.line_spacing = 1.15
    p_memb.paragraph_format.space_after = Pt(0)
    p_memb.paragraph_format.space_before = Pt(altura_util_pt * porcentaje_membrete * 0.3)

    lineas_membrete = [
        "REPÚBLICA BOLIVARIANA DE VENEZUELA",
        "MINISTERIO DEL PODER POPULAR PARA LA EDUCACIÓN UNIVERSITARIA",
        "INSTITUTO UNIVERSITARIO DE TECNOLOGÍA",
        "\"ELÍAS CALIXTO POMPA\" (IUTECP)",
        "EL TIGRE, ESTADO ANZOÁTEGUI"
    ]

    for i, linea in enumerate(lineas_membrete):
        r = p_memb.add_run(linea)
        r.font.bold = True
        r.font.size = Pt(12)
        if i < len(lineas_membrete) - 1:
            r.add_break()

    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_before = Pt(altura_util_pt * porcentaje_espacio1)
    p_titulo.paragraph_format.space_after = Pt(0)

    r_title = p_titulo.add_run("DISEÑO DE UN SISTEMA PARA EL CONTROL, TRAZABILIDAD Y REPORTE DE MOVIMIENTOS DOCUMENTALES EN LUBRICANTES Y EQUIPOS VARYNA, C.A.")
    r_title.font.bold = True
    r_title.font.size = Pt(12)

    p_datos = doc.add_paragraph()
    p_datos.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_datos.paragraph_format.space_before = Pt(altura_util_pt * porcentaje_espacio2)
    p_datos.paragraph_format.space_after = Pt(0)

    r_datos = p_datos.add_run("Autor: Guzmán, Keidy\nC.I.: 28.706.352")
    r_datos.font.bold = True

    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pie.paragraph_format.space_before = Pt(altura_util_pt * porcentaje_espacio3)
    p_pie.paragraph_format.space_after = Pt(0)

    r_pie = p_pie.add_run("El Tigre, julio de 2026")
    r_pie.font.size = Pt(12)

def construir_cuerpo_documento(doc, carpeta_imagenes="imagenes"):
    iniciar_capitulo(doc, "I", "REALIDAD ORGANIZACIONAL")

    agregar_titulo_nivel2(doc, "IDENTIFICACIÓN DE LA EMPRESA")

    agregar_titulo_nivel3(doc, "1.1.1 Razón social")
    agregar_parrafo_normado(doc, "Lubricantes y Equipos Varyna, C.A.", sangria=False)

    agregar_titulo_nivel3(doc, "1.1.2 Reseña histórica")
    agregar_parrafo_normado(doc, "Lubricantes y Equipos Varyná C.A. es una empresa venezolana con más de treinta y seis (36) años de trayectoria en el sector petrolero, industrial y de construcción, consolidándose como una organización de amplia experiencia y reconocimiento dentro del mercado nacional. Desde sus inicios, la empresa ha orientando sus esfuerzos al desarrollo de soluciones integrales que contribuyan al fortalecimiento de las actividades productivas del país, especialmente en áreas estratégicas vinculadas con la industria energética.")
    agregar_parrafo_normado(doc, "Como empresa integrante y pilar fundamental del Grupo Corporativo VTC, Lubricantes y Equipos Varyná C.A. ha mantenido un crecimiento sostenido basado en la innovación, la calidad de sus servicios y el compromiso con sus clientes. Su actividad principal se enfoca en el procesamiento y suministro de productos químicos especializados, el tratamiento de crudo y la provisión de maquinaria pesada para operaciones industriales y petroleras.")
    agregar_parrafo_normado(doc, "A lo largo de su trayectoria, la organización ha participado activamente en proyectos destinados al crecimiento, recuperación y mantenimiento de la producción de pozos petroleros, aportando recursos técnicos y operativos que contribuyen al desarrollo del sector energético venezolano. Asimismo, ha fortalecido su presencia en la región oriental del país, donde mantiene una sólida base de operaciones y una participación constante en ferias y eventos empresariales de carácter regional y nacional.")
    agregar_parrafo_normado(doc, "En la actualidad, Lubricantes y Equipos Varyná C.A. continúa posicionándose como una empresa comprometida con la excelencia, la eficiencia operativa y el desarrollo sostenible, ofreciendo productos y servicios que responden a las necesidades de los sectores petrolero, industrial y de construcción, contribuyendo de manera significativa al progreso económico de Venezuela.")

    agregar_titulo_nivel3(doc, "1.1.3 Misión")
    agregar_parrafo_normado(doc, "Brindar a nuestros clientes objetivos, soluciones de calidad en las áreas en las cuales nos desempeñamos, para contribuir de manera significativa en sus resultados. Aportando valor con nuestras respuestas a sus requerimientos. En la búsqueda de un mejor país y de una mejor humanidad.", cursiva=True)

    agregar_titulo_nivel3(doc, "1.1.4 Visión")
    agregar_parrafo_normado(doc, "Ser el conglomerado de empresas líderes en cada una de las categorías en las que participamos, generando modelos de negocios altamente competitivos, atendiendo a nuestros distintos beneficiarios con productos y servicios de calidad, a la vez que contribuimos con el desarrollo social y ambiental del planeta. Buscando cubrir y trascender los mercados locales, teniendo presencia importante en el mercado global.", cursiva=True)

    agregar_titulo_nivel3(doc, "1.1.5 Valores")
    agregar_parrafo_normado(doc, "Los valores que orientan las actividades de Lubricantes y Equipos Varyná C.A. están enfocados en el compromiso con la calidad, la excelencia y la satisfacción de sus clientes. Entre los principales valores de la organización destacan:")

    valores = [
        ("Responsabilidad", "Cumplimiento eficiente de los compromisos adquiridos con clientes, proveedores y trabajadores."),
        ("Integridad", "Actuar con honestidad, ética y transparencia en todas las operaciones de la empresa."),
        ("Calidad", "Garantizar productos y servicios que satisfagan las necesidades y expectativas de los clientes."),
        ("Compromiso", "Mantener una actitud de dedicación y esfuerzo para alcanzar los objetivos organizacionales."),
        ("Trabajo en equipo", "Fomentar la cooperación y el respeto entre los trabajadores para lograr resultados exitosos."),
        ("Innovación", "Promover la mejora continua de los procesos, productos y servicios ofrecidos."),
        ("Seguridad", "Velar por la protección de los trabajadores, las instalaciones y el medio ambiente en cada actividad realizada.")
    ]

    for i, (valor, descripcion) in enumerate(valores, 1):
        agregar_item_lista(doc, i, descripcion, valor)

    agregar_titulo_nivel3(doc, "1.1.6 Objetivos Organizacionales")
    agregar_parrafo_normado(doc, "Lubricantes y Equipos Varyná C.A. tiene como objetivo principal contribuir al desarrollo de los sectores petrolero, industrial y de construcción mediante la prestación de servicios especializados y el suministro de productos de alta calidad. Entre sus objetivos organizacionales se encuentran:")

    objetivos = [
        "Proporcionar soluciones eficientes y oportunas a las necesidades de sus clientes.",
        "Mantener altos estándares de calidad en los productos y servicios ofrecidos.",
        "Impulsar el crecimiento y fortalecimiento de la industria petrolera nacional.",
        "Garantizar la satisfacción de los clientes mediante la mejora continua de los procesos.",
        "Promover el desarrollo profesional y personal de sus trabajadores.",
        "Contribuir al desarrollo económico, social y ambiental del país mediante una gestión responsable y sostenible."
    ]

    for i, objetivo in enumerate(objetivos, 1):
        agregar_item_lista(doc, i, objetivo)

    agregar_titulo_nivel3(doc, "1.1.7 Ubicación geográfica")
    agregar_parrafo_normado(doc, "Calle 23 de enero entre calle principal el palomar y calle la paz sector vista al sol, San José de Guanipa Edo. Anzoátegui.", sangria=False)

    # --- CAMBIO: Imagen 1 reducida a 5 cm ---
    ruta_imagen_1 = buscar_imagen_por_numero(carpeta_imagenes, 1)
    agregar_imagen(doc, ruta_imagen_1, "Gráfico 1. Representación cartográfica y ubicación espacial de la empresa.", ancho=Cm(5))
    # ----------------------------------------

    agregar_titulo_nivel3(doc, "1.1.8 Población de los trabajadores de la empresa")
    agregar_parrafo_normado(doc, "La población de trabajadores de Lubricantes y Equipos Varyná C.A. está conformada por personal administrativo, técnico y operativo, quienes desempeñan funciones esenciales para el cumplimiento de las actividades de la organización. El talento humano constituye uno de los recursos más importantes de la empresa, ya que contribuye al desarrollo eficiente de los procesos relacionados con el sector petrolero, industrial y de construcción.")
    agregar_parrafo_normado(doc, "La estructura laboral se encuentra distribuida en diferentes áreas funcionales, entre las que destacan administración, operaciones, logística, mantenimiento, seguridad industrial y atención al cliente. Todo el personal trabaja de manera coordinada para garantizar la calidad de los servicios prestados y el logro de los objetivos organizacionales.")

    agregar_titulo_nivel3(doc, "1.1.9 Estructura organizacional de la empresa")
    ruta_imagen_2 = buscar_imagen_por_numero(carpeta_imagenes, 2)
    agregar_imagen(doc, ruta_imagen_2, "Gráfico 2. Organigrama estructural y niveles jerárquicos de la organización.")


    iniciar_capitulo(doc, "II", "DIAGNÓSTICO SITUACIONAL")

    agregar_titulo_nivel2(doc, "Identificación de la Situación Problemática")
    agregar_parrafo_normado(doc, "En el área de presidencia de Lubricantes y Equipos Varyna, C.A., se ha detectado una problemática significativa relacionada con la gestión y el flujo de los movimientos documentales. En la actualidad, el control de oficios, memorandos, contratos, facturas y correspondencia en general se realiza de manera empírica, apoyándose en registros manuales, hojas de cálculo no integradas y archivos físicos dispersos. Esta carencia tecnológica genera retrasos considerables en la búsqueda y recuperación de información, duplicidad de esfuerzos, pérdida de documentos críticos y una alta vulnerabilidad ante auditorías internas.")
    agregar_parrafo_normado(doc, "La falta de un mecanismo formal para la trazabilidad documental impide conocer con exactitud la fecha de recepción, el responsable actual, el estado de trámite y el destino final de cada expediente. Como consecuencia, la toma de decisiones estratégicas por parte de la alta gerencia se ve entorpecida, afectando la eficiencia operativa, la transparencia administrativa y la capacidad de respuesta ante clientes y entes gubernamentales. Por tanto, resulta imperante el diseño de un sistema que centralice y automatice estos procesos.")

    agregar_titulo_nivel2(doc, "Objetivo General")
    agregar_parrafo_normado(doc, "Diseñar un sistema para el control, trazabilidad y reporte de movimientos documentales en la presidencia de Lubricantes y Equipos Varyna, C.A., con el propósito de optimizar la gestión de la información, garantizar la seguridad de los archivos y mejorar la eficiencia de los procesos administrativos de la organización.")

    agregar_titulo_nivel2(doc, "Objetivos Específicos")
    objetivos_esp = [
        "Diagnosticar la situación actual del manejo de los movimientos documentales en el área de presidencia de Lubricantes y Equipos Varyna, C.A., para identificar las fallas y cuellos de botella existentes.",
        "Identificar los requerimientos funcionales, técnicos y de seguridad necesarios para la estructuración del nuevo sistema de control documental.",
        "Proponer la arquitectura lógica y los componentes del sistema automatizado que permitan la trazabilidad en tiempo real y la generación de reportes gerenciales."
    ]
    for i, obj in enumerate(objetivos_esp, 1):
        agregar_item_lista(doc, i, obj)

    agregar_titulo_nivel2(doc, "Planificación integral de objetivos")
    agregar_parrafo_normado(doc, "La planificación integral de objetivos establece la relación directa entre cada objetivo específico, las actividades estratégicas a ejecutar, los recursos requeridos y los indicadores de logro que permitirán verificar el cumplimiento de cada etapa del proyecto. A continuación, se presenta la matriz de operacionalización:")

    datos_planificacion = [
        (
            "Diagnosticar la situación actual del manejo de los movimientos documentales en presidencia.",
            "Aplicación de entrevistas al personal administrativo y observación directa de los flujos de trabajo actuales.",
            "Guía de entrevista, libreta de notas, cámara fotográfica.",
            "Informe de diagnóstico con fallas y cuellos de botella identificados."
        ),
        (
            "Identificar los requerimientos funcionales, técnicos y de seguridad del sistema.",
            "Levantamiento de requerimientos mediante técnicas de ingeniería de software y análisis de los formatos documentales vigentes.",
            "Herramientas CASE, plantillas de requerimientos, acceso a documentos internos.",
            "Documento de especificación de requerimientos del sistema (ERS) validado."
        ),
        (
            "Proponer la arquitectura lógica y los componentes del sistema automatizado.",
            "Diseño de diagramas de flujo, modelo entidad-relación de la base de datos y prototipado de interfaces de usuario.",
            "Software de modelado (draw.io / Lucidchart), entorno de desarrollo, PC.",
            "Prototipo funcional del sistema y documento de diseño técnico aprobado."
        ),
    ]
    agregar_tabla_planificacion(doc, datos_planificacion)
    doc.add_paragraph()

    agregar_titulo_nivel2(doc, "Cronograma de actividades")
    agregar_parrafo_normado(doc, "El cronograma de actividades estructura temporalmente las tareas a ejecutar durante el periodo de pasantías, garantizando el cumplimiento de los objetivos planteados en el tiempo establecido. Las celdas marcadas (✓) indican la semana de ejecución de cada actividad:")

    gantt_datos = [
        ("Inducción institucional y reconocimiento del área de presidencia.",                                [True,  False, False, False, False]),
        ("Planteamiento formal de la situación problemática.",                                              [True,  False, False, False, False]),
        ("Recolección de datos mediante observación directa de los flujos documentales.",                   [False, True,  False, False, False]),
        ("Aplicación de entrevistas al personal administrativo.",                                           [False, True,  False, False, False]),
        ("Elaboración del informe de diagnóstico de los procesos actuales.",                                [False, True,  False, False, False]),
        ("Análisis de requerimientos funcionales, técnicos y de seguridad.",                                [False, False, True,  False, False]),
        ("Definición de las reglas de negocio para la trazabilidad documental.",                            [False, False, True,  False, False]),
        ("Diseño de la arquitectura del sistema y modelado de la base de datos.",                           [False, False, False, True,  False]),
        ("Prototipado de interfaces de usuario del sistema propuesto.",                                     [False, False, False, True,  False]),
        ("Elaboración del manual de procedimientos y presentación de la propuesta a la gerencia.",          [False, False, False, False, True ]),
    ]
    agregar_gantt(doc, gantt_datos)
    doc.add_paragraph()


# ================================================================
#  EJECUCIÓN PRINCIPAL
# ================================================================

def generar_reporte_completo(carpeta_imagenes="imagenes"):
    print("» Inicializando documento bajo Normas IUTECP 2025...")
    doc = setup_iutecp_document()

    construir_portada(doc)
    construir_cuerpo_documento(doc, carpeta_imagenes=carpeta_imagenes)
    agregar_numeracion_pie(doc)

    docx_output = "Informe_Pasantia_Varyna_IUTECP.docx"
    doc.save(docx_output)
    print(f"✔ Archivo Word generado: {docx_output}")

    print("» Renderizando PDF usando LibreOffice...")

    soffice_cmd = 'libreoffice'
    if platform.system() == 'Windows':
        possible_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for path in possible_paths:
            if os.path.exists(path):
                soffice_cmd = path
                break
        else:
            soffice_cmd = 'soffice'

    try:
        subprocess.run([soffice_cmd, '--headless', '--convert-to', 'pdf', docx_output], check=True, stdout=subprocess.DEVNULL)
        print("✔ ¡PDF académico generado con éxito!")
    except FileNotFoundError:
        print("❌ Error: LibreOffice no está instalado o no se encontró la ruta. Abre el Word y guárdalo como PDF manualmente.")
    except subprocess.CalledProcessError as e:
        print(f"❌ Error en la conversión a PDF: {e}")

if __name__ == "__main__":
    CARPETA_IMAGENES = "imagenes"
    generar_reporte_completo(carpeta_imagenes=CARPETA_IMAGENES)
