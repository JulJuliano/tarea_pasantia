import os
import subprocess
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ================================================================
#  FUNCIONES DE CONFIGURACIÓN Y FORMATO (ESTÁTICAS - NORMA IUTECP)
# ================================================================

def setup_page(doc):
    """Márgenes Art. 6 (4cm Izquierdo para encuadernación, 3cm los demás)"""
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)

def add_paragraph_normado(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)

    run_l = p.add_run(label)
    run_l.font.name = 'Times New Roman'
    run_l.font.size = Pt(12)
    run_l.font.bold = True

    run_v = p.add_run(value)
    run_v.font.name = 'Times New Roman'
    run_v.font.size = Pt(12)
    return p

def set_cell_format(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        borders.append(border)
    tblPr.append(borders)

# ================================================================
#  PROCESADOR CORE: COMPILA UNA SEMANA ESPECÍFICA
# ================================================================

def generar_documento_semana(datos):
    doc = Document()
    setup_page(doc)

    # 1. Título Centralizado
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_after = Pt(12)
    run_t = p_titulo.add_run("PLAN SEMANAL DE PASANTÍAS")
    run_t.font.name = 'Times New Roman'
    run_t.font.size = Pt(14)
    run_t.font.bold = True

    # 2. Datos Informativos de la Empresa y la Pasante
    add_label_value(doc, "EMPRESA: ", "Lubricantes y Equipos Varyna, C.A.")
    add_label_value(doc, "DEPARTAMENTO: ", "Departamento Administrativo / Área de Procura")
    add_label_value(doc, "ESPECIALIDAD: ", "Administración / Ciencias Comerciales")
    add_label_value(doc, "SEMANA N°: ", datos["num_semana"])
    add_label_value(doc, "PERÍODO: ", datos["periodo"])
    add_label_value(doc, "PASANTE: ", "Keidy Guzmán |  C.I.: 28.706.352")
    add_label_value(doc, "TUTOR INDUSTRIAL: ", "Martina Rondón | C.I.: 12.208.768")

    doc.add_paragraph()

    # 3. Objetivo de la semana
    p_obj_t = doc.add_paragraph()
    p_obj_t.paragraph_format.space_after = Pt(4)
    run_ot = p_obj_t.add_run("OBJETIVO DE LA SEMANA:")
    run_ot.font.name = 'Times New Roman'
    run_ot.font.size = Pt(12)
    run_ot.font.bold = True

    add_paragraph_normado(doc, datos["objetivo"])

    doc.add_paragraph()

    # 4. Tabla de Actividades Planificadas (Dimensiones fijas para evitar desbordes)
    p_cuarto_t = doc.add_paragraph()
    p_cuarto_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cuarto_t.paragraph_format.space_after = Pt(6)
    run_ct = p_cuarto_t.add_run(f"CUADRO 1. Actividades Planificadas de la Semana {datos['num_semana']}")
    run_ct.font.name = 'Times New Roman'
    run_ct.font.size = Pt(10)
    run_ct.font.bold = True

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False

    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    table._tbl.tblPr.append(tblLayout)

    # Encabezados con Sombreado Azul Pastel
    headers = ["Nro", "Actividad", "Descripción", "Recursos"]
    for i, h in enumerate(headers):
        set_cell_format(table.rows[0].cells[i], h, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), "D9E2F3")
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_elm)

    # Inyección de actividades semanales
    for nro, act, desc, rec in datos["actividades"]:
        row = table.add_row()
        set_cell_format(row.cells[0], nro, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_format(row.cells[1], act, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_format(row.cells[2], desc, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_cell_format(row.cells[3], rec, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Ancho total útil: 14.59 cm
    widths = [Cm(1.2), Cm(3.5), Cm(6.5), Cm(3.39)]
    for idx, width in enumerate(widths):
        table.columns[idx].width = width
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.add_paragraph()

    # 5. Entregables
    p_ent_t = doc.add_paragraph()
    p_ent_t.paragraph_format.space_after = Pt(4)
    run_et = p_ent_t.add_run("ENTREGABLES DE LA SEMANA:")
    run_et.font.name = 'Times New Roman'
    run_et.font.size = Pt(12)
    run_et.font.bold = True

    add_paragraph_normado(doc, datos["entregables"])

    doc.add_paragraph()

    # 6. Observaciones del Tutor Industrial
    p_obs_t = doc.add_paragraph()
    p_obs_t.paragraph_format.space_after = Pt(4)
    run_ot2 = p_obs_t.add_run("OBSERVACIONES DEL TUTOR INDUSTRIAL:")
    run_ot2.font.name = 'Times New Roman'
    run_ot2.font.size = Pt(12)
    run_ot2.font.bold = True

    obs_texto = datos.get("observaciones", "")
    if obs_texto:
        add_paragraph_normado(doc, obs_texto)
    else:
        p_obs = doc.add_paragraph()
        p_obs.paragraph_format.line_spacing = 1.5
        run_obs = p_obs.add_run("(Espacio reservado para la evaluación cualitativa del tutor sobre control de procesos, puntualidad, manejo de solicitudes y cumplimiento de directrices contables)")
        run_obs.font.name = 'Times New Roman'
        run_obs.font.size = Pt(11)
        run_obs.font.italic = True

        for _ in range(4):
            p_line = doc.add_paragraph()
            p_line.paragraph_format.space_before = Pt(8)
            p_line.paragraph_format.space_after = Pt(0)
            run_line = p_line.add_run("_" * 75)
            run_line.font.name = 'Times New Roman'
            run_line.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # 7. Firmas (Estructura de Tabla Invisible)
    sign_table = doc.add_table(rows=2, cols=2)
    remove_table_borders(sign_table)
    sign_table.allow_autofit = False
    for row in sign_table.rows:
        row.cells[0].width = Cm(7.29)
        row.cells[1].width = Cm(7.30)

    set_cell_format(sign_table.rows[0].cells[0], "", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_format(sign_table.rows[0].cells[1], "", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Insertar firma del pasante si existe
    ruta_firma_keidy = os.path.join("imagenes", "firma_keidy.png")
    if os.path.exists(ruta_firma_keidy):
        cell_fk = sign_table.rows[0].cells[0]
        cell_fk.text = ""
        cell_fk.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_fk = cell_fk.paragraphs[0].add_run()
        run_fk.add_picture(ruta_firma_keidy, width=Cm(3.5))
    else:
        set_cell_format(sign_table.rows[0].cells[0], "_________________________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    ruta_firma_tutor = os.path.join("imagenes", "firma_tutor_keidy.png")
    if os.path.exists(ruta_firma_tutor):
        cell_ft = sign_table.rows[0].cells[1]
        cell_ft.text = ""
        cell_ft.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_ft = cell_ft.paragraphs[0].add_run()
        run_ft.add_picture(ruta_firma_tutor, width=Cm(3.5))
    else:
        set_cell_format(sign_table.rows[0].cells[1], "_________________________", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_format(sign_table.rows[1].cells[0], "Firma del Pasante:\nKeidy Guzmán |  C.I.: 28.706.352", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_format(sign_table.rows[1].cells[1], "Firma del Tutor Industrial:\nMartina Rondón | C.I.: 12.208.768", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # 8. Bloque Sello Corporativo
    ruta_sello = os.path.join("imagenes", "sello_keidy.jpg")
    if os.path.exists(ruta_sello):
        p_sello = doc.add_paragraph()
        p_sello.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sello.paragraph_format.space_before = Pt(12)
        run_s = p_sello.add_run()
        run_s.add_picture(ruta_sello, width=Cm(4))
    else:
        p_sello = doc.add_paragraph()
        p_sello.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sello.paragraph_format.space_before = Pt(12)
        run_s = p_sello.add_run("SELLO DE LA EMPRESA")
        run_s.font.name = 'Times New Roman'
        run_s.font.size = Pt(10)
        run_s.font.bold = True

    output_dir = "cronogramas generados"
    os.makedirs(output_dir, exist_ok=True)
    fn_word = os.path.join(output_dir, f"Cronograma_Procura_Semana{datos['num_semana']}_IUTECP.docx")
    doc.save(fn_word)
    print(f"✔ Generado Word: {fn_word}")

    try:
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, fn_word], check=True, stdout=subprocess.DEVNULL)
        print(f"✔ Convertido PDF: {fn_word.replace('.docx', '.pdf')}")
    except Exception:
        print(f"⚠️ Word creado. LibreOffice headless no se ejecutó.")

# ================================================================
#  COLECCIÓN DE DATOS MAESTRA (CRONOGRAMA ADAPTADO A TSU - 10 SEMANAS)
# ================================================================

reportes_pasantias = [{'num_semana': '1',
  'periodo': 'Del 08 al 12 de junio de 2026',
  'objetivo': 'Formalizar el inicio de las pasantías, reconocer el Departamento Administrativo y observar el '
              'flujo inicial de las requisiciones de compra.',
  'entregables': 'Registro de inducción y primer levantamiento del flujo de requisiciones.',
  'observaciones': '',
  'actividades': [('1',
                   'Formalización e inducción',
                   'Consignar los formatos de inicio, recibir la inducción de la empresa y conocer los '
                   'lineamientos del Departamento Administrativo.',
                   'Formatos institucionales, cuaderno de notas'),
                  ('2',
                   'Recorrido por el área',
                   'Reconocer las áreas físicas y los actores que intervienen en el proceso de procura.',
                   'Agenda, cuaderno de apuntes'),
                  ('3',
                   'Presentación ante el equipo',
                   'Integrarse al equipo de trabajo y conocer las responsabilidades generales del área de compras '
                   'y logística.',
                   'Documentos de presentación'),
                  ('4',
                   'Observación de requisiciones',
                   'Observar cómo ingresa y se procesa inicialmente una solicitud de compra.',
                   'Formatos de requisición, cuaderno de campo'),
                  ('5',
                   'Registro de situación inicial',
                   'Anotar las primeras oportunidades de mejora detectadas en recepción, seguimiento y control de '
                   'solicitudes.',
                   'Libreta de notas')]},
 {'num_semana': '2',
  'periodo': 'Del 15 al 19 de junio de 2026',
  'objetivo': 'Apoyar la recepción y registro de solicitudes de compra y levantar el flujo secuencial del proceso '
              'actual de procura.',
  'entregables': 'Registro de solicitudes y esquema preliminar del flujo actual de procura.',
  'observaciones': '',
  'actividades': [('1',
                   'Recepción de solicitudes',
                   'Apoyar la recepción y clasificación de las solicitudes de compra pendientes.',
                   'Formatos de requisición, archiveros'),
                  ('2',
                   'Registro de documentos',
                   'Registrar la correspondencia de compras entrante en el control diario del departamento.',
                   'Libro de control, PC'),
                  ('3',
                   'Reconocimiento del flujo',
                   'Conversar con el personal encargado para precisar el recorrido de cada requisición.',
                   'Cuaderno de apuntes'),
                  ('4',
                   'Seguimiento de cotizaciones',
                   'Revisar el mecanismo utilizado para conocer el estado de las cotizaciones en curso.',
                   'Expedientes y registros del área'),
                  ('5',
                   'Mapeo del proceso',
                   'Documentar las etapas, responsables y puntos de decisión del proceso observado.',
                   'Procesador de textos, papel')]},
 {'num_semana': '3',
  'periodo': 'Del 22 al 26 de junio de 2026',
  'objetivo': 'Tramitar requisiciones y aplicar una entrevista estructurada al personal para identificar '
              'dificultades de seguimiento y aprobación.',
  'entregables': 'Requisiciones tramitadas y guía de entrevista aplicada con respuestas registradas.',
  'observaciones': '',
  'actividades': [('1',
                   'Tramitación de requisiciones',
                   'Colaborar en la gestión de requisiciones de materiales e insumos recibidas durante la semana.',
                   'Formatos de requisición'),
                  ('2',
                   'Comunicación con proveedores',
                   'Apoyar el seguimiento de cotizaciones solicitadas a proveedores.',
                   'Teléfono, correo de oficina'),
                  ('3',
                   'Actualización de cotizaciones',
                   'Registrar el estatus de las cotizaciones activas.',
                   'PC, hojas de seguimiento'),
                  ('4',
                   'Aplicación de entrevista',
                   'Aplicar la guía de entrevista estructurada al personal relacionado con compras y '
                   'administración.',
                   'Guía de entrevista'),
                  ('5',
                   'Registro de hallazgos',
                   'Anotar las demoras, duplicidades y dificultades de trazabilidad señaladas durante la '
                   'entrevista.',
                   'Cuaderno de campo')]},
 {'num_semana': '4',
  'periodo': 'Del 29 de junio al 03 de julio de 2026',
  'objetivo': 'Dar seguimiento a órdenes de compra y revisar expedientes históricos para determinar tiempos de '
              'respuesta y etapas críticas.',
  'entregables': 'Registro actualizado de órdenes abiertas y tabla de tiempos por fase.',
  'observaciones': '',
  'actividades': [('1',
                   'Seguimiento de órdenes abiertas',
                   'Verificar el estado de las órdenes de compra en proceso.',
                   'Expedientes activos, PC'),
                  ('2',
                   'Cotizaciones pendientes',
                   'Identificar cotizaciones con mayor tiempo de espera.',
                   'Hojas de seguimiento'),
                  ('3',
                   'Actualización de proveedores',
                   'Depurar información básica del registro de proveedores.',
                   'PC, base de proveedores'),
                  ('4',
                   'Revisión histórica',
                   'Revisar expedientes de compras anteriores para calcular tiempos aproximados de respuesta.',
                   'Archivo histórico, Excel'),
                  ('5',
                   'Etapas críticas',
                   'Determinar las fases donde se concentran los mayores retrasos.',
                   'Hoja de cálculo, notas')]},
 {'num_semana': '5',
  'periodo': 'Del 06 al 10 de julio de 2026',
  'objetivo': 'Analizar las causas de las demoras del proceso de procura y elaborar el diagrama de Ishikawa.',
  'entregables': 'Diagrama de Ishikawa y jerarquización de causas del proceso.',
  'observaciones': '',
  'actividades': [('1',
                   'Tramitación de solicitudes',
                   'Apoyar la gestión de solicitudes de compra del período.',
                   'Formatos de requisición'),
                  ('2',
                   'Cuadros comparativos',
                   'Colaborar en la comparación de cotizaciones de proveedores.',
                   'Excel'),
                  ('3',
                   'Actualización de proveedores',
                   'Incorporar información necesaria al registro de proveedores.',
                   'PC'),
                  ('4',
                   'Diagrama de Ishikawa',
                   'Organizar las causas detectadas en categorías relacionadas con procedimiento, documentación, '
                   'responsabilidades y seguimiento.',
                   'Herramienta de diagramación'),
                  ('5',
                   'Jerarquización de causas',
                   'Clasificar los factores según su incidencia sobre los retrasos y la trazabilidad.',
                   'Matriz de causas')]},
 {'num_semana': '6',
  'periodo': 'Del 13 al 17 de julio de 2026',
  'objetivo': 'Diseñar el flujo simplificado del proceso de procura, definiendo etapas, responsables y puntos de '
              'control.',
  'entregables': 'Flujo simplificado con responsables y puntos de control definidos.',
  'observaciones': '',
  'actividades': [('1',
                   'Seguimiento a cotizaciones',
                   'Actualizar expedientes de cotizaciones pendientes.',
                   'Hojas de seguimiento, PC'),
                  ('2',
                   'Actualización de expedientes',
                   'Registrar avances de compras en proceso.',
                   'Archiveros, formatos'),
                  ('3',
                   'Diseño del flujo',
                   'Redactar y diagramar el nuevo recorrido simplificado de las requisiciones.',
                   'Procesador de textos, diagramador'),
                  ('4',
                   'Asignación de responsables',
                   'Definir el responsable de cada etapa del procedimiento propuesto.',
                   'Organigrama, procesador de textos'),
                  ('5',
                   'Puntos de control',
                   'Establecer verificaciones mínimas y tiempos de referencia por fase.',
                   'Hoja de trabajo')]},
 {'num_semana': '7',
  'periodo': 'Del 20 al 24 de julio de 2026',
  'objetivo': 'Diseñar los formatos estandarizados, la matriz de autorización por monto y los indicadores básicos '
              'de seguimiento de la procura.',
  'entregables': 'Formatos de solicitud y orden de compra, matriz de autorización e indicadores propuestos.',
  'observaciones': '',
  'actividades': [('1',
                   'Apoyo en requisiciones',
                   'Procesar solicitudes de compra recibidas durante la semana.',
                   'Formatos de requisición'),
                  ('2',
                   'Formato de solicitud de cotización',
                   'Diseñar un formato uniforme para requerir cotizaciones a proveedores.',
                   'PC, procesador de textos'),
                  ('3',
                   'Formato de orden de compra',
                   'Diseñar un formato con los campos esenciales para el control de la adquisición.',
                   'PC, procesador de textos'),
                  ('4',
                   'Matriz de autorización',
                   'Definir niveles de autorización por rangos de monto de compra para revisión de la gerencia.',
                   'Hoja de cálculo'),
                  ('5',
                   'Indicadores de seguimiento',
                   'Proponer indicadores básicos de tiempo, solicitudes pendientes y cumplimiento de plazos.',
                   'Excel')]},
 {'num_semana': '8',
  'periodo': 'Del 27 al 31 de julio de 2026',
  'objetivo': 'Redactar la propuesta de simplificación administrativa integrando el flujo, los formatos y los '
              'mecanismos de control diseñados.',
  'entregables': 'Borrador completo de la propuesta de simplificación administrativa.',
  'observaciones': '',
  'actividades': [('1',
                   'Seguimiento a órdenes',
                   'Actualizar el estatus de órdenes de compra activas.',
                   'Expedientes, PC'),
                  ('2',
                   'Actualización de proveedores',
                   'Revisar y completar registros de proveedores del período.',
                   'Base de proveedores'),
                  ('3',
                   'Redacción de propuesta',
                   'Integrar el flujo simplificado y la descripción del procedimiento.',
                   'Procesador de textos'),
                  ('4',
                   'Integración de formatos',
                   'Incorporar los formatos y la matriz de autorización como componentes de la propuesta.',
                   'Archivos digitales'),
                  ('5',
                   'Indicadores y control',
                   'Documentar la forma propuesta de seguimiento mediante indicadores.',
                   'Hoja de cálculo, procesador de textos')]},
 {'num_semana': '9',
  'periodo': 'Del 03 al 07 de agosto de 2026',
  'objetivo': 'Validar la propuesta con el tutor industrial e incorporar las observaciones recibidas.',
  'entregables': 'Propuesta revisada y ajustada con registro de observaciones.',
  'observaciones': '',
  'actividades': [('1',
                   'Cierre administrativo',
                   'Apoyar la organización de expedientes y órdenes correspondientes al período.',
                   'Archiveros'),
                  ('2',
                   'Archivo de órdenes',
                   'Organizar las órdenes de compra gestionadas durante las pasantías.',
                   'Carpetas físicas'),
                  ('3',
                   'Presentación al tutor',
                   'Presentar la propuesta para revisión del tutor industrial.',
                   'Documento de propuesta'),
                  ('4',
                   'Registro de observaciones',
                   'Documentar las correcciones y sugerencias indicadas.',
                   'Libreta de notas'),
                  ('5',
                   'Ajustes finales',
                   'Incorporar las observaciones pertinentes al documento.',
                   'Procesador de textos')]},
 {'num_semana': '10',
  'periodo': 'Del 10 al 14 de agosto de 2026',
  'objetivo': 'Presentar la propuesta y consolidar el informe académico de pasantías para su entrega '
              'institucional.',
  'entregables': 'Propuesta presentada, recaudos de cierre e informe académico final.',
  'observaciones': '',
  'actividades': [('1',
                   'Presentación de propuesta',
                   'Exponer la propuesta de simplificación ante la gerencia del área.',
                   'Documento final, equipo de apoyo'),
                  ('2',
                   'Recaudos de cierre',
                   'Gestionar los formatos y firmas correspondientes al cierre de las pasantías.',
                   'Formatos institucionales'),
                  ('3',
                   'Consolidación del informe',
                   'Integrar capítulos, referencias y anexos en el documento definitivo.',
                   'PC, procesador de textos'),
                  ('4',
                   'Revisión normativa',
                   'Verificar ortografía, referencias, márgenes, paginación e índices.',
                   'Normativa IUTECP'),
                  ('5',
                   'Preparación de entrega',
                   'Organizar la versión digital e impresa para su consignación institucional.',
                   'Pendrive, impresora, carpeta')]}]

# ================================================================
#  EJECUCIÓN GENERAL
# ================================================================

if __name__ == "__main__":
    print("🚀 Iniciando motor de automatización de cronogramas para Procura (Keidy)...")
    for idx, reporte in enumerate(reportes_pasantias):
        print(f"\n[Compilando {idx+1}/{len(reportes_pasantias)}] Semana {reporte['num_semana']}...")
        generar_documento_semana(reporte)
    print("\n✨ ¡Proceso culminado con éxito! Se han generado los 10 archivos independientes en Word y PDF.")
